"""
账单管理系统
企业级发票、报价单、付款和客户管理系统
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import json
import os
import re

# 文件拖拽支持
try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DRAG_DROP_AVAILABLE = True
except ImportError:
    DRAG_DROP_AVAILABLE = False

# Word文档和PDF导出支持
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# PyMuPDF支持（用于PDF盖章）
try:
    import fitz  # PyMuPDF
    MUPDF_AVAILABLE = True
except ImportError:
    MUPDF_AVAILABLE = False

# OCR支持
try:
    import pytesseract
    from PIL import Image
    # 设置Tesseract OCR路径
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# AI模型支持
try:
    import requests
    import base64
    import json
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

# 日期处理
try:
    from dateutil.relativedelta import relativedelta
    DATEUTIL_AVAILABLE = True
except ImportError:
    DATEUTIL_AVAILABLE = False


class InvoiceManager:
    """账单管理系统主类"""

    def __init__(self):
        # 使用 TkinterDnD 支持文件拖拽
        if DRAG_DROP_AVAILABLE:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()
        self.root.title("账单管理系统")

        # 数据存储
        self.monthly_invoices = []  # 月度账单列表
        self.quotations = []        # 报价单列表
        self.payments = []           # 付款记录列表
        self.customers = []          # 客户列表

        # 系统设置属性（初始化默认值）
        self.system_name = ''
        self.company_name = ''
        self.company_address = ''
        self.company_phone = ''
        self.bank_account = ''
        self.bank_name = ''
        self.paper_size = 'A4'
        self.logo_path = ''
        self.stamp_path = ''         # 公章图片路径
        self.quoter = ''             # 报价人
        self.handler = ''            # 经办人
        
        # AI模型设置属性（初始化默认值）
        self.doubao_model = '豆包-Seed-2.0-pro'
        self.doubao_model_name = 'doubao-seed-2.0-pro-260215'
        self.doubao_base_url = 'https://ark.cn-beijing.volces.com/api/v3'
        self.doubao_api_key = ''
        
        # 模型预设配置
        self.MODEL_PRESETS = {
            "doubao-lite": ("豆包-Seed-2.0-lite", "https://ark.cn-beijing.volces.com/api/v3", "doubao-seed-2-0-lite-260215"),
            "doubao-pro": ("豆包-Seed-2.0-pro", "https://ark.cn-beijing.volces.com/api/v3", "doubao-seed-2-0-pro-260215"),
        }

        # 数据文件保存路径（与脚本同一目录）
        self.data_dir = os.path.dirname(os.path.abspath(__file__))

        # 加载数据
        self.load_data()
        self.load_settings()

        # 创建界面
        self.create_menu()
        self.create_main_interface()

        # 窗口最大化
        self.root.state('zoomed')

        # 应用保存的系统名称到窗口标题
        self.update_title()

    def load_data(self):
        """加载所有数据文件"""
        try:
            monthly_invoices_path = os.path.join(self.data_dir, 'monthly_invoices.json')
            quotations_path = os.path.join(self.data_dir, 'quotations.json')
            customers_path = os.path.join(self.data_dir, 'customers.json')
            payments_path = os.path.join(self.data_dir, 'payments.json')

            if os.path.exists(monthly_invoices_path):
                with open(monthly_invoices_path, 'r', encoding='utf-8') as f:
                    self.monthly_invoices = json.load(f)
            if os.path.exists(quotations_path):
                with open(quotations_path, 'r', encoding='utf-8') as f:
                    self.quotations = json.load(f)
            if os.path.exists(customers_path):
                with open(customers_path, 'r', encoding='utf-8') as f:
                    self.customers = json.load(f)
            if os.path.exists(payments_path):
                with open(payments_path, 'r', encoding='utf-8') as f:
                    self.payments = json.load(f)
        except Exception as e:
            messagebox.showerror("错误", f"加载数据失败: {str(e)}")

    def load_settings(self):
        """加载系统设置"""
        try:
            settings_path = os.path.join(self.data_dir, 'settings.txt')
            if os.path.exists(settings_path):
                with open(settings_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if ':' in line:
                            key, value = line.strip().split(':', 1)
                            if key == '系统名称':
                                self.system_name = value
                            elif key == '公司名称':
                                self.company_name = value
                            elif key == '公司地址':
                                self.company_address = value
                            elif key == '联系电话':
                                self.company_phone = value
                            elif key == '银行账号':
                                self.bank_account = value
                            elif key == '开户银行':
                                self.bank_name = value
                            elif key == '纸张大小':
                                self.paper_size = value
                            elif key == 'LOGO路径':
                                self.logo_path = value
                            elif key == '公章路径':
                                self.stamp_path = value
                            elif key == '报价人':
                                self.quoter = value
                            elif key == '经办人':
                                self.handler = value
                            elif key == 'AI模型':
                                self.doubao_model = value
                            elif key == '模型名称':
                                self.doubao_model_name = value
                            elif key == 'BaseURL':
                                self.doubao_base_url = value
                            elif key == 'API Key':
                                self.doubao_api_key = value
        except Exception as e:
            print(f"加载设置失败: {e}")

    def save_data(self):
        """保存数据"""
        try:
            monthly_invoices_path = os.path.join(self.data_dir, 'monthly_invoices.json')
            quotations_path = os.path.join(self.data_dir, 'quotations.json')

            with open(monthly_invoices_path, 'w', encoding='utf-8') as f:
                json.dump(self.monthly_invoices, f, ensure_ascii=False, indent=2)
            with open(quotations_path, 'w', encoding='utf-8') as f:
                json.dump(self.quotations, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("错误", f"保存数据失败: {str(e)}")

    def save_customers(self):
        """保存客户数据"""
        try:
            customers_path = os.path.join(self.data_dir, 'customers.json')
            with open(customers_path, 'w', encoding='utf-8') as f:
                json.dump(self.customers, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("错误", f"保存客户数据失败: {str(e)}")

    def save_payments(self):
        """保存付款记录"""
        try:
            payments_path = os.path.join(self.data_dir, 'payments.json')
            with open(payments_path, 'w', encoding='utf-8') as f:
                json.dump(self.payments, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("错误", f"保存付款记录失败: {str(e)}")

    def save_settings_to_file(self):
        """保存设置到文件"""
        try:
            settings_path = os.path.join(self.data_dir, 'settings.txt')
            with open(settings_path, 'w', encoding='utf-8') as f:
                f.write(f"系统名称:{getattr(self, 'system_name', '')}\n")
                f.write(f"公司名称:{self.company_name}\n")
                f.write(f"公司地址:{getattr(self, 'company_address', '')}\n")
                f.write(f"联系电话:{getattr(self, 'company_phone', '')}\n")
                f.write(f"银行账号:{getattr(self, 'bank_account', '')}\n")
                f.write(f"开户银行:{getattr(self, 'bank_name', '')}\n")
                f.write(f"纸张大小:{getattr(self, 'paper_size', 'A4')}\n")
                f.write(f"LOGO路径:{getattr(self, 'logo_path', '')}\n")
                f.write(f"公章路径:{getattr(self, 'stamp_path', '')}\n")
                f.write(f"报价人:{getattr(self, 'quoter', '')}\n")
                f.write(f"经办人:{getattr(self, 'handler', '')}\n")
                f.write(f"AI模型:{getattr(self, 'doubao_model', '豆包2.0LITE')}\n")
                f.write(f"模型名称:{getattr(self, 'doubao_model_name', '')}\n")
                f.write(f"BaseURL:{getattr(self, 'doubao_base_url', 'https://api.doubao.com')}\n")
                f.write(f"API Key:{getattr(self, 'doubao_api_key', '')}\n")
        except Exception as e:
            print(f"保存设置失败: {e}")

    def create_menu(self):
        """创建菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # 对账单菜单
        monthly_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="对账单", menu=monthly_menu)
        monthly_menu.add_command(label="创建对账单", command=self.create_new_monthly_invoice)
        monthly_menu.add_command(label="查看对账单列表", command=self.show_monthly_invoices)

        # 报价单菜单
        quotation_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="报价管理", menu=quotation_menu)
        quotation_menu.add_command(label="创建新报价", command=self.create_new_quotation)
        quotation_menu.add_command(label="查看报价列表", command=self.show_quotations)

        # 客户管理菜单
        customer_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="客户管理", menu=customer_menu)
        customer_menu.add_command(label="客户列表", command=self.show_customers)

        # 付款管理菜单
        payment_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="付款管理", menu=payment_menu)
        payment_menu.add_command(label="付款列表", command=self.show_payments)
        payment_menu.add_command(label="添加付款记录", command=self.add_payment)
        payment_menu.add_separator()
        payment_menu.add_command(label="导出付款记录", command=self.export_payments)

        # 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="帮助", menu=help_menu)
        help_menu.add_command(label="关于", command=self.show_about)
        help_menu.add_command(label="使用帮助", command=self.show_help)


    def _add_stamp_to_pdf(self, pdf_path, output_path, stamp_path):
        """使用PyMuPDF在PDF上添加公章
        
        Args:
            pdf_path: 源PDF文件路径
            output_path: 输出PDF文件路径
            stamp_path: 公章图片路径
        
        公章将加盖在经办人签名位置，大小为4CM*4CM
        """
        import fitz  # PyMuPDF
        from reportlab.lib.units import cm
        from PIL import Image
        import numpy as np
        import tempfile
        import shutil
        
        # 打开PDF
        doc = fitz.open(pdf_path)
        page = doc[0]  # 获取第一页
        
        # 获取页面尺寸
        page_width = page.rect.width
        page_height = page.rect.height
        
        # ========== 公章位置计算（加盖在经办人签名位置）==========
        # 对账单是横版A4纸（29.7cm x 21cm）
        # 经办人在右下角，位置参照PDF中的sign_info_table布局
        stamp_size = 4 * cm
        
        # 经办人签名区域位置（横版A4右下角）
        # PDF布局：colWidths=[12*cm, 10*cm]，经办人在右侧10cm区域
        # 页边距1cm，所以经办人x坐标约在 29.7 - 1 - 5 = 23.7cm 处
        sign_x_approx = 23 * cm  # 经办人x坐标（右下角）
        sign_y_approx = 3 * cm   # 经办人y坐标（距底部）- 再下移1厘米
        
        # overlap控制印章与文字重合程度：0.3表示印章30%压在文字上
        overlap = 0.3
        stamp_y = sign_y_approx + (stamp_size * overlap)
        stamp_x = sign_x_approx - (stamp_size / 2)
        
        # 确保印章在页面范围内
        if stamp_x < 1 * cm:
            stamp_x = 1 * cm
        if stamp_x + stamp_size > page_width - 1 * cm:
            stamp_x = page_width - stamp_size - 1 * cm
        if stamp_y < 1 * cm:
            stamp_y = 1 * cm
        if stamp_y + stamp_size > page_height - 1 * cm:
            stamp_y = page_height - stamp_size - 1 * cm
        
        # ========== 关键修复：处理透明公章图片 ==========
        # 参照报价单导出PDF的盖章逻辑，使用Alpha通道处理透明度
        stamp_img = Image.open(stamp_path).convert('RGBA')
        
        # 应用透明度：0.7 表示70%不透明度
        alpha = stamp_img.split()[3]  # 获取Alpha通道
        alpha = alpha.point(lambda p: int(p * 0.7))  # 70%强度
        stamp_img.putalpha(alpha)
        
        # 保存透明图片到临时PNG文件
        transparent_temp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        transparent_path = transparent_temp.name
        transparent_temp.close()
        stamp_img.save(transparent_path, 'PNG')
        
        # 使用PyMuPDF添加公章
        # PyMuPDF使用左下角为原点，y轴向上
        mupdf_stamp_x = stamp_x
        mupdf_stamp_y = page_height - stamp_y - stamp_size  # 转换y坐标
        
        # 创建图片矩形区域
        stamp_rect = fitz.Rect(
            mupdf_stamp_x, 
            mupdf_stamp_y, 
            mupdf_stamp_x + stamp_size, 
            mupdf_stamp_y + stamp_size
        )
        
        # 使用insert_image方法添加图片
        page.insert_image(stamp_rect, filename=transparent_path)
        
        # 参照报价单的保存方式：先保存到临时文件，再替换原文件
        doc.save(output_path + '.tmp')
        doc.close()
        
        # 替换原文件
        shutil.move(output_path + '.tmp', output_path)
        
        # 删除临时透明图片
        try:
            os.remove(transparent_path)
        except:
            pass
        
        print(f"✓ PyMuPDF公章添加成功: 位置=({stamp_x/cm:.2f}cm, {stamp_y/cm:.2f}cm), 大小={stamp_size/cm:.1f}cm×{stamp_size/cm:.1f}cm")

    def update_title(self):
        """更新系统标题"""
        system_name = getattr(self, 'system_name', '')
        if not system_name:
            system_name = "账单管理系统"
        self.root.title(system_name)
        if hasattr(self, 'title_label'):
            self.title_label.config(text=system_name)

    def create_main_interface(self):
        """创建主界面"""
        # 标题栏
        title_frame = tk.Frame(self.root, bg='#2c3e50', height=80)
        title_frame.pack(fill=tk.X)

        self.title_label = tk.Label(
            title_frame,
            text=getattr(self, 'system_name', '') or "账单管理系统",
            font=("Microsoft YaHei", 24, "bold"),
            bg='#2c3e50',
            fg='white'
        )
        self.title_label.pack(pady=20)

        # 内容区域容器
        self.content_container = tk.Frame(self.root, bg='#f5f5f5')
        self.content_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # 左侧区域 - 快速操作
        self.left_frame = tk.Frame(self.content_container, width=300, bg='#ffffff')
        self.left_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 20))

        # 主内容区域 - 默认显示欢迎页面
        self.main_content_frame = tk.Frame(self.content_container, bg='#ffffff')
        self.main_content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.show_welcome_page()

        tk.Label(self.left_frame, text="快速操作", font=("Microsoft YaHei", 14, "bold"), bg='#ffffff').pack(anchor=tk.W, pady=(0, 10))

        btn_style = {'font': ("Microsoft YaHei", 11), 'width': 20, 'height': 2}
        tk.Button(self.left_frame, text="创建对账单", command=self.create_new_monthly_invoice, **btn_style, bg='#3498db', fg='white').pack(pady=5)
        tk.Button(self.left_frame, text="创建新报价", command=self.create_new_quotation, **btn_style, bg='#9b59b6', fg='white').pack(pady=5)
        tk.Button(self.left_frame, text="查看对账单列表", command=self.show_monthly_invoices, **btn_style, bg='#27ae60', fg='white').pack(pady=5)
        tk.Button(self.left_frame, text="查看报价列表", command=self.show_quotations, **btn_style, bg='#e67e22', fg='white').pack(pady=5)

        # 统计信息
        stats_frame = tk.LabelFrame(self.left_frame, text="统计信息", font=("Microsoft YaHei", 12), padx=10, pady=10, bg='#ffffff')
        stats_frame.pack(fill=tk.X, pady=(30, 0))
        self.stats_label = tk.Label(stats_frame, text=f"对账单数: {len(self.monthly_invoices)}\n报价单数: {len(self.quotations)}",
                                    font=("Microsoft YaHei", 10), justify=tk.LEFT, bg='#ffffff')
        self.stats_label.pack()

        # 月付款信息按钮（彩色样式）
        tk.Button(self.left_frame, text="📅 月付款信息", command=self.show_payments_page,
                  font=("Microsoft YaHei", 12, "bold"), width=20, height=2,
                  bg='#FF6B6B', fg='white', relief=tk.RAISED, bd=3,
                  activebackground='#FF4757', activeforeground='white').pack(pady=8)

        # 系统设置按钮
        tk.Button(self.left_frame, text="系统设置", command=self.show_system_settings, **btn_style, bg='#34495e', fg='white').pack(side=tk.BOTTOM, pady=(20, 0))

        # 状态栏
        status_frame = tk.Frame(self.root, relief=tk.SUNKEN)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X)
        self.status_label = tk.Label(status_frame, text=f"就绪 | 当前时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", anchor=tk.W)
        self.status_label.pack(side=tk.LEFT, padx=5)

    def show_welcome_page(self):
        """显示欢迎页面"""
        for widget in self.main_content_frame.winfo_children():
            widget.destroy()
        system_name = getattr(self, 'system_name', '') or "账单管理系统"
        welcome_label = tk.Label(
            self.main_content_frame,
            text=f"欢迎使用{system_name}",
            font=("Microsoft YaHei", 24, "bold"),
            bg='#ffffff',
            fg='#2c3e50'
        )
        welcome_label.pack(pady=50)

    def clear_main_content(self):
        """清空主内容区域"""
        for widget in self.main_content_frame.winfo_children():
            widget.destroy()

    # -------------------- 对账单相关 --------------------
    def create_new_monthly_invoice(self, insert_index=None):
        """创建对账单 - 在主窗口内打开
        insert_index: 如果指定，在列表的指定位置插入新账单
        """
        self.status_label.config(text="正在创建对账单...")
        self.clear_main_content()

        form_frame = tk.Frame(self.main_content_frame, padx=20, pady=15)
        form_frame.pack(fill=tk.BOTH, expand=True)

        # 返回按钮
        back_frame = tk.Frame(form_frame)
        back_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Button(back_frame, text="← 返回主页", command=self.show_welcome_page,
                  font=("Microsoft YaHei", 10), bg='#95a5a6', fg='white',
                  relief=tk.FLAT, padx=15, pady=5).pack(side=tk.LEFT)

        # 标题行
        title_frame = tk.Frame(form_frame, bg='#f5f5f5', height=60)
        title_frame.pack(fill=tk.X)

        info_row1 = tk.Frame(title_frame, bg='#f5f5f5')
        info_row1.pack(fill=tk.X, pady=15)
        center_frame = tk.Frame(info_row1, bg='#f5f5f5')
        center_frame.pack(side=tk.TOP, anchor='center')

        year_month_var = tk.StringVar()
        year_month_combo = ttk.Combobox(center_frame, textvariable=year_month_var, width=12, state='readonly', font=("Microsoft YaHei", 22))
        months = []
        if DATEUTIL_AVAILABLE:
            for i in range(12):
                month_date = datetime.now().replace(day=1) - relativedelta(months=i)
                months.append(month_date.strftime("%Y年%m月"))
        else:
            for i in range(12):
                year = datetime.now().year
                month = datetime.now().month - i
                if month <= 0:
                    month += 12
                    year -= 1
                months.append(f"{year}年{month:02d}月")
        year_month_combo['values'] = months
        year_month_combo.current(0)
        year_month_combo.pack(side=tk.LEFT, padx=(0, 10))

        header_label = tk.Label(center_frame, text="对账单", font=("Microsoft YaHei", 22, "bold"), fg='#e74c3c', bg='#f5f5f5')
        header_label.pack(side=tk.LEFT)

        # 左右布局
        content_frame = tk.Frame(form_frame)
        content_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        left_panel = tk.Frame(content_frame, width=600)
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        left_panel.pack_propagate(False)

        photo_panel = tk.Frame(content_frame, bg='#f8f9fa', width=240)
        photo_panel.pack(side=tk.RIGHT, fill=tk.BOTH)
        photo_panel.pack_propagate(False)

        # 图片分析区域
        tk.Label(photo_panel, text="📷 智能图片分析", font=("Microsoft YaHei", 12, "bold"),
                 bg='#f8f9fa', fg='#2c3e50').pack(pady=(10, 5))
        tk.Label(photo_panel, text="拖拽送货单图片自动识别", font=("Microsoft YaHei", 9),
                 bg='#f8f9fa', fg='#7f8c8d').pack(pady=(0, 5))
        
        # AI模型选择
        ai_model_frame = tk.Frame(photo_panel, bg='#f8f9fa')
        ai_model_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Label(ai_model_frame, text="AI模型:", font=("Microsoft YaHei", 9),
                 bg='#f8f9fa', fg='#2c3e50').pack(side=tk.LEFT)
        ai_model_var = tk.StringVar(value=getattr(self, 'doubao_model', '豆包2.0PRO'))
        ai_model_combo = ttk.Combobox(ai_model_frame, textvariable=ai_model_var, width=15, state='readonly')
        ai_model_combo['values'] = ['豆包2.0LITE', '豆包2.0PRO']
        # 默认选择PRO模型
        if ai_model_var.get() == '豆包2.0PRO':
            ai_model_combo.current(1)
        else:
            ai_model_combo.current(0)
        ai_model_combo.pack(side=tk.LEFT, padx=(5, 0))

        drop_border_frame = tk.Frame(photo_panel, bg='#3498db', padx=2, pady=2)
        drop_border_frame.pack(padx=15, pady=5)
        photo_drop_frame = tk.Frame(drop_border_frame, bg='#ecf0f1', relief=tk.FLAT)
        photo_drop_frame.pack(fill=tk.BOTH, expand=True)

        # A5比例: 148mm x 210mm = 1:1.42 (约)
        tk.Label(photo_drop_frame, text="🖼️", font=("Arial", 28), bg='#ecf0f1').pack(pady=(10, 3))
        photo_text = tk.Text(photo_drop_frame, width=18, height=2, font=("Microsoft YaHei", 9),
                             bg='#ecf0f1', fg='#2c3e50', wrap=tk.WORD, relief=tk.FLAT,
                             bd=0, highlightthickness=0, takefocus=0)
        photo_text.pack(padx=10, pady=3)
        photo_text.insert('1.0', '将送货单图片拖拽到此处\n或点击选择图片文件')
        photo_text.config(state='disabled', cursor='hand2')

        dropped_photo_path = [None]
        preview_label = [None]

        def handle_drop(event):
            """处理拖拽到图片区域的图片文件"""
            try:
                # tkinterdnd2 会传递文件路径，可能有多个文件（用空格分隔）
                data = event.data.strip()
                
                # 处理可能的大括号格式 {path1} {path2}
                import re
                # 匹配文件路径（可能包含空格但在大括号内或以空格分隔）
                paths = []
                # 先尝试匹配大括号包裹的路径
                brace_pattern = r'\{([^}]+)\}'
                matches = re.findall(brace_pattern, data)
                if matches:
                    paths = matches
                else:
                    # 否则按空格分割
                    paths = data.split()
                
                # 获取第一个图片文件
                file_path = None
                for p in paths:
                    p = p.strip().strip('"').strip("'")
                    if p.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp')):
                        file_path = p
                        break
                
                if not file_path:
                    messagebox.showwarning("警告", "请拖拽图片文件(jpg/png/bmp)")
                    return
                
                dropped_photo_path[0] = file_path
                photo_text.config(state='normal')
                photo_text.delete('1.0', tk.END)
                photo_text.insert('1.0', f'已选择:\n{os.path.basename(file_path)}')
                photo_text.config(state='disabled')
                
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(file_path)
                    # 检测图片方向，如果是竖图则旋转90度变为横图
                    width, height = img.size
                    if height > width:
                        img = img.rotate(90, expand=True)
                    # A5比例预览 148:210 (保持A5纸张比例)
                    display_height = 150
                    display_width = int(display_height * 210 / 148)
                    img.thumbnail((display_width, display_height))
                    photo_preview = ImageTk.PhotoImage(img)
                    if preview_label[0]:
                        preview_label[0].destroy()
                    preview_label[0] = tk.Label(photo_drop_frame, image=photo_preview, bg='#ecf0f1', relief=tk.SOLID, bd=1)
                    preview_label[0].image = photo_preview
                    preview_label[0].pack(pady=5)
                    # 双击打开图片
                    def open_with_default_app(event, filepath=dropped_photo_path[0]):
                        import subprocess
                        import platform
                        if platform.system() == 'Windows':
                            subprocess.Popen(['start', '', filepath], shell=True)
                        elif platform.system() == 'Darwin':
                            subprocess.Popen(['open', filepath])
                        else:
                            subprocess.Popen(['xdg-open', filepath])
                    preview_label[0].bind('<Double-Button-1>', open_with_default_app)
                    preview_label[0].config(cursor='hand2')
                except Exception as img_error:
                    print(f"图片预览失败: {img_error}")
            except Exception as e:
                messagebox.showerror("错误", f"处理图片失败: {str(e)}")

        def handle_click(event):
            """点击打开文件选择对话框"""
            file_path = filedialog.askopenfilename(
                title="选择送货单图片",
                filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif *.webp"), ("所有文件", "*.*")]
            )
            if file_path:
                class FakeEvent:
                    def __init__(self, data):
                        self.data = data
                handle_drop(FakeEvent(file_path))

        # 使用 tkinterdnd2 的拖拽功能
        if DRAG_DROP_AVAILABLE:
            # 注册为拖放目标
            photo_drop_frame.drop_target_register(DND_FILES)
            drop_border_frame.drop_target_register(DND_FILES)
            photo_text.drop_target_register(DND_FILES)
            
            # 绑定拖拽事件
            photo_drop_frame.dnd_bind('<<Drop>>', handle_drop)
            drop_border_frame.dnd_bind('<<Drop>>', handle_drop)
            photo_text.dnd_bind('<<Drop>>', handle_drop)
        else:
            # 备用：使用点击选择
            messagebox.showinfo("提示", "拖拽功能需要安装 tkinterdnd2 库\n您可以点击图片区域选择文件")

        # 点击事件用于选择文件
        for widget in [photo_drop_frame, drop_border_frame, photo_text]:
            widget.bind('<Button-1>', handle_click)

        def analyze_photo():
            """分析送货单图片并直接填充到当前表单"""
            if not dropped_photo_path[0]:
                self.status_label.config(text="⚠️ 请先拖拽或选择一张送货单图片")
                return
            
            # 获取当前选择的AI模型
            selected_model = ai_model_var.get()
            self.status_label.config(text=f"正在使用{selected_model}分析图片...")
            
            # 临时保存选择的模型到系统设置
            original_model = getattr(self, 'doubao_model', '豆包2.0LITE')
            self.doubao_model = selected_model
            
            try:
                analyzed_data = self._analyze_delivery_photo(dropped_photo_path[0])
                if analyzed_data and analyzed_data.get('items'):
                    # 直接填充到当前表单
                    detected_customer = analyzed_data.get('客户名称', '')
                    if detected_customer:
                        # 设置客户名称
                        customer_var.set(detected_customer)
                    
                    # 清空现有表格数据
                    for item in tree.get_children():
                        tree.delete(item)
                    
                    # 填充分析结果
                    items = analyzed_data.get('items', [])
                    for idx, item_data in enumerate(items, 1):
                        tree.insert('', tk.END, values=(
                            idx,
                            item_data.get('订单号', ''),
                            item_data.get('指令号', ''),
                            item_data.get('送货日期', ''),
                            item_data.get('送货单号', ''),
                            item_data.get('品名', ''),
                            item_data.get('规格', ''),
                            item_data.get('颜色', ''),
                            item_data.get('单位', ''),
                            item_data.get('数量', ''),
                            item_data.get('单价', ''),
                            item_data.get('染费', ''),
                            item_data.get('金额', '')
                        ))
                    
                    # 添加空行以便继续输入
                    for i in range(3):
                        tree.insert('', tk.END, values=(len(items) + i + 1, '', '', '', '', '', '', '', '', '', '', '', ''))
                    
                    # 更新合计金额
                    self._update_total_from_tree(tree, total_label)
                    
                    self.status_label.config(text=f"✓ 分析完成，已识别 {len(items)} 条记录")
                elif analyzed_data and not analyzed_data.get('items'):
                    self.status_label.config(text="💡 未能识别出送货明细，请手动输入")
                else:
                    self.status_label.config(text="⚠️ 未能识别出有效信息，请确保图片清晰")
            except Exception as e:
                self.status_label.config(text=f"✗ 分析出错: {str(e)}")
            finally:
                # 恢复原始模型设置
                self.doubao_model = original_model
                self.status_label.config(text="就绪")

        analyze_btn = tk.Button(photo_panel, text="🔍 开始分析", command=analyze_photo,
                                width=18, height=2, bg='#e74c3c', fg='white',
                                font=("Microsoft YaHei", 11, "bold"), relief=tk.FLAT, cursor='hand2')
        analyze_btn.pack(pady=15)
        tk.Label(photo_panel, text="提示: 支持JPG/PNG/BMP格式", font=("Microsoft YaHei", 8),
                 bg='#f8f9fa', fg='#95a5a6').pack(pady=(5, 10))

        # 左侧客户信息
        info_row2 = tk.Frame(left_panel)
        info_row2.pack(fill=tk.X, pady=(0, 5))

        left_frame_cust = tk.Frame(info_row2)
        left_frame_cust.pack(side=tk.LEFT)

        tk.Label(left_frame_cust, text="客户名称:", font=("Microsoft YaHei", 11)).pack(side=tk.LEFT)
        customer_var = tk.StringVar()
        customer_combo = ttk.Combobox(left_frame_cust, textvariable=customer_var, width=25, state='readonly')

        def refresh_customer_combo():
            try:
                customers_path = os.path.join(self.data_dir, 'customers.json')
                if os.path.exists(customers_path):
                    with open(customers_path, 'r', encoding='utf-8') as f:
                        self.customers = json.load(f)
            except:
                pass
            customer_names = [c.get('name', '') for c in self.customers if c.get('name')]
            customer_combo['values'] = customer_names if customer_names else ['暂无客户']

        refresh_customer_combo()
        customer_combo.pack(side=tk.LEFT, padx=(5, 0))

        def add_quick_customer():
            quick_window = tk.Toplevel(self.root)
            quick_window.title("添加客户")
            quick_window.geometry("350x150")
            quick_window.transient(self.root)
            quick_window.grab_set()
            tk.Label(quick_window, text="客户名称:", font=("Microsoft YaHei", 11)).pack(pady=10)
            quick_name_entry = tk.Entry(quick_window, width=30, font=("Microsoft YaHei", 11))
            quick_name_entry.pack(pady=5)
            quick_name_entry.focus()
            def save_quick_customer():
                name = quick_name_entry.get().strip()
                if not name:
                    messagebox.showwarning("警告", "请输入客户名称")
                    return
                self.customers.append({'name': name, 'contact': '', 'phone': '', 'address': '', 'remark': ''})
                self.save_customers()
                refresh_customer_combo()
                customer_var.set(name)
                quick_window.destroy()
            btn_frame = tk.Frame(quick_window)
            btn_frame.pack(pady=10)
            tk.Button(btn_frame, text="保存并选择", command=save_quick_customer, width=12, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="取消", command=quick_window.destroy, width=12, bg='#95a5a6').pack(side=tk.LEFT, padx=5)
            quick_name_entry.bind('<Return>', lambda e: save_quick_customer())


        right_frame_date = tk.Frame(info_row2)
        right_frame_date.pack(side=tk.RIGHT)
        tk.Label(right_frame_date, text="制表日期:", font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=(30, 5))
        current_date = datetime.now().strftime("%Y年%m月%d日")
        date_entry = tk.Entry(right_frame_date, width=18)
        date_entry.insert(0, current_date)
        date_entry.pack(side=tk.LEFT)

        # 商品明细表格
        columns = ('序号', '订单号', '指令号', '送货日期', '送货单号', '品名', '规格', '颜色', '单位', '数量', '单价', '染费', '金额')
        tree = ttk.Treeview(left_panel, columns=columns, show='headings', height=12)

        tree.heading('序号', text='序号')
        tree.heading('订单号', text='订单号')
        tree.heading('指令号', text='指令号')
        tree.heading('送货日期', text='送货日期')
        tree.heading('送货单号', text='送货单号')
        tree.heading('品名', text='品名')
        tree.heading('规格', text='规格')
        tree.heading('颜色', text='颜色')
        tree.heading('单位', text='单位')
        tree.heading('数量', text='数量')
        tree.heading('单价', text='单价(¥)')
        tree.heading('染费', text='染费(¥)')
        tree.heading('金额', text='金额(¥)')

        tree.column('序号', width=50, anchor='center')
        tree.column('订单号', width=80, anchor='center')
        tree.column('指令号', width=80, anchor='center')
        tree.column('送货日期', width=85, anchor='center')
        tree.column('送货单号', width=85, anchor='center')
        tree.column('品名', width=80, anchor='center')
        tree.column('规格', width=70, anchor='center')
        tree.column('颜色', width=60, anchor='center')
        tree.column('单位', width=50, anchor='center')
        tree.column('数量', width=50, anchor='center')
        tree.column('单价', width=70, anchor='center')
        tree.column('染费', width=70, anchor='center')
        tree.column('金额', width=80, anchor='center')

        tree.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(left_panel, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)

        # 右键菜单
        context_menu = tk.Menu(tree, tearoff=0)
        def edit_selected():
            selection = tree.selection()
            if selection:
                item_values = tree.item(selection[0])['values']
                self._edit_item_dialog(item_values, tree, selection[0], self._update_total_from_tree, total_label)
        def delete_selected():
            selection = tree.selection()
            if selection and messagebox.askyesno("确认", "确定要删除这条记录吗?"):
                parent = tree.parent(selection[0])
                tree.delete(selection[0])
                # 重新编号
                for i, item_id in enumerate(tree.get_children(parent), 1):
                    values = list(tree.item(item_id)['values'])
                    values[0] = i
                    tree.item(item_id, values=values)
                self._update_total_from_tree(tree, total_label)
        def insert_row_above():
            selection = tree.selection()
            if selection:
                parent = tree.parent(selection[0])
                siblings = tree.get_children(parent)
                idx = siblings.index(selection[0])
                new_values = list(tree.item(selection[0])['values'])
                # 清空除序号外的所有字段
                for i in range(len(new_values)):
                    if i != 0:  # 保留序号位置，后面会重新编号
                        new_values[i] = ''
                tree.insert(parent, idx, values=new_values)
                # 重新编号
                for i, item_id in enumerate(tree.get_children(parent), 1):
                    values = list(tree.item(item_id)['values'])
                    values[0] = i
                    tree.item(item_id, values=values)
                self._update_total_from_tree(tree, total_label)
        def insert_row_below():
            selection = tree.selection()
            if selection:
                parent = tree.parent(selection[0])
                siblings = tree.get_children(parent)
                idx = siblings.index(selection[0]) + 1
                new_values = list(tree.item(selection[0])['values'])
                # 清空除序号外的所有字段
                for i in range(len(new_values)):
                    if i != 0:  # 保留序号位置，后面会重新编号
                        new_values[i] = ''
                tree.insert(parent, idx, values=new_values)
                # 重新编号
                for i, item_id in enumerate(tree.get_children(parent), 1):
                    values = list(tree.item(item_id)['values'])
                    values[0] = i
                    tree.item(item_id, values=values)
                self._update_total_from_tree(tree, total_label)
        context_menu.add_command(label="编辑", command=edit_selected)
        context_menu.add_command(label="删除", command=delete_selected)
        context_menu.add_separator()
        context_menu.add_command(label="向上插入一行", command=insert_row_above)
        context_menu.add_command(label="向下插入一行", command=insert_row_below)

        def on_double_click(event):
            item_id = tree.selection()
            if item_id:
                item_values = tree.item(item_id[0])['values']
                self._edit_item_dialog(item_values, tree, item_id[0], self._update_total_from_tree, total_label)

        def show_context_menu(event):
            item_id = tree.identify_row(event.y)
            if item_id:
                tree.selection_set(item_id)
                context_menu.post(event.x_root, event.y_root)

        tree.bind("<Double-Button-1>", on_double_click)
        tree.bind("<Button-3>", show_context_menu)

        # 初始6行空数据
        for i in range(6):
            tree.insert('', tk.END, values=(i+1, '', '', '', '', '', '', '', '', '', '', ''))

        # 备注区域
        remark_frame = tk.Frame(left_panel)
        remark_frame.pack(fill=tk.X, pady=10)
        tk.Label(remark_frame, text="备注:", font=("Microsoft YaHei", 11)).pack(side=tk.LEFT)
        remark_entry = tk.Entry(remark_frame, width=60)
        remark_entry.pack(side=tk.LEFT, padx=5)

        # 底部合计
        bottom_frame = tk.Frame(left_panel)
        bottom_frame.pack(fill=tk.X, pady=10)
        total_frame = tk.Frame(bottom_frame)
        total_frame.pack(side=tk.RIGHT)
        tk.Label(total_frame, text="金额合计(¥): ", font=("Microsoft YaHei", 12, "bold")).pack(side=tk.LEFT)
        total_label = tk.Label(total_frame, text="0.00", font=("Microsoft YaHei", 14, "bold"), fg='red')
        total_label.pack(side=tk.LEFT)

        # 账户信息占位
        account_frame = tk.Frame(bottom_frame)
        account_frame.pack(side=tk.LEFT)
        tk.Label(account_frame, text="开户行:", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)
        tk.Label(account_frame, text="账号:", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)
        tk.Label(account_frame, text="户名:", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)

        # 按钮区域
        btn_frame = tk.Frame(left_panel)
        btn_frame.pack(pady=15)

        def save_invoice():
            invoice_data = {
                'year_month': year_month_var.get(),
                'customer': customer_var.get(),
                'items': [],
                'total': 0.0,
                'remark': remark_entry.get(),
                'date': date_entry.get(),
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            for item in tree.get_children():
                values = tree.item(item)['values']
                if values[1] or values[2]:
                    invoice_data['items'].append({
                        '序号': values[0],
                        '订单号': values[1],
                        '指令号': values[2],
                        '送货日期': values[3],
                        '送货单号': values[4],
                        '品名': values[5],
                        '规格': values[6],
                        '颜色': values[7],
                        '单位': values[8],
                        '数量': values[9],
                        '单价': values[10],
                        '染费': values[11],
                        '金额': values[12]
                    })
                    try:
                        invoice_data['total'] += float(values[12]) if values[12] else 0
                    except:
                        pass
            self.monthly_invoices.append(invoice_data)
            self.save_data()
            
            # 同步付款记录到月付款信息列表 - 按年份+客户+月份更新
            year_month_str = year_month_var.get()  # 格式: "2026年04月"
            customer = customer_var.get()
            total_amount = invoice_data.get('total', 0)
            
            if year_month_str and customer and total_amount > 0:
                try:
                    # 解析年月 (格式: "2026年04月")
                    import re
                    match = re.match(r'(\d{4})年(\d{2})月', year_month_str)
                    if match:
                        year = int(match.group(1))
                        month = int(match.group(2))
                        payment_date = f"{year}-{month:02d}-01"  # 转换为日期格式
                        
                        # 按年份+客户查找付款记录（用于年月+客户交叉定位）
                        key = (year, customer)
                        
                        # 查找是否存在该年份+客户的付款记录
                        existing_payment = None
                        for p in self.payments:
                            p_date = p.get('payment_date', '')
                            p_customer = p.get('customer', '')
                            p_year = None
                            p_month = None
                            try:
                                if '-' in p_date:
                                    p_year = int(p_date.split('-')[0])
                                    p_month = int(p_date.split('-')[1])
                            except:
                                pass
                            
                            if p_customer == customer and p_year == year:
                                existing_payment = p
                                # 更新该月份对应的金额
                                if p_month == month:
                                    p['amount'] = total_amount
                                    p['payment_date'] = payment_date
                                break
                        
                        if not existing_payment:
                            # 创建新付款记录
                            new_payment = {
                                'payment_no': f"INV-{year}-{month:02d}-{customer[:3]}",
                                'customer': customer,
                                'amount': total_amount,
                                'payment_date': payment_date,
                                'payment_method': '对账单转入',
                                'status': '已付款',
                                'remark': f'对账单自动同步'
                            }
                            self.payments.append(new_payment)
                        
                        self.save_payments()
                except Exception as e:
                    print(f"同步付款记录失败: {e}")
            
            self.update_stats()
            messagebox.showinfo("成功", "对账单保存成功")
            self.show_welcome_page()
            self.status_label.config(text="对账单已保存")

        tk.Button(btn_frame, text="保存账单", command=save_invoice, width=15, height=2, bg='#27ae60', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=self.show_welcome_page, width=15, height=2, bg='#e74c3c', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)

    def _edit_item_dialog(self, item_values, tree, item_id, update_total_callback, total_label):
        """编辑明细项对话框"""
        edit_window = tk.Toplevel(self.root)
        edit_window.title("送货单编辑")
        edit_window.geometry("650x400")
        
        # 窗口居中显示
        edit_window.update_idletasks()
        screen_width = edit_window.winfo_screenwidth()
        screen_height = edit_window.winfo_screenheight()
        x = (screen_width - 650) // 2
        y = (screen_height - 400) // 2
        edit_window.geometry(f"650x400+{x}+{y}")
        
        edit_frame = tk.Frame(edit_window, padx=20, pady=20)
        edit_frame.pack(fill=tk.BOTH, expand=True)

        # 订单号
        tk.Label(edit_frame, text="订单号:", width=10, anchor=tk.W).grid(row=0, column=0, sticky=tk.W, pady=5)
        order_entry = tk.Entry(edit_frame, width=25)
        order_entry.grid(row=0, column=1, pady=5, padx=(0, 15))
        order_entry.insert(0, item_values[1])

        tk.Label(edit_frame, text="指令号:", width=10, anchor=tk.W).grid(row=0, column=2, sticky=tk.W, pady=5)
        cmd_entry = tk.Entry(edit_frame, width=25)
        cmd_entry.grid(row=0, column=3, pady=5, padx=(0, 15))
        cmd_entry.insert(0, item_values[2])

        # 送货日期
        tk.Label(edit_frame, text="送货日期:", width=10, anchor=tk.W).grid(row=1, column=0, sticky=tk.W, pady=5)
        delivery_date_entry = tk.Entry(edit_frame, width=25)
        delivery_date_entry.grid(row=1, column=1, pady=5, padx=(0, 15))
        delivery_date_entry.insert(0, item_values[3])

        tk.Label(edit_frame, text="送货单号:", width=10, anchor=tk.W).grid(row=1, column=2, sticky=tk.W, pady=5)
        delivery_entry = tk.Entry(edit_frame, width=25)
        delivery_entry.grid(row=1, column=3, pady=5, padx=(0, 15))
        delivery_entry.insert(0, item_values[4])

        # 品名规格颜色
        tk.Label(edit_frame, text="品名:", width=10, anchor=tk.W).grid(row=2, column=0, sticky=tk.W, pady=5)
        name_entry = tk.Entry(edit_frame, width=25)
        name_entry.grid(row=2, column=1, pady=5, padx=(0, 15))
        name_entry.insert(0, item_values[5])

        tk.Label(edit_frame, text="规格:", width=10, anchor=tk.W).grid(row=2, column=2, sticky=tk.W, pady=5)
        spec_entry = tk.Entry(edit_frame, width=25)
        spec_entry.grid(row=2, column=3, pady=5, padx=(0, 15))
        spec_entry.insert(0, item_values[6])

        tk.Label(edit_frame, text="颜色:", width=10, anchor=tk.W).grid(row=3, column=0, sticky=tk.W, pady=5)
        color_entry = tk.Entry(edit_frame, width=25)
        color_entry.grid(row=3, column=1, pady=5, padx=(0, 15))
        color_entry.insert(0, item_values[7])

        # 单位数量
        tk.Label(edit_frame, text="单位:", width=10, anchor=tk.W).grid(row=3, column=2, sticky=tk.W, pady=5)
        unit_entry = tk.Entry(edit_frame, width=25)
        unit_entry.grid(row=3, column=3, pady=5, padx=(0, 15))
        unit_entry.insert(0, item_values[8])

        tk.Label(edit_frame, text="数量:", width=10, anchor=tk.W).grid(row=4, column=0, sticky=tk.W, pady=5)
        qty_entry = tk.Entry(edit_frame, width=25)
        qty_entry.grid(row=4, column=1, pady=5, padx=(0, 15))
        qty_entry.insert(0, item_values[9])

        # 单价染费金额
        tk.Label(edit_frame, text="单价:", width=10, anchor=tk.W).grid(row=4, column=2, sticky=tk.W, pady=5)
        price_entry = tk.Entry(edit_frame, width=25)
        price_entry.grid(row=4, column=3, pady=5, padx=(0, 15))
        price_val = item_values[10]
        if price_val:
            try:
                price_val = f"￥{float(price_val):.2f}"
            except:
                pass
        price_entry.insert(0, price_val)

        tk.Label(edit_frame, text="染费:", width=10, anchor=tk.W).grid(row=5, column=0, sticky=tk.W, pady=5)
        dye_fee_entry = tk.Entry(edit_frame, width=25, state='normal')
        dye_fee_entry.grid(row=5, column=1, pady=5, padx=(0, 15))
        dye_fee_entry.insert(0, item_values[11] if len(item_values) > 11 else '')
        dye_fee_entry.focus_set()
        dye_fee_entry.select_range(0, tk.END)

        tk.Label(edit_frame, text="金额:", width=10, anchor=tk.W).grid(row=6, column=0, sticky=tk.W, pady=5)
        amount_var = tk.StringVar()
        amount_entry = tk.Entry(edit_frame, width=25, textvariable=amount_var, state='readonly')
        amount_entry.grid(row=6, column=1, pady=5, padx=(0, 15))

        def calculate_amount(*args):
            try:
                qty = float(qty_entry.get())
                price_str = price_entry.get().replace('￥', '')
                price = float(price_str)
                dye_fee = float(dye_fee_entry.get()) if dye_fee_entry.get() else 0
                amount = qty * price + dye_fee
                amount_var.set(f"￥{amount:.2f}")
            except:
                amount_var.set("￥0.00")

        if len(item_values) > 11 and item_values[11]:
            try:
                amount_var.set(f"￥{float(item_values[11]):.2f}")
            except:
                amount_var.set("￥0.00")
        else:
            amount_var.set("￥0.00")

        qty_entry.bind('<KeyRelease>', calculate_amount)
        price_entry.bind('<KeyRelease>', calculate_amount)
        dye_fee_entry.bind('<KeyRelease>', calculate_amount)

        def save_edit():
            new_values = [item_values[0],
                          order_entry.get(),
                          cmd_entry.get(),
                          delivery_date_entry.get(),
                          delivery_entry.get(),
                          name_entry.get(),
                          spec_entry.get(),
                          color_entry.get(),
                          unit_entry.get(),
                          qty_entry.get(),
                          price_entry.get().replace('￥', ''),
                          dye_fee_entry.get(),
                          amount_var.get().replace('￥', '')]
            tree.item(item_id, values=new_values)
            if update_total_callback:
                update_total_callback(tree, total_label)
            edit_window.destroy()

        btn_frame = tk.Frame(edit_window)
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="保存", command=save_edit, width=15, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="取消", command=edit_window.destroy, width=15, bg='#95a5a6').pack(side=tk.LEFT, padx=10)

    def _update_total_from_tree(self, tree, total_label):
        """从Treeview更新合计金额"""
        total = 0.0
        for item_id in tree.get_children():
            values = tree.item(item_id)['values']
            if len(values) > 10:
                try:
                    amount_str = str(values[10]).replace('￥', '').replace(',', '')
                    if amount_str:
                        total += float(amount_str)
                except:
                    pass
        total_label.config(text=f"{total:.2f}")

    def _show_analyze_result_page(self, analyzed_data, photo_path=None):
        """显示图片分析结果 - 使用图片中的样式：左侧图片预览+按钮，右侧表单"""
        result_window = tk.Toplevel(self.root)
        result_window.title("送货单识别")
        result_window.geometry("1100x650")
        
        detected_customer = analyzed_data.get('客户名称', '')
        items = analyzed_data.get('items', [])
        is_modified = [False]
        
        # 主容器 - 左右布局
        main_container = tk.Frame(result_window)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # ==================== 左侧面板：图片预览 ====================
        left_panel = tk.Frame(main_container, bg='#ecf0f1', width=300)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=10)
        left_panel.pack_propagate(False)
        
        # 图片预览区域
        preview_frame = tk.Frame(left_panel, bg='white', relief=tk.SOLID, bd=1)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 5))
        
        preview_label = tk.Label(preview_frame, text="📷\n\n已识别送货单信息\n请确认并修改", 
                                  font=("Microsoft YaHei", 12), bg='white', fg='#7f8c8d',
                                  justify=tk.CENTER)
        preview_label.pack(fill=tk.BOTH, expand=True, pady=20)
        
        # 如果有图片路径，显示图片预览
        if photo_path and os.path.exists(photo_path):
            try:
                from PIL import Image, ImageTk
                img = Image.open(photo_path)
                width, height = img.size
                if height > width:
                    img = img.rotate(90, expand=True)
                # 按比例缩放图片
                display_height = 300
                display_width = int(display_height * width / height)
                img.thumbnail((display_width, display_height))
                photo_preview = ImageTk.PhotoImage(img)
                preview_label.config(image=photo_preview, text='', bg='white')
                preview_label.image = photo_preview
            except Exception as e:
                print(f"预览图片失败: {e}")
        
        # 左侧按钮区域
        btn_panel = tk.Frame(left_panel, bg='#ecf0f1')
        btn_panel.pack(pady=10, fill=tk.X)
        
        # 左侧按钮区域（只保留清空和重置）
        # 清空按钮
        clear_btn = tk.Button(btn_panel, text="🗑️ 清空", 
                              width=20, height=2, bg='#e67e22', fg='white',
                              font=("Microsoft YaHei", 11), relief=tk.FLAT, cursor='hand2',
                              command=lambda: self._clear_analyze_result(tree))
        clear_btn.pack(pady=5)
        
        # 重置按钮
        reset_btn = tk.Button(btn_panel, text="🔄 重置", 
                              width=20, height=2, bg='#95a5a6', fg='white',
                              font=("Microsoft YaHei", 11), relief=tk.FLAT, cursor='hand2',
                              command=lambda: self._reset_analyze_result(analyzed_data, tree))
        reset_btn.pack(pady=5)
        
        # ==================== 右侧面板：表单内容 ====================
        right_panel = tk.Frame(main_container)
        right_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10), pady=10)
        
        # 标题区域
        title_frame = tk.Frame(right_panel, bg='#ffffff')
        title_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 账单年月选择
        year_month_var = tk.StringVar()
        year_month_combo = ttk.Combobox(title_frame, textvariable=year_month_var, width=12, 
                                         state='readonly', font=("Microsoft YaHei", 18))
        months = []
        if DATEUTIL_AVAILABLE:
            for i in range(12):
                month_date = datetime.now().replace(day=1) - relativedelta(months=i)
                months.append(month_date.strftime("%Y年%m月"))
        else:
            for i in range(12):
                year = datetime.now().year
                month = datetime.now().month - i
                if month <= 0:
                    month += 12
                    year -= 1
                months.append(f"{year}年{month:02d}月")
        year_month_combo['values'] = months
        year_month_combo.current(0)
        year_month_combo.pack(side=tk.LEFT, padx=(0, 10))
        year_month_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 客户和日期
        info_row = tk.Frame(right_panel, bg='#ffffff')
        info_row.pack(fill=tk.X, pady=(0, 10))
        
        # 客户名称
        left_frame = tk.Frame(info_row, bg='#ffffff')
        left_frame.pack(side=tk.LEFT)
        tk.Label(left_frame, text="客户名称:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        customer_var = tk.StringVar(value=detected_customer)
        customer_entry = tk.Entry(left_frame, textvariable=customer_var, width=25, font=("Microsoft YaHei", 11))
        customer_entry.pack(side=tk.LEFT, padx=(5, 0))
        customer_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 制表日期
        right_frame = tk.Frame(info_row, bg='#ffffff')
        right_frame.pack(side=tk.RIGHT)
        tk.Label(right_frame, text="制表日期:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT, padx=(30, 5))
        date_var = tk.StringVar(value=datetime.now().strftime("%Y年%m月%d日"))
        date_entry = tk.Entry(right_frame, textvariable=date_var, width=18, font=("Microsoft YaHei", 11))
        date_entry.pack(side=tk.LEFT)
        date_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 提示信息
        hint_label = tk.Label(right_panel, text="💡 双击列表项可编辑 | 右键可打开编辑菜单", 
                              font=("Microsoft YaHei", 9), fg='#3498db', bg='#ffffff')
        hint_label.pack(anchor=tk.W, pady=(0, 5))
        
        # 商品明细表格
        columns = ('序号', '订单号', '品名', '规格', '数量', '单价', '金额')
        tree = ttk.Treeview(right_panel, columns=columns, show='headings', height=10)
        
        tree.heading('序号', text='序号')
        tree.heading('订单号', text='订单号')
        tree.heading('品名', text='品名')
        tree.heading('规格', text='规格')
        tree.heading('数量', text='数量')
        tree.heading('单价', text='单价(¥)')
        tree.heading('金额', text='金额(¥)')
        
        tree.column('序号', width=50, anchor='center')
        tree.column('订单号', width=100, anchor='center')
        tree.column('品名', width=150, anchor='center')
        tree.column('规格', width=100, anchor='center')
        tree.column('数量', width=80, anchor='center')
        tree.column('单价', width=100, anchor='center')
        tree.column('金额', width=120, anchor='center')
        
        tree.pack(fill=tk.BOTH, expand=True, pady=5)
        scrollbar = ttk.Scrollbar(right_panel, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        
        # 右键菜单
        context_menu = tk.Menu(tree, tearoff=0)
        total_label_for_callback = [None]
        
        def edit_selected():
            selection = tree.selection()
            if selection:
                item_values = tree.item(selection[0])['values']
                self._edit_item_dialog(item_values, tree, selection[0], self._update_total_from_tree, total_label_for_callback[0])
        def delete_selected():
            selection = tree.selection()
            if selection and messagebox.askyesno("确认", "确定要删除这条记录吗?"):
                tree.delete(selection[0])
                is_modified[0] = True
                self._update_total_from_tree(tree, total_label_for_callback[0])
        context_menu.add_command(label="向上插入一行", command=insert_row_above)
        context_menu.add_command(label="向下插入一行", command=insert_row_below)
        context_menu.add_separator()
        context_menu.add_command(label="上面加一行", command=insert_row_above)
        context_menu.add_command(label="下面加一行", command=insert_row_below)
        context_menu.add_command(label="编辑", command=edit_selected)
        context_menu.add_command(label="删除", command=delete_selected)
        context_menu.add_separator()
        context_menu.add_command(label="向上加一行", command=insert_row_above)
        context_menu.add_command(label="向下加一行", command=insert_row_below)

        def on_double_click(event):
            item_id = tree.selection()
            if item_id:
                item_values = tree.item(item_id[0])['values']
                self._edit_item_dialog(item_values, tree, item_id[0], self._update_total_from_tree, total_label_for_callback[0])
        def show_context_menu_tree(event):
            item_id = tree.identify_row(event.y)
            if item_id:
                tree.selection_set(item_id)
                context_menu.post(event.x_root, event.y_root)
        tree.bind("<Double-Button-1>", on_double_click)
        tree.bind("<Button-3>", show_context_menu_tree)
        
        # 填充分析结果数据
        for idx, item_data in enumerate(items, 1):
            tree.insert('', tk.END, values=(
                idx,
                item_data.get('订单号', ''),
                item_data.get('品名', ''),
                item_data.get('规格', ''),
                item_data.get('数量', ''),
                item_data.get('单价', ''),
                item_data.get('金额', '')
            ))
        
        # 如果没有数据，添加空行（默认6行）
        if not items:
            for i in range(6):
                tree.insert('', tk.END, values=(i+1, '', '', '', '', '', ''))
        
        # 合计区域
        bottom_frame = tk.Frame(right_panel, bg='#ffffff')
        bottom_frame.pack(fill=tk.X, pady=10)
        
        total_frame = tk.Frame(bottom_frame, bg='#ffffff')
        total_frame.pack(side=tk.RIGHT)
        tk.Label(total_frame, text="金额合计(¥): ", font=("Microsoft YaHei", 12, "bold"), bg='#ffffff').pack(side=tk.LEFT)
        total_label = tk.Label(total_frame, text="0.00", font=("Microsoft YaHei", 14, "bold"), fg='red', bg='#ffffff')
        total_label.pack(side=tk.LEFT)
        total_label_for_callback[0] = total_label
        
        # 自动计算合计
        self._update_total_from_tree(tree, total_label)
        
        # 按钮区域
        btn_frame = tk.Frame(right_panel, bg='#ffffff')
        btn_frame.pack(pady=10)
        
        def save_to_invoice():
            invoice_data = {
                'year_month': year_month_var.get(),
                'customer': customer_var.get(),
                'date': date_var.get(),
                'items': [],
                'total': 0.0,
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            for item_id in tree.get_children():
                values = tree.item(item_id)['values']
                if values[1] or values[2]:
                    invoice_data['items'].append({
                        '序号': values[0],
                        '订单号': values[1],
                        '品名': values[2],
                        '规格': values[3],
                        '数量': values[4],
                        '单价': values[5],
                        '金额': values[6]
                    })
                    try:
                        invoice_data['total'] += float(values[6]) if values[6] else 0
                    except:
                        pass
            self.monthly_invoices.append(invoice_data)
            self.save_data()
            
            # 同步付款记录 - 提取年份、客户、月份，更新对应付款记录
            year_month = invoice_data.get('year_month', '')
            customer = invoice_data.get('customer', '')
            total_amount = invoice_data.get('total', 0)
            
            if year_month and customer and total_amount > 0:
                try:
                    # 解析年月 (格式: YYYY-MM)
                    parts = year_month.split('-')
                    year = int(parts[0])
                    month = int(parts[1])
                    payment_date = year_month + '-01'  # 转换为完整日期
                    
                    # 查找是否存在对应的付款记录
                    payment_no = f"INV-{year_month}-{customer[:3]}"
                    existing_payment = None
                    for p in self.payments:
                        if p.get('payment_no') == payment_no and p.get('customer') == customer:
                            existing_payment = p
                            break
                    
                    if existing_payment:
                        # 更新现有付款记录
                        existing_payment['amount'] = total_amount
                        existing_payment['payment_date'] = payment_date
                    else:
                        # 创建新付款记录
                        new_payment = {
                            'payment_no': payment_no,
                            'customer': customer,
                            'amount': total_amount,
                            'payment_date': payment_date,
                            'payment_method': '对账单转入',
                            'status': '已付款',
                            'remark': f'对账单自动同步'
                        }
                        self.payments.append(new_payment)
                    
                    self.save_payments()
                except Exception as e:
                    print(f"同步付款记录失败: {e}")
            
            self.update_stats()
            messagebox.showinfo("成功", "已保存到对账单列表!")
            result_window.destroy()
        
        def cancel():
            result_window.destroy()
        
        tk.Button(btn_frame, text="💾 保存到对账单", command=save_to_invoice, width=18, height=2, 
                  bg='#27ae60', fg='white', font=("Microsoft YaHei", 11), relief=tk.FLAT).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="❌ 取消", command=cancel, width=15, height=2, 
                  bg='#95a5a6', fg='white', font=("Microsoft YaHei", 11), relief=tk.FLAT).pack(side=tk.LEFT, padx=10)
    
    def _re_analyze_photo(self, window, photo_path, tree, customer_var, year_month_var):
        """重新分析图片"""
        if not photo_path:
            messagebox.showwarning("警告", "请先拖拽一张送货单图片")
            return
        self.status_label.config(text="正在重新分析图片...")
        try:
            analyzed_data = self._analyze_delivery_photo(photo_path)
            if analyzed_data and analyzed_data.get('items'):
                # 清空现有数据
                for item in tree.get_children():
                    tree.delete(item)
                # 填充新数据
                items = analyzed_data.get('items', [])
                for idx, item_data in enumerate(items, 1):
                    tree.insert('', tk.END, values=(
                        idx,
                        item_data.get('订单号', ''),
                        item_data.get('品名', ''),
                        item_data.get('规格', ''),
                        item_data.get('数量', ''),
                        item_data.get('单价', ''),
                        item_data.get('金额', '')
                    ))
                # 更新客户名称
                if analyzed_data.get('客户名称'):
                    customer_var.set(analyzed_data.get('客户名称'))
                messagebox.showinfo("成功", f"分析完成，识别到 {len(items)} 条记录")
            else:
                messagebox.showinfo("提示", "未能从图片中识别出送货明细")
        except Exception as e:
            messagebox.showerror("错误", f"分析图片时出错\n{str(e)}")
        finally:
            self.status_label.config(text="就绪")
    
    def _clear_analyze_result(self, tree):
        """清空分析结果"""
        if tree.get_children() and messagebox.askyesno("确认", "确定要清空所有记录吗?"):
            for item in tree.get_children():
                tree.delete(item)
    
    def _reset_analyze_result(self, original_data, tree):
        """重置为原始分析结果"""
        items = original_data.get('items', [])
        for item in tree.get_children():
            tree.delete(item)
        for idx, item_data in enumerate(items, 1):
            tree.insert('', tk.END, values=(
                idx,
                item_data.get('订单号', ''),
                item_data.get('品名', ''),
                item_data.get('规格', ''),
                item_data.get('数量', ''),
                item_data.get('单价', ''),
                item_data.get('金额', '')
            ))

    # -------------------- 报价单相关 --------------------
    def create_new_quotation(self):
        """创建新报价单"""
        self.status_label.config(text="正在创建报价单...")
        self.clear_main_content()

        container = tk.Frame(self.main_content_frame, bg='#ffffff', padx=20, pady=15)
        container.pack(fill=tk.BOTH, expand=True)

        # 返回按钮
        back_frame = tk.Frame(container, bg='#ffffff')
        back_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Button(back_frame, text="← 返回主页", command=self.show_welcome_page,
                  font=("Microsoft YaHei", 10), bg='#95a5a6', fg='white',
                  relief=tk.FLAT, padx=15, pady=5).pack(side=tk.LEFT)

        # 公司名称（来自系统设置）
        company_name = self.company_name or "公司名称"
        company_label = tk.Label(container, text=company_name, font=("Microsoft YaHei", 16, "bold"),
                                  bg='#ffffff', fg='#2c3e50')
        company_label.pack(pady=(5, 5))

        # 报价单标题
        title_label = tk.Label(container, text="报价单", font=("Microsoft YaHei", 18, "bold"),
                               bg='#ffffff', fg='#e74c3c')
        title_label.pack(pady=(0, 15))

        # 客户、联系人、日期选择区域
        info_frame = tk.Frame(container, bg='#ffffff')
        info_frame.pack(fill=tk.X, pady=(0, 15))

        # 客户下拉选择
        customer_frame = tk.Frame(info_frame, bg='#ffffff')
        customer_frame.pack(side=tk.LEFT, padx=(0, 20))
        tk.Label(customer_frame, text="客户:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        customer_var = tk.StringVar()
        customer_combo = ttk.Combobox(customer_frame, textvariable=customer_var, width=20, state='readonly')
        
        def refresh_customer_combo():
            try:
                customers_path = os.path.join(self.data_dir, 'customers.json')
                if os.path.exists(customers_path):
                    with open(customers_path, 'r', encoding='utf-8') as f:
                        customers_data = json.load(f)
                    customer_names = [c.get('name', '') for c in customers_data if c.get('name')]
                    customer_combo['values'] = customer_names if customer_names else ['暂无客户']
                else:
                    customer_combo['values'] = ['暂无客户']
            except:
                customer_combo['values'] = ['暂无客户']
        
        refresh_customer_combo()
        if customer_combo['values']:
            customer_combo.current(0)
        customer_combo.pack(side=tk.LEFT, padx=(5, 0))

        # 客户选择后自动填充联系人
        def on_customer_selected(event):
            selected_customer = customer_var.get()
            for c in self.customers:
                if c.get('name') == selected_customer:
                    contact_var.set(c.get('contact', ''))
                    break

        customer_combo.bind('<<ComboboxSelected>>', on_customer_selected)

        # 联系人文本框
        contact_frame = tk.Frame(info_frame, bg='#ffffff')
        contact_frame.pack(side=tk.LEFT, padx=(0, 20))
        tk.Label(contact_frame, text="联系人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        contact_var = tk.StringVar()
        contact_entry = tk.Entry(contact_frame, textvariable=contact_var, width=15, font=("Microsoft YaHei", 11))
        contact_entry.pack(side=tk.LEFT, padx=(5, 0))

        # 日期文本框（右对齐）
        date_frame = tk.Frame(info_frame, bg='#ffffff')
        date_frame.pack(side=tk.RIGHT)
        tk.Label(date_frame, text="日期:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        date_entry = tk.Entry(date_frame, width=15, font=("Microsoft YaHei", 11), justify=tk.RIGHT)
        date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        date_entry.pack(side=tk.LEFT, padx=(5, 0))

        # 商品列表（只显示4行）
        columns = ('编号', '品名', '规格', '单价/码')
        tree = ttk.Treeview(container, columns=columns, show='headings', height=4)

        tree.heading('编号', text='编号')
        tree.heading('品名', text='品名')
        tree.heading('规格', text='规格')
        tree.heading('单价/码', text='单价/码')

        tree.column('编号', width=80, anchor='center')
        tree.column('品名', width=250, anchor='center')
        tree.column('规格', width=150, anchor='center')
        tree.column('单价/码', width=150, anchor='center')

        tree.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)

        # 初始4行空数据
        for i in range(4):
            tree.insert('', tk.END, values=(i+1, '', '', ''))

        # 右键菜单
        context_menu = tk.Menu(tree, tearoff=0)
        def edit_selected():
            selection = tree.selection()
            if selection:
                item_values = tree.item(selection[0])['values']
                self._edit_quotation_item(item_values, tree, selection[0])
        def delete_selected():
            selection = tree.selection()
            if selection:
                tree.delete(selection[0])
                # 重新编号
                for idx, item_id in enumerate(tree.get_children(), 1):
                    values = tree.item(item_id)['values']
                    tree.item(item_id, values=(idx,) + values[1:])

        context_menu.add_command(label="编辑", command=edit_selected)
        context_menu.add_command(label="删除", command=delete_selected)
        context_menu.add_separator()
        context_menu.add_command(label="添加行", command=lambda: tree.insert('', tk.END, values=(len(tree.get_children())+1, '', '', '')))

        def show_context_menu(event):
            item_id = tree.identify_row(event.y)
            if item_id:
                tree.selection_set(item_id)
                context_menu.post(event.x_root, event.y_root)

        tree.bind("<Button-3>", show_context_menu)

        # 双击直接编辑
        edit_entry = [None]
        edit_col = [None]
        
        def on_tree_double_click(event):
            nonlocal edit_entry, edit_col
            if edit_entry[0]:
                edit_entry[0].destroy()
                edit_entry[0] = None
            
            item_id = tree.identify_row(event.y)
            if not item_id:
                return
            
            column = tree.identify_column(event.x)
            col_map = {'#1': 0, '#2': 1, '#3': 2, '#4': 3}
            col_idx = col_map.get(column, 1)
            if col_idx == 0:
                return
            
            edit_col[0] = col_idx
            bbox = tree.bbox(item_id, column)
            if not bbox:
                return
            x, y, width, height = bbox
            
            values = tree.item(item_id)['values']
            
            # 获取树视图的父容器坐标
            tree_x = tree.winfo_rootx()
            tree_y = tree.winfo_rooty()
            entry = tk.Entry(tree, font=("Microsoft YaHei", 10))
            entry.insert(0, values[col_idx])
            entry.place(x=x, y=y, width=width, height=height)
            entry.select_range(0, tk.END)
            entry.focus()
            edit_entry[0] = entry
            
            def save_edit(e=None):
                new_val = entry.get()
                new_values = list(values)
                new_values[col_idx] = new_val
                tree.item(item_id, values=new_values)
                entry.destroy()
                edit_entry[0] = None
            
            def cancel_edit(e=None):
                entry.destroy()
                edit_entry[0] = None
            
            entry.bind('<Return>', save_edit)
            entry.bind('<KP_Enter>', save_edit)
            entry.bind('<Escape>', cancel_edit)
            entry.bind('<FocusOut>', lambda e: save_edit())
        
        tree.bind("<Double-Button-1>", on_tree_double_click)

        # 上方第一个备注文本框（新增）
        remark_frame_1 = tk.Frame(container, bg='#ffffff')
        remark_frame_1.pack(fill=tk.X, pady=(10, 5))
        tk.Label(remark_frame_1, text="备注1:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        remark_entry_1 = tk.Entry(remark_frame_1, width=60, font=("Microsoft YaHei", 11))
        remark_entry_1.pack(side=tk.LEFT, padx=(5, 0), fill=tk.X, expand=True)

        # 上方第二个备注文本框（原有）
        top_remark_frame = tk.Frame(container, bg='#ffffff')
        top_remark_frame.pack(fill=tk.X, pady=(5, 10))

        # 下方多行备注文本框
        remark_frame = tk.Frame(container, bg='#ffffff')
        remark_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        remark_text = tk.Text(remark_frame, width=80, height=10, font=("Microsoft YaHei", 11), wrap=tk.WORD)
        remark_text.pack(fill=tk.BOTH, expand=True, padx=(0, 0))
        remark_text.insert('1.0', "1. 以上为本公司统一价格\nAll prices listed above are our company's standard quotations.\n\n2. 颜色可以根据客户要求生产，一般不另加价。\nColors can be produced as per customer's requirements, with generally no additional charge.\n\n3. 本公司负责为客户送到指定地点。\nWe will arrange delivery of the goods to the designated location specified by the customer.\n4. 本公司结算方式按约定。\nPayment terms shall be as mutually agreed upon by both parties.\n5. 本报价单一经客户确认后，即作为定单附件。\nOnce confirmed by the customer, this quotation shall become an integral part of the corresponding purchase order (PO).")
        remark_text.config(state='disabled')

        # 客户负责人和报价人（向下移动）
        bottom_frame = tk.Frame(container, bg='#ffffff')
        bottom_frame.pack(fill=tk.X, pady=(50, 15))

        # 客户负责人（左对齐，文本框带下划线）
        manager_frame = tk.Frame(bottom_frame, bg='#ffffff')
        manager_frame.pack(side=tk.LEFT)
        tk.Label(manager_frame, text="客户负责人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        manager_entry = tk.Entry(manager_frame, width=20, font=("Microsoft YaHei", 11), justify=tk.LEFT)
        manager_entry.insert(0, "____________")
        manager_entry.pack(side=tk.LEFT, padx=(5, 0))

        # 报价人（右对齐，文本框带下划线）
        quoter_frame = tk.Frame(bottom_frame, bg='#ffffff')
        quoter_frame.pack(side=tk.RIGHT)
        tk.Label(quoter_frame, text="报价人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        quoter_entry = tk.Entry(quoter_frame, width=20, font=("Microsoft YaHei", 11), justify=tk.CENTER)
        quoter_entry.insert(0, getattr(self, 'quoter', '') or "____________")
        quoter_entry.pack(side=tk.LEFT, padx=(5, 0))

        # 按钮区域
        btn_frame = tk.Frame(container, bg='#ffffff')
        btn_frame.pack(pady=15)

        def save_quotation():
            quotation_data = {
                'customer': customer_var.get(),
                'contact': contact_var.get(),
                'date': date_entry.get(),
                'items': [],
                'remark': remark_text.get('1.0', tk.END).strip(),
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            for item_id in tree.get_children():
                values = tree.item(item_id)['values']
                if values[1]:  # 品名不为空
                    quotation_data['items'].append({
                        '编号': values[0],
                        '品名': values[1],
                        '规格': values[2],
                        '单价/码': values[3]
                    })
            self.quotations.append(quotation_data)
            self.save_data()
            self.update_stats()
            messagebox.showinfo("成功", "报价单保存成功")
            self.show_welcome_page()
            self.status_label.config(text="报价单已保存")

        tk.Button(btn_frame, text="保存报价单", command=save_quotation, width=15, height=2,
                  bg='#27ae60', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="取消", command=self.show_welcome_page, width=15, height=2,
                  bg='#95a5a6', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=10)

    def _edit_quotation_item(self, item_values, tree, item_id):
        """编辑报价单项"""
        edit_window = tk.Toplevel(self.root)
        edit_window.title("编辑商品")
        edit_window.geometry("400x200")
        edit_frame = tk.Frame(edit_window, padx=20, pady=20)
        edit_frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(edit_frame, text="品名:", width=10, anchor=tk.W).grid(row=1, column=0, sticky=tk.W, pady=10)
        name_entry = tk.Entry(edit_frame, width=25)
        name_entry.grid(row=1, column=1, pady=10)
        name_entry.insert(0, item_values[1])

        tk.Label(edit_frame, text="规格:", width=10, anchor=tk.W).grid(row=2, column=0, sticky=tk.W, pady=10)
        spec_entry = tk.Entry(edit_frame, width=25)
        spec_entry.grid(row=2, column=1, pady=10)
        spec_entry.insert(0, item_values[2])

        tk.Label(edit_frame, text="单价/码:", width=10, anchor=tk.W).grid(row=3, column=0, sticky=tk.W, pady=10)
        price_entry = tk.Entry(edit_frame, width=25)
        price_entry.grid(row=3, column=1, pady=10)
        price_entry.insert(0, item_values[3])

        def save_edit():
            new_values = (item_values[0], name_entry.get(), spec_entry.get(), price_entry.get())
            tree.item(item_id, values=new_values)
            edit_window.destroy()

        btn_frame = tk.Frame(edit_window)
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="保存", command=save_edit, width=12, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="取消", command=edit_window.destroy, width=12, bg='#95a5a6').pack(side=tk.LEFT, padx=10)

    def _create_quotation_row(self, parent, row_num, all_rows, total_label, update_total_cb):
        """创建报价单的一行（内部使用）"""
        # 实际已在create_new_quotation中实现，此处占位
        pass

    def _save_excel_quotation(self, quote_info, excel_rows, remark_entry, total_label):
        """保存Excel风格的报价单"""
        quotation_data = {
            'quote_no': quote_info['quote_no'].get(),
            'customer': quote_info['customer'].get(),
            'contact': quote_info['contact'].get(),
            'phone': quote_info['phone'].get(),
            'date': quote_info['date'].get(),
            'validity': quote_info['validity'].get(),
            'items': [],
            'total': 0.0,
            'remark': remark_entry.get(),
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        for idx, row in enumerate(excel_rows, 1):
            if row[1].get():
                item_data = {
                    '序号': idx,
                    '项目名称': row[1].get(),
                    '规格': row[2].get(),
                    '单位': row[3].get(),
                    '数量': row[4].get(),
                    '单价': row[5].get(),
                    '折扣': row[6].get(),
                    '合计': row[7].get()
                }
                quotation_data['items'].append(item_data)
                try:
                    quotation_data['total'] += float(row[7].get() or 0)
                except:
                    pass
        self.quotations.append(quotation_data)
        self.save_data()
        self.update_stats()
        messagebox.showinfo("成功", "报价单保存成功")
        self.show_welcome_page()
        self.status_label.config(text="报价单已保存")

    # -------------------- 显示列表 --------------------
    def show_monthly_invoices(self):
        """显示对账单列表 - 标签页显示"""
        self.status_label.config(text="正在加载对账单列表...")
        self.clear_main_content()
        
        container = tk.Frame(self.main_content_frame, bg='#ffffff', padx=20, pady=15)
        container.pack(fill=tk.BOTH, expand=True)
        
        # 顶部标题区域
        header_frame = tk.Frame(container, bg='#ffffff')
        header_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Label(header_frame, text="对账单管理", font=("Microsoft YaHei", 18, "bold"), 
                bg='#ffffff', fg='#2c3e50').pack(side=tk.LEFT)
        tk.Button(header_frame, text="← 返回主页", command=self.show_welcome_page,
                  font=("Microsoft YaHei", 10), bg='#95a5a6', fg='white',
                  relief=tk.FLAT, padx=15, pady=5).pack(side=tk.RIGHT)

        # 创建标签页
        notebook = ttk.Notebook(container)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # ========== 标签页1: 对账单列表 ==========
        list_tab = tk.Frame(notebook, bg='#ffffff', padx=15, pady=15)
        notebook.add(list_tab, text="📋 对账单列表")
        
        # 搜索区域
        search_frame = tk.Frame(list_tab, bg='#f5f5f5', padx=10, pady=10)
        search_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(search_frame, text="搜索客户:", font=("Microsoft YaHei", 11), bg='#f5f5f5').pack(side=tk.LEFT)
        search_var = tk.StringVar()
        search_combo = ttk.Combobox(search_frame, textvariable=search_var, width=25)
        def refresh_customer_combo():
            try:
                customers_path = os.path.join(self.data_dir, 'customers.json')
                if os.path.exists(customers_path):
                    with open(customers_path, 'r', encoding='utf-8') as f:
                        customers_data = json.load(f)
                    customer_names = [c.get('name', '') for c in customers_data if c.get('name')]
                    search_combo['values'] = ['全部'] + customer_names if customer_names else ['全部']
                else:
                    search_combo['values'] = ['全部']
            except:
                search_combo['values'] = ['全部']
        refresh_customer_combo()
        search_combo.current(0)
        search_combo.pack(side=tk.LEFT, padx=5)

        # 表格
        columns = ('序号', '账单年月', '客户名称', '品名列表', '总金额')
        tree = ttk.Treeview(list_tab, columns=columns, show='headings', height=15)
        tree.heading('序号', text='序号')
        tree.heading('账单年月', text='账单年月')
        tree.heading('客户名称', text='客户名称')
        tree.heading('品名列表', text='品名列表')
        tree.heading('总金额', text='总金额')
        tree.column('序号', width=60, anchor='center')
        tree.column('账单年月', width=120, anchor='center')
        tree.column('客户名称', width=180, anchor='center')
        tree.column('品名列表', width=300, anchor='center')
        tree.column('总金额', width=120, anchor='center')
        tree.pack(fill=tk.BOTH, expand=True, pady=10)
        scrollbar = ttk.Scrollbar(list_tab, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)

        def refresh_list():
            for item in tree.get_children():
                tree.delete(item)
            search_text = search_var.get()
            sorted_invoices = sorted(self.monthly_invoices, key=lambda x: x.get('year_month', ''), reverse=True)
            for idx, inv in enumerate(sorted_invoices, 1):
                if search_text == '全部' or not search_text or search_text in inv.get('customer', ''):
                    products_list = inv.get('items', [])
                    if products_list:
                        product_names = [p.get('品名', '') + ' ' + p.get('规格', '') for p in products_list if p.get('品名')]
                        product_display = ', '.join(product_names[:5])
                    else:
                        product_display = ''
                    tree.insert('', tk.END, values=(idx, inv.get('year_month', ''), inv.get('customer', ''), product_display, f"¥{inv.get('total', 0):.2f}"))

        search_combo.bind('<<ComboboxSelected>>', lambda e: refresh_list())
        refresh_list()

        # 详情标签页框架（初始为空）
        detail_tab = tk.Frame(notebook, bg='#ffffff', padx=15, pady=15)
        notebook.add(detail_tab, text="📄 账单详情")
        
        tk.Label(detail_tab, text="请从左侧列表选择一条对账单查看详情", 
                 font=("Microsoft YaHei", 12), bg='#ffffff', fg='#999999').pack(pady=50)
        
        # 用于存储详情内容的frame引用
        detail_content_frame = [detail_tab]

        def view_detail():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请选择对账单")
                return
            item = tree.item(selection[0])
            idx = item['values'][0] - 1
            self._view_invoice_detail_in_tab(idx, notebook, detail_content_frame)
            notebook.select(1)  # 切换到详情标签页

        def delete_invoice():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            if messagebox.askyesno("确认", "确定要删除这条账单吗?"):
                item = tree.item(selection[0])
                idx = item['values'][0] - 1
                del self.monthly_invoices[idx]
                self.save_data()
                self.update_stats()
                refresh_list()
                messagebox.showinfo("成功", "删除成功!")

        context_menu = tk.Menu(tree, tearoff=0)
        context_menu.add_command(label="查看详情", command=view_detail)
        context_menu.add_command(label="删除", command=delete_invoice)
        context_menu.add_separator()
        def insert_invoice_above():
            # 在选中项上方插入新对账单
            selection = tree.selection()
            if selection:
                idx = tree.index(selection[0])
            else:
                idx = 0
            self.create_new_monthly_invoice(insert_index=idx)
        
        def insert_invoice_below():
            # 在选中项下方插入新对账单
            selection = tree.selection()
            if selection:
                idx = tree.index(selection[0]) + 1
            else:
                idx = 0
            self.create_new_monthly_invoice(insert_index=idx)
        
        def show_context_menu(event):
            item_id = tree.identify_row(event.y)
            if item_id:
                tree.selection_set(item_id)
                context_menu.post(event.x_root, event.y_root)
        tree.bind("<Double-Button-1>", lambda e: view_detail())
        tree.bind("<Button-3>", show_context_menu)

        # 按钮区域
        btn_frame = tk.Frame(list_tab, bg='#ffffff')
        btn_frame.pack(pady=10)
        tk.Label(btn_frame, text="⚠️ 请选择对账单", font=("Microsoft YaHei", 10), fg='#f39c12').pack(side=tk.LEFT, padx=(0, 10))
        tk.Button(btn_frame, text="查看详情", command=view_detail, width=15, height=2, bg='#3498db', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="删除", command=delete_invoice, width=15, height=2, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        self.status_label.config(text="对账单列表已加载")
    def _view_invoice_detail_in_tab(self, idx, notebook, detail_content_frame):
        """在标签页中显示对账单详情"""
        invoice = self.monthly_invoices[idx]
        is_modified = [False]
        
        # 清空详情内容框架
        for widget in detail_content_frame[0].winfo_children():
            widget.destroy()
        
        # 创建内容容器
        container = tk.Frame(detail_content_frame[0], bg='#ffffff', padx=15, pady=15)
        container.pack(fill=tk.BOTH, expand=True)
        
        # 顶部标题和操作区域
        header_frame = tk.Frame(container, bg='#ffffff')
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 标题
        tk.Label(header_frame, text=f"{invoice.get('year_month', '')} 对账单", 
                font=("Microsoft YaHei", 18, "bold"), bg='#ffffff', fg='#2c3e50').pack(side=tk.LEFT)
        
        # 操作按钮
        btn_frame_header = tk.Frame(header_frame, bg='#ffffff')
        btn_frame_header.pack(side=tk.RIGHT)

        # 客户和日期信息区域
        info_row2 = tk.Frame(container, bg='#ffffff')
        info_row2.pack(fill=tk.X, pady=10)
        left_frame = tk.Frame(info_row2, bg='#ffffff')
        left_frame.pack(side=tk.LEFT)
        tk.Label(left_frame, text="客户名称:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        customer_var = tk.StringVar(value=invoice.get('customer', ''))
        customer_entry = tk.Entry(left_frame, textvariable=customer_var, width=25, font=("Microsoft YaHei", 11))
        customer_entry.pack(side=tk.LEFT, padx=(5, 0))
        customer_var.trace('w', lambda *a: setattr(is_modified, '0', True))

        right_frame = tk.Frame(info_row2, bg='#ffffff')
        right_frame.pack(side=tk.RIGHT)
        tk.Label(right_frame, text="制表日期:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT, padx=(30, 5))
        date_var = tk.StringVar(value=invoice.get('date', '') or datetime.now().strftime("%Y年%m月%d日"))
        date_entry = tk.Entry(right_frame, textvariable=date_var, width=18, font=("Microsoft YaHei", 11))
        date_entry.pack(side=tk.LEFT)
        date_var.trace('w', lambda *a: setattr(is_modified, '0', True))

        hint_label = tk.Label(container, text="💡 双击列表项可编辑，右键可打开编辑菜单", font=("Microsoft YaHei", 9), fg='#3498db', bg='#ffffff')
        hint_label.pack(anchor=tk.W, pady=(0, 5))

        # 表格
        columns = ('序号', '订单号', '指令号', '送货日期', '送货单号', '品名', '规格', '颜色', '单位', '数量', '单价', '染费', '金额')
        detail_tree = ttk.Treeview(container, columns=columns, show='headings', height=12)
        for col in columns:
            detail_tree.heading(col, text=col)
            width = 50 if col == '序号' else 80
            detail_tree.column(col, width=width, anchor='center')
        detail_tree.pack(fill=tk.BOTH, expand=True, pady=10)
        scrollbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=detail_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        detail_tree.configure(yscrollcommand=scrollbar.set)

        # 合计标签（用于回调）
        total_label = [None]

        # 右键菜单
        context_menu = tk.Menu(detail_tree, tearoff=0)
        def edit_selected():
            selection = detail_tree.selection()
            if selection:
                item_values = detail_tree.item(selection[0])['values']
                self._edit_item_dialog(item_values, detail_tree, selection[0], self._update_total_from_tree, total_label[0])
        def delete_selected():
            selection = detail_tree.selection()
            if selection and messagebox.askyesno("确认", "确定要删除这条记录吗?"):
                detail_tree.delete(selection[0])
                is_modified[0] = True
                self._update_total_from_tree(detail_tree, total_label[0])
        def insert_row_above():
            selection = detail_tree.selection()
            if selection:
                parent = detail_tree.parent(selection[0])
                siblings = detail_tree.get_children(parent)
                idx = siblings.index(selection[0])
                new_values = list(detail_tree.item(selection[0])['values'])
                new_values[0] = ''  # 清空序号
                for i in range(1, 13):
                    new_values[i] = ''
                detail_tree.insert(parent, idx, values=new_values)
                is_modified[0] = True
                # 重新编号
                for i, item_id in enumerate(detail_tree.get_children(parent), 1):
                    values = list(detail_tree.item(item_id)['values'])
                    values[0] = i
                    detail_tree.item(item_id, values=values)
        def insert_row_below():
            selection = detail_tree.selection()
            if selection:
                parent = detail_tree.parent(selection[0])
                siblings = detail_tree.get_children(parent)
                idx = siblings.index(selection[0]) + 1
                new_values = list(detail_tree.item(selection[0])['values'])
                new_values[0] = ''  # 清空序号
                for i in range(1, 13):
                    new_values[i] = ''
                detail_tree.insert(parent, idx, values=new_values)
                is_modified[0] = True
                # 重新编号
                for i, item_id in enumerate(detail_tree.get_children(parent), 1):
                    values = list(detail_tree.item(item_id)['values'])
                    values[0] = i
                    detail_tree.item(item_id, values=values)
        context_menu.add_command(label="编辑", command=edit_selected)
        context_menu.add_command(label="删除", command=delete_selected)
        context_menu.add_separator()
        context_menu.add_command(label="向上加一行", command=insert_row_above)
        context_menu.add_command(label="向下加一行", command=insert_row_below)

        def on_double_click(event):
            item_id = detail_tree.selection()
            if item_id:
                item_values = detail_tree.item(item_id[0])['values']
                self._edit_item_dialog(item_values, detail_tree, item_id[0], self._update_total_from_tree, total_label[0])
        def show_context_menu_tree(event):
            item_id = detail_tree.identify_row(event.y)
            if item_id:
                detail_tree.selection_set(item_id)
                context_menu.post(event.x_root, event.y_root)
        detail_tree.bind("<Double-Button-1>", on_double_click)
        detail_tree.bind("<Button-3>", show_context_menu_tree)

        # 填充数据
        existing_count = len(detail_tree.get_children())
        for item_data in invoice.get('items', []):
            detail_tree.insert('', tk.END, values=(
                item_data.get('序号', ''), item_data.get('订单号', ''), item_data.get('指令号', ''),
                item_data.get('送货日期', ''), item_data.get('送货单号', ''), item_data.get('品名', ''),
                item_data.get('规格', ''), item_data.get('颜色', ''), item_data.get('单位', ''),
                item_data.get('数量', ''), item_data.get('单价', ''), item_data.get('染费', ''),
                item_data.get('金额', '')
            ))

        # 确保至少有6行数据
        current_count = len(detail_tree.get_children())
        for i in range(current_count, 6):
            detail_tree.insert('', tk.END, values=(
                i + 1, '', '', '', '', '', '', '', '', '', '', '', ''
            ))

        # 备注和合计金额同一行，备注在左，合计在右
        bottom_frame = tk.Frame(container, bg='#ffffff')
        bottom_frame.pack(fill=tk.X, pady=10)
        tk.Label(bottom_frame, text="备注:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        remark_var = tk.StringVar(value=invoice.get('remark', ''))
        remark_entry = tk.Entry(bottom_frame, textvariable=remark_var, width=60, font=("Microsoft YaHei", 11))
        remark_entry.pack(side=tk.LEFT, padx=5)
        remark_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        # 金额合计靠右对齐
        total_frame = tk.Frame(bottom_frame, bg='#ffffff')
        total_frame.pack(side=tk.RIGHT)
        tk.Label(total_frame, text="金额合计: ¥", font=("Microsoft YaHei", 12, "bold"), bg='#ffffff').pack(side=tk.LEFT)
        t_label = tk.Label(total_frame, text=f"{invoice.get('total', 0):.2f}", font=("Microsoft YaHei", 14, "bold"), fg='red', bg='#ffffff')
        t_label.pack(side=tk.LEFT)
        total_label[0] = t_label

        # 客户核签区域 - 客户核签在左，经办人在右
        sign_frame = tk.Frame(container, bg='#ffffff')
        sign_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 客户核签和经办人同一行
        top_sign_row = tk.Frame(sign_frame, bg='#ffffff')
        top_sign_row.pack(fill=tk.X)
        
        # 客户核签（左侧）
        customer_sign_frame = tk.Frame(top_sign_row, bg='#ffffff')
        customer_sign_frame.pack(side=tk.LEFT)
        tk.Label(customer_sign_frame, text="客户核签:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        customer_sign_var = tk.StringVar(value=invoice.get('customer_sign', ''))
        customer_sign_entry = tk.Entry(customer_sign_frame, textvariable=customer_sign_var, width=20, font=("Microsoft YaHei", 11))
        customer_sign_entry.pack(side=tk.LEFT, padx=5)
        customer_sign_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 经办人（右侧）
        handler_sign_frame = tk.Frame(top_sign_row, bg='#ffffff')
        handler_sign_frame.pack(side=tk.RIGHT)
        tk.Label(handler_sign_frame, text="经办人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        handler_sign_var = tk.StringVar(value=invoice.get('handler', getattr(self, 'handler', '')))
        handler_sign_entry = tk.Entry(handler_sign_frame, textvariable=handler_sign_var, width=15, font=("Microsoft YaHei", 11))
        handler_sign_entry.pack(side=tk.LEFT, padx=5)
        handler_sign_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 核签日期和手机号同一行
        bottom_sign_row = tk.Frame(sign_frame, bg='#ffffff')
        bottom_sign_row.pack(fill=tk.X, pady=(5, 0))
        
        # 核签日期（左侧，带下划线）
        sign_date_frame = tk.Frame(bottom_sign_row, bg='#ffffff')
        sign_date_frame.pack(side=tk.LEFT)
        tk.Label(sign_date_frame, text="核签日期:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        sign_date_var = tk.StringVar(value=invoice.get('sign_date', ''))
        sign_date_entry = tk.Entry(sign_date_frame, textvariable=sign_date_var, width=15, font=("Microsoft YaHei", 11))
        sign_date_entry.pack(side=tk.LEFT, padx=5)
        # 添加下划线效果
        sign_date_entry.config(relief=tk.SUNKEN, bd=1)
        sign_date_var.trace('w', lambda *a: setattr(is_modified, '0', True))
        
        # 手机号（右侧）
        phone_frame = tk.Frame(bottom_sign_row, bg='#ffffff')
        phone_frame.pack(side=tk.RIGHT)
        tk.Label(phone_frame, text="手机号：", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        phone_var = tk.StringVar(value=invoice.get('phone', '13924371586'))
        phone_entry = tk.Entry(phone_frame, textvariable=phone_var, width=15, font=("Microsoft YaHei", 11))
        phone_entry.pack(side=tk.LEFT, padx=5)
        phone_var.trace('w', lambda *a: setattr(is_modified, '0', True))

        def auto_save():
            invoice['customer'] = customer_var.get()
            invoice['date'] = date_var.get()
            invoice['remark'] = remark_var.get()
            invoice['customer_sign'] = customer_sign_var.get()
            invoice['handler'] = handler_sign_var.get()
            invoice['sign_date'] = sign_date_var.get()
            invoice['phone'] = phone_var.get()
            items = []
            total = 0.0
            for item_id in detail_tree.get_children():
                values = detail_tree.item(item_id)['values']
                if values[1] or values[2]:
                    items.append({
                        '序号': values[0], '订单号': values[1], '指令号': values[2],
                        '送货日期': values[3], '送货单号': values[4], '品名': values[5], '规格': values[6],
                        '颜色': values[7], '单位': values[8], '数量': values[9],
                        '单价': values[10], '染费': values[11], '金额': values[12]
                    })
                    try:
                        total += float(values[12]) if values[12] else 0
                    except:
                        pass
            invoice['items'] = items
            invoice['total'] = total
            self.save_data()
            self.update_stats()

        # 导出PDF功能
        def export_invoice_to_pdf():
            if not PDF_AVAILABLE:
                messagebox.showwarning("警告", "请先安装reportlab库: pip install reportlab")
                return
            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF文件", "*.pdf")], 
                                                     initialfile=f"{invoice.get('year_month', '')}_{invoice.get('customer', '账单')}")
            if file_path:
                # 导出前先保存最新数据
                auto_save()
                try:
                    from reportlab.lib.pagesizes import A4, landscape
                    from reportlab.lib.units import cm
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib import colors
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont
                    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
                    
                    # 注册中文字体
                    font_registered = False
                    font_paths = [
                        'C:/Windows/Fonts/simhei.ttf',
                        'C:/Windows/Fonts/simsun.ttc',
                        'C:/Windows/Fonts/msyh.ttc',
                        'C:/Windows/Fonts/simfang.ttf',
                        'C:/Windows/Fonts/simkai.ttf',
                        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                        '/System/Library/Fonts/PingFang.ttc',
                        '/Library/Fonts/Arial Unicode.ttf',
                    ]
                    
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                if font_path.endswith('.ttc'):
                                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                                else:
                                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                                font_registered = True
                                break
                            except:
                                continue
                    
                    # 创建字体样式
                    if font_registered:
                        title_style = ParagraphStyle('Title', fontName='ChineseFont', fontSize=22, leading=28, alignment=TA_CENTER, spaceAfter=8)
                        subtitle_style = ParagraphStyle('Subtitle', fontName='ChineseFont', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=6)
                        normal_style = ParagraphStyle('Normal', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                        header_style = ParagraphStyle('Header', fontName='ChineseFont', fontSize=10, leading=14, alignment=TA_CENTER)
                        cell_style = ParagraphStyle('Cell', fontName='ChineseFont', fontSize=9, leading=12, alignment=TA_CENTER)
                        right_style = ParagraphStyle('Right', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_RIGHT)
                    else:
                        title_style = ParagraphStyle('Title', fontSize=22, leading=28, alignment=TA_CENTER, spaceAfter=8)
                        subtitle_style = ParagraphStyle('Subtitle', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=6)
                        normal_style = ParagraphStyle('Normal', fontSize=11, leading=16, alignment=TA_LEFT)
                        header_style = ParagraphStyle('Header', fontSize=10, leading=14, alignment=TA_CENTER)
                        cell_style = ParagraphStyle('Cell', fontSize=9, leading=12, alignment=TA_CENTER)
                        right_style = ParagraphStyle('Right', fontSize=11, leading=16, alignment=TA_RIGHT)
                    
                    # 使用横版A4，上页边距1cm，下页边距0.1cm
                    doc = SimpleDocTemplate(file_path, pagesize=landscape(A4), encoding='utf-8',
                                          leftMargin=1*cm, rightMargin=1*cm, 
                                          topMargin=1*cm, bottomMargin=0.1*cm)
                    elements = []
                    
                    # 标题区域 - LOGO在左，公司名称在右
                    logo_path = getattr(self, 'logo_path', '')
                    logo_element = None
                    if logo_path and os.path.exists(logo_path):
                        try:
                            logo_element = RLImage(logo_path, width=2*cm, height=2*cm)
                        except:
                            pass
                    
                    # 公司名称 + LOGO 并排
                    if logo_element:
                        header_table = Table([[logo_element, Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style)]], 
                                            colWidths=[3*cm, 18*cm])
                        header_table.setStyle(TableStyle([
                            ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('LEFTPADDING', (0, 0), (0, 0), 0),
                            ('RIGHTPADDING', (0, 0), (0, 0), 10),
                        ]))
                        elements.append(header_table)
                    else:
                        elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                    
                    elements.append(Paragraph(f"<b>{invoice.get('year_month', '')} 对账单</b>", subtitle_style))
                    elements.append(Spacer(1, 0.3*cm))
                    
                    # 信息行 - 客户在左，制表日期在右
                    info_data = [
                        [Paragraph(f"客户: {invoice.get('customer', '')}", normal_style),
                         Paragraph(f"制表日期: {datetime.now().strftime('%Y年%m月%d日')}", normal_style)]
                    ]
                    info_table = Table(info_data, colWidths=[12*cm, 10*cm])
                    info_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    elements.append(info_table)
                    elements.append(Spacer(1, 0.3*cm))
                    
                    # 表格 - 横版A4使用更大的列宽
                    table_data = []
                    columns = ['序号', '订单号', '指令号', '送货日期', '送货单号', '品名', '规格', '颜色', '单位', '数量', '单价', '染费', '金额']
                    header_row = [Paragraph(col, header_style) for col in columns]
                    table_data.append(header_row)
                    
                    for item in invoice.get('items', []):
                        row = [Paragraph(str(item.get(col, '')), cell_style) for col in columns]
                        table_data.append(row)
                    
                    # 确保最少显示6行数据
                    for i in range(max(0, 6 - len(invoice.get('items', [])))):
                        empty_row = [Paragraph('', cell_style) for _ in columns]
                        table_data.append(empty_row)
                    
                    # 横版A4宽度约29.7cm，减去左右边距
                    col_widths = [1.2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2*cm, 1.2*cm, 1.8*cm, 1.8*cm, 1.8*cm, 2.5*cm]
                    # 统一行高：所有行高度相同
                    row_height = 0.8 * cm
                    row_heights = [row_height] * len(table_data)
                    t = Table(table_data, colWidths=col_widths, rowHeights=row_heights)
                    
                    # 根据字体注册情况选择合适的字体名
                    header_font = 'Helvetica-Bold' if not font_registered else 'ChineseFont'
                    t.setStyle(TableStyle([
                        # 表头背景色改为白色，文字黑色
                        ('BACKGROUND', (0, 0), (-1, 0), colors.white),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, 0), header_font),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        # 统一设置所有单元格上下内边距，保证行高一致
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                        ('BACKGROUND', (0, 1), (-1, -2), colors.HexColor('#f5f5f5')),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#f5f5f5')]),
                        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#ecf0f1')),
                        ('FONTNAME', (0, -1), (-1, -1), header_font),
                    ]))
                    elements.append(t)
                    
                    # 合计金额和备注同一行 - 备注左对齐，合计金额右对齐
                    elements.append(Spacer(1, 0.3*cm))
                    # 备注样式（左对齐）
                    remark_pdf_style = ParagraphStyle('RemarkPDF', fontName='Helvetica-Bold', fontSize=11, leading=16, alignment=TA_LEFT)
                    if font_registered:
                        remark_pdf_style = ParagraphStyle('RemarkPDF', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                    # 合计金额样式（右对齐）
                    total_style = ParagraphStyle('Total', fontName='Helvetica-Bold', fontSize=14, leading=18, alignment=TA_RIGHT)
                    if font_registered:
                        total_style = ParagraphStyle('Total', fontName='ChineseFont', fontSize=14, leading=18, alignment=TA_RIGHT)
                    total_amount = invoice.get('total', 0)
                    # 同一行表格：左侧备注文本框，右侧合计金额
                    remark_text = invoice.get('remark', '')
                    # 使用Table实现同一行布局：备注在左（约17cm），合计在右（约10cm）
                    remark_total_data = [[
                        Paragraph(f"<b>备注:</b> {remark_text}", remark_pdf_style),
                        Paragraph(f"金额合计: CNY {total_amount:.2f}", total_style)
                    ]]
                    remark_total_table = Table(remark_total_data, colWidths=[17*cm, 10*cm])
                    remark_total_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),  # 备注左对齐
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  # 合计右对齐
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (0, 0), 11),
                        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (1, 0), (1, 0), 14),
                        ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor('#e74c3c')),
                        ('TOPPADDING', (0, 0), (-1, -1), 5),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ]))
                    if font_registered:
                        remark_total_table.setStyle(TableStyle([
                            ('FONTNAME', (0, 0), (0, 0), 'ChineseFont'),
                            ('FONTNAME', (1, 0), (1, 0), 'ChineseFont'),
                        ]))
                    elements.append(remark_total_table)
                    
                    # 客户核签区域 - 向上移动到第一页，减小间距
                    elements.append(Spacer(1, 1*cm))
                    sign_title_style = ParagraphStyle('SignTitle', fontName='Helvetica-Bold', fontSize=11, leading=16, alignment=TA_LEFT)
                    if font_registered:
                        sign_title_style = ParagraphStyle('SignTitle', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                    
                    # 获取经办人名称，默认为王平
                    handler_name = invoice.get('handler', '') or '王平'
                    # 获取手机号，默认为13924371586
                    phone_number = invoice.get('phone', '') or '13924371586'
                    
                    # 客户核签信息行 - 经办人
                    sign_info_data = [
                        [Paragraph(f"<b>客户核签:</b> {invoice.get('customer_sign', '')}", sign_title_style),
                         Paragraph(f"<b>经办人:</b> {handler_name}", right_style)]
                    ]
                    sign_info_table = Table(sign_info_data, colWidths=[12*cm, 10*cm])
                    sign_info_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    elements.append(sign_info_table)
                    
                    # 核签日期和手机号 - 手机号右对齐
                    elements.append(Spacer(1, 0.2*cm))
                    sign_detail_data = [
                        [Paragraph(f"核签日期: {invoice.get('sign_date', '')}", normal_style),
                         Paragraph(f"手机号: {phone_number}", right_style)]
                    ]
                    sign_detail_table = Table(sign_detail_data, colWidths=[12*cm, 10*cm])
                    sign_detail_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    elements.append(sign_detail_table)
                    
                    doc.build(elements)
                    
                    # 使用PyMuPDF盖章（加盖在经办人位置）
                    stamp_path = getattr(self, 'stamp_path', '')
                    if stamp_path and os.path.exists(stamp_path) and MUPDF_AVAILABLE:
                        try:
                            self._add_stamp_to_pdf(file_path, file_path, stamp_path)
                        except Exception as stamp_error:
                            print(f"盖章失败: {stamp_error}")
                    
                    messagebox.showinfo("成功", f"已导出到: {file_path}")
                    import subprocess
                    subprocess.Popen(f'"{file_path}"', shell=True)
                except Exception as e:
                    messagebox.showerror("错误", f"导出失败: {str(e)}")
        
        # 打印功能 - 生成PDF并打开，60秒后自动删除
        def print_invoice():
            if not PDF_AVAILABLE:
                messagebox.showwarning("警告", "请先安装reportlab库: pip install reportlab")
                return
            
            try:
                import tempfile
                import subprocess
                import threading
                import time as time_module
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
                
                # 注册中文字体
                font_registered = False
                font_paths = [
                    'C:/Windows/Fonts/simhei.ttf',
                    'C:/Windows/Fonts/simsun.ttc',
                    'C:/Windows/Fonts/msyh.ttc',
                    'C:/Windows/Fonts/simfang.ttf',
                    'C:/Windows/Fonts/simkai.ttf',
                    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                    '/System/Library/Fonts/PingFang.ttc',
                    '/Library/Fonts/Arial Unicode.ttf',
                ]
                
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            if font_path.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                            else:
                                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_registered = True
                            break
                        except:
                            continue
                
                # 创建字体样式（参照导出PDF按钮的样式）
                if font_registered:
                    title_style = ParagraphStyle('Title', fontName='ChineseFont', fontSize=22, leading=28, alignment=TA_CENTER, spaceAfter=8)
                    subtitle_style = ParagraphStyle('Subtitle', fontName='ChineseFont', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=6)
                    normal_style = ParagraphStyle('Normal', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                    header_style = ParagraphStyle('Header', fontName='ChineseFont', fontSize=10, leading=14, alignment=TA_CENTER)
                    cell_style = ParagraphStyle('Cell', fontName='ChineseFont', fontSize=9, leading=12, alignment=TA_CENTER)
                    right_style = ParagraphStyle('Right', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_RIGHT)
                else:
                    title_style = ParagraphStyle('Title', fontSize=22, leading=28, alignment=TA_CENTER, spaceAfter=8)
                    subtitle_style = ParagraphStyle('Subtitle', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=6)
                    normal_style = ParagraphStyle('Normal', fontSize=11, leading=16, alignment=TA_LEFT)
                    header_style = ParagraphStyle('Header', fontSize=10, leading=14, alignment=TA_CENTER)
                    cell_style = ParagraphStyle('Cell', fontSize=9, leading=12, alignment=TA_CENTER)
                    right_style = ParagraphStyle('Right', fontSize=11, leading=16, alignment=TA_RIGHT)
                
                # 使用横版A4
                temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_path = temp_file.name
                temp_file.close()
                
                doc = SimpleDocTemplate(temp_path, pagesize=landscape(A4), encoding='utf-8',
                                        leftMargin=1*cm, rightMargin=1*cm, 
                                        topMargin=1*cm, bottomMargin=0.1*cm)
                elements = []
                
                # 标题区域 - LOGO在左，公司名称在右
                logo_path = getattr(self, 'logo_path', '')
                logo_element = None
                if logo_path and os.path.exists(logo_path):
                    try:
                        logo_element = RLImage(logo_path, width=2*cm, height=2*cm)
                    except:
                        pass
                
                if logo_element:
                    header_table = Table([[logo_element, Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style)]], 
                                        colWidths=[3*cm, 18*cm])
                    header_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (0, 0), 0),
                        ('RIGHTPADDING', (0, 0), (0, 0), 10),
                    ]))
                    elements.append(header_table)
                else:
                    elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                
                elements.append(Paragraph(f"<b>{invoice.get('year_month', '')} 对账单</b>", subtitle_style))
                elements.append(Spacer(1, 0.3*cm))
                
                # 信息行
                info_data = [
                    [Paragraph(f"客户: {invoice.get('customer', '')}", normal_style),
                     Paragraph(f"制表日期: {datetime.now().strftime('%Y年%m月%d日')}", normal_style)]
                ]
                info_table = Table(info_data, colWidths=[12*cm, 10*cm])
                info_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                elements.append(info_table)
                elements.append(Spacer(1, 0.3*cm))
                
                # 表格
                table_data = []
                columns = ['序号', '订单号', '指令号', '送货日期', '送货单号', '品名', '规格', '颜色', '单位', '数量', '单价', '染费', '金额']
                header_row = [Paragraph(col, header_style) for col in columns]
                table_data.append(header_row)
                
                for item in invoice.get('items', []):
                    row = [Paragraph(str(item.get(col, '')), cell_style) for col in columns]
                    table_data.append(row)
                
                # 确保最少显示6行数据
                for i in range(max(0, 6 - len(invoice.get('items', [])))):
                    empty_row = [Paragraph('', cell_style) for _ in columns]
                    table_data.append(empty_row)
                
                col_widths = [1.2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2*cm, 1.2*cm, 1.8*cm, 1.8*cm, 1.8*cm, 2.5*cm]
                row_height = 0.8 * cm
                row_heights = [row_height] * len(table_data)
                t = Table(table_data, colWidths=col_widths, rowHeights=row_heights)
                
                header_font = 'Helvetica-Bold' if not font_registered else 'ChineseFont'
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.white),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, 0), header_font),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#f5f5f5')]),
                    ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#ecf0f1')),
                    ('FONTNAME', (0, -1), (-1, -1), header_font),
                ]))
                elements.append(t)
                
                # 备注和合计金额
                elements.append(Spacer(1, 0.3*cm))
                remark_pdf_style = ParagraphStyle('RemarkPDF', fontName='Helvetica-Bold', fontSize=11, leading=16, alignment=TA_LEFT)
                if font_registered:
                    remark_pdf_style = ParagraphStyle('RemarkPDF', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                total_style = ParagraphStyle('Total', fontName='Helvetica-Bold', fontSize=14, leading=18, alignment=TA_RIGHT)
                if font_registered:
                    total_style = ParagraphStyle('Total', fontName='ChineseFont', fontSize=14, leading=18, alignment=TA_RIGHT)
                total_amount = invoice.get('total', 0)
                remark_text = invoice.get('remark', '')
                remark_total_data = [[
                    Paragraph(f"<b>备注:</b> {remark_text}", remark_pdf_style),
                    Paragraph(f"金额合计: CNY {total_amount:.2f}", total_style)
                ]]
                remark_total_table = Table(remark_total_data, colWidths=[17*cm, 10*cm])
                remark_total_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, 0), 11),
                    ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (1, 0), (1, 0), 14),
                    ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor('#e74c3c')),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ]))
                if font_registered:
                    remark_total_table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (0, 0), 'ChineseFont'),
                        ('FONTNAME', (1, 0), (1, 0), 'ChineseFont'),
                    ]))
                elements.append(remark_total_table)
                
                # 客户核签区域
                elements.append(Spacer(1, 1*cm))
                sign_title_style = ParagraphStyle('SignTitle', fontName='Helvetica-Bold', fontSize=11, leading=16, alignment=TA_LEFT)
                if font_registered:
                    sign_title_style = ParagraphStyle('SignTitle', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT)
                
                handler_name = invoice.get('handler', '') or getattr(self, 'handler', '') or '王平'
                phone_number = invoice.get('phone', '') or '13924371586'
                
                sign_info_data = [
                    [Paragraph(f"<b>客户核签:</b> {invoice.get('customer_sign', '')}", sign_title_style),
                     Paragraph(f"<b>经办人:</b> {handler_name}", right_style)]
                ]
                sign_info_table = Table(sign_info_data, colWidths=[12*cm, 10*cm])
                sign_info_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                elements.append(sign_info_table)
                
                elements.append(Spacer(1, 0.2*cm))
                sign_detail_data = [
                    [Paragraph(f"核签日期: {invoice.get('sign_date', '')}", normal_style),
                     Paragraph(f"手机号: {phone_number}", right_style)]
                ]
                sign_detail_table = Table(sign_detail_data, colWidths=[12*cm, 10*cm])
                sign_detail_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                elements.append(sign_detail_table)
                
                doc.build(elements)
                
                # 打开PDF文档
                subprocess.Popen(f'"{temp_path}"', shell=True)
                
                # 60秒后自动删除临时文件
                def auto_delete_pdf():
                    time_module.sleep(60)
                    try:
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    except:
                        pass
                
                threading.Thread(target=auto_delete_pdf, daemon=True).start()
            except:
                pass
        
        # 按钮区域
        btn_frame = tk.Frame(container, bg='#ffffff')
        btn_frame.pack(pady=15)
        def save_invoice():
            auto_save()
            messagebox.showinfo("成功", "保存成功!")
        tk.Button(btn_frame, text="💾 保存", command=save_invoice, width=15, height=2, bg='#27ae60', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="📄 导出PDF", command=export_invoice_to_pdf, width=15, height=2, bg='#e74c3c', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="🖨️ 打印", command=print_invoice, width=15, height=2, bg='#3498db', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)

    def show_quotations(self):
        """显示报价单列表"""
        self.status_label.config(text="正在加载报价单列表...")
        self.clear_main_content()
        
        container = tk.Frame(self.main_content_frame, bg='#ffffff')
        container.pack(fill=tk.BOTH, expand=True)
        
        # 创建标签页
        notebook = ttk.Notebook(container)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # ========== 标签页1: 报价单列表 ==========
        list_tab = tk.Frame(notebook, bg='#ffffff', padx=15, pady=15)
        notebook.add(list_tab, text="📋 报价单列表")
        
        tk.Label(list_tab, text="报价单列表", font=("Microsoft YaHei", 16, "bold"), bg='#ffffff').pack(anchor=tk.W, pady=(0, 10))

        search_frame = tk.Frame(list_tab, bg='#ffffff')
        search_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Label(search_frame, text="客户:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        search_var = tk.StringVar()
        search_combo = ttk.Combobox(search_frame, textvariable=search_var, width=20, state='readonly')
        
        def refresh_customer_combo():
            try:
                customers_path = os.path.join(self.data_dir, 'customers.json')
                if os.path.exists(customers_path):
                    with open(customers_path, 'r', encoding='utf-8') as f:
                        customers_data = json.load(f)
                    customer_names = ['全部'] + [c.get('name', '') for c in customers_data if c.get('name')]
                    search_combo['values'] = customer_names
                else:
                    search_combo['values'] = ['全部']
            except:
                search_combo['values'] = ['全部']
        
        refresh_customer_combo()
        search_combo.current(0)
        search_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        columns = ('序号', '报价单号', '客户名称', '联系人', '电话', '报价日期', '总金额')
        tree = ttk.Treeview(list_tab, columns=columns, show='headings')
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=120, anchor='w')
        tree.column('序号', width=50, anchor='center')
        tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        scrollbar = ttk.Scrollbar(list_tab, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        
        def refresh_list():
            for item in tree.get_children():
                tree.delete(item)
            search_text = search_var.get()
            display_idx = 0
            for quote in self.quotations:
                if search_text == '全部' or not search_text or search_text == quote.get('customer', ''):
                    display_idx += 1
                    tree.insert('', tk.END, values=(
                        display_idx, quote.get('quote_no', ''), quote.get('customer', ''),
                        quote.get('contact', ''), quote.get('phone', ''), quote.get('date', ''),
                        f"¥{quote.get('total', 0):.2f}"
                    ))
        
        search_combo.bind('<<ComboboxSelected>>', lambda e: refresh_list())
        refresh_list()

        # 详情标签页
        detail_tab = tk.Frame(notebook, bg='#ffffff', padx=15, pady=15)
        notebook.add(detail_tab, text="📄 报价单详情")
        
        detail_content_frame = tk.Frame(detail_tab, bg='#ffffff')
        detail_content_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(detail_content_frame, text="请从左侧列表选择一条报价单查看详情", 
                 font=("Microsoft YaHei", 12), bg='#ffffff', fg='#999999').pack(pady=50)

        def view_detail():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            item = tree.item(selection[0])
            idx = item['values'][0] - 1
            self._view_quotation_detail_in_tab(idx, detail_content_frame, notebook, close_callback=lambda: notebook.select(0))
            notebook.select(1)  # 切换到详情标签页

        def delete_quotation():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            if messagebox.askyesno("确认", "确定要删除这条报价单吗?"):
                item = tree.item(selection[0])
                idx = item['values'][0] - 1
                del self.quotations[idx]
                self.save_data()
                self.update_stats()
                refresh_list()
                messagebox.showinfo("成功", "删除成功!")

        btn_frame = tk.Frame(list_tab)
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="查看详情", command=view_detail, width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="删除报价", command=delete_quotation, width=15, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
        
        # 双击打开详情
        tree.bind("<Double-Button-1>", lambda e: view_detail())

    def _view_quotation_detail_in_tab(self, idx, parent_frame, notebook, close_callback=None):
        """在标签页内显示报价单详情（左右布局，左6右2）"""
        # 清空父框架
        for widget in parent_frame.winfo_children():
            widget.destroy()
        
        quote = self.quotations[idx]
        
        # 创建左右分栏布局容器
        content_container = tk.Frame(parent_frame, bg='#ffffff')
        content_container.pack(fill=tk.BOTH, expand=True)
        
        # 左侧主内容区域（占6份）
        left_frame = tk.Frame(content_container, bg='#ffffff')
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 右侧辅助区域（占2份）- 已签报价单图片
        right_frame = tk.Frame(content_container, bg='#f8f9fa', width=250)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=(10, 0))
        right_frame.pack_propagate(False)
        
        # 已签报价单标题
        signed_header = tk.Label(right_frame, text="📎 已签报价单", font=("Microsoft YaHei", 11, "bold"),
                                bg='#f8f9fa', fg='#2c3e50')
        signed_header.pack(pady=(10, 5))
        
        # 已签报价单路径变量
        signed_quotation_path = tk.StringVar(value=quote.get('signed_image', ''))
        
        # A4比例预览区域 (1:1.414)
        preview_border = tk.Frame(right_frame, bg='#e0e0e0', padx=2, pady=2)
        preview_border.pack(padx=10, pady=5)
        preview_inner = tk.Frame(preview_border, bg='#ffffff')
        preview_inner.pack(fill=tk.BOTH, expand=True)
        
        # A4比例预览标签
        signed_preview_label = tk.Label(preview_inner, text="暂无图片\n\n双击打开图片",
                                       font=("Microsoft YaHei", 9), bg='#ffffff', fg='#999999')
        signed_preview_label.pack(pady=10)
        
        # 加载已有图片
        def load_signed_image():
            img_path = quote.get('signed_image', '')
            if img_path and os.path.exists(img_path):
                signed_quotation_path.set(img_path)
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(img_path)
                    # A4比例预览 (大约200:283像素)
                    img.thumbnail((180, 254), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    signed_preview_label.config(image=photo, text="", bg='#ffffff')
                    signed_preview_label.image = photo
                except Exception as e:
                    print(f"加载图片失败: {e}")
        
        load_signed_image()
        
        # 双击打开图片
        def open_signed_image(event):
            img_path = signed_quotation_path.get()
            if img_path and os.path.exists(img_path):
                import subprocess
                subprocess.Popen(f'"{img_path}"', shell=True)
        
        signed_preview_label.bind('<Double-Button-1>', open_signed_image)
        preview_inner.bind('<Double-Button-1>', open_signed_image)
        preview_border.bind('<Double-Button-1>', open_signed_image)
        
        # 选择图片按钮
        def select_signed_quotation():
            file_path = filedialog.askopenfilename(
                title="选择已签报价单图片",
                filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif *.webp"), ("所有文件", "*.*")]
            )
            if file_path:
                signed_quotation_path.set(file_path)
                quote['signed_image'] = file_path
                self.save_data()
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(file_path)
                    img.thumbnail((180, 254), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    signed_preview_label.config(image=photo, text="", bg='#ffffff')
                    signed_preview_label.image = photo
                except Exception as e:
                    print(f"加载图片失败: {e}")
        
        # 图片路径文本框
        path_frame = tk.Frame(right_frame, bg='#f8f9fa')
        path_frame.pack(fill=tk.X, padx=10, pady=(5, 5))
        tk.Label(path_frame, text="路径:", font=("Microsoft YaHei", 8), bg='#f8f9fa', fg='#7f8c8d').pack(anchor=tk.W)
        path_entry = tk.Entry(path_frame, textvariable=signed_quotation_path, width=28,
                              font=("Microsoft YaHei", 8))
        path_entry.pack(fill=tk.X)
        
        # 选择图片按钮
        def select_signed_quotation():
            file_path = filedialog.askopenfilename(
                title="选择已签报价单图片",
                filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif *.webp"), ("所有文件", "*.*")]
            )
            if file_path:
                signed_quotation_path.set(file_path)
                quote['signed_image'] = file_path
                self.save_data()
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(file_path)
                    img.thumbnail((180, 254), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    signed_preview_label.config(image=photo, text="", bg='#ffffff')
                    signed_preview_label.image = photo
                except Exception as e:
                    print(f"加载图片失败: {e}")
        
        tk.Button(right_frame, text="📷 选择图片", command=select_signed_quotation,
                 font=("Microsoft YaHei", 10), bg='#3498db', fg='white',
                 relief=tk.FLAT, padx=15, pady=8, cursor='hand2').pack(pady=(5, 3))
        
        # 打开图片按钮
        def open_signed_image_action():
            img_path = signed_quotation_path.get()
            if img_path and os.path.exists(img_path):
                import subprocess
                subprocess.Popen(f'"{img_path}"', shell=True)
            else:
                messagebox.showwarning("警告", "请先选择已签报价单图片")
        
        tk.Button(right_frame, text="🟢 打开图片", command=open_signed_image_action,
                 font=("Microsoft YaHei", 10), bg='#27ae60', fg='white',
                 relief=tk.FLAT, padx=15, pady=8, cursor='hand2').pack(pady=(5, 3))
        
        # 清除图片按钮
        def clear_signed_image():
            result = messagebox.askyesno("确认", "是否清除已签报价单图片？")
            if result:
                signed_quotation_path.set('')
                quote['signed_image'] = ''
                self.save_data()
                signed_preview_label.config(image='', text="暂无图片\n\n双击打开图片", bg='#ffffff')
                signed_preview_label.image = None
        
        tk.Button(right_frame, text="🗑️ 清除图片", command=clear_signed_image,
                 font=("Microsoft YaHei", 10), bg='#e74c3c', fg='white',
                 relief=tk.FLAT, padx=15, pady=8, cursor='hand2').pack(pady=(3, 10))
        
        # 图片预览区域右键菜单
        img_context_menu = tk.Menu(preview_inner, tearoff=0)
        img_context_menu.add_command(label="📂 打开图片", command=open_signed_image_action)
        img_context_menu.add_separator()
        img_context_menu.add_command(label="🗑️ 清除图片", command=clear_signed_image)
        
        def show_img_context_menu(event):
            img_context_menu.post(event.x_root, event.y_root)
        
        signed_preview_label.bind('<Button-3>', show_img_context_menu)
        preview_inner.bind('<Button-3>', show_img_context_menu)
        preview_border.bind('<Button-3>', show_img_context_menu)
        
        # 顶部操作栏（包含按钮）
        top_bar = tk.Frame(left_frame, bg='#ffffff')
        top_bar.pack(fill=tk.X, pady=(0, 5))
        
        # 左侧功能按钮
        btn_left = tk.Frame(top_bar, bg='#ffffff')
        btn_left.pack(side=tk.LEFT)
        
        def export_to_excel():
            if not WORD_AVAILABLE:
                messagebox.showwarning("警告", "请先安装python-docx库: pip install python-docx")
                return
            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel文件", "*.xlsx")])
            if file_path:
                try:
                    from docx import Document
                    from docx.shared import Pt, Cm
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    # 标题
                    title = doc.add_heading(f"{self.company_name or '公司'}报价单", 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 基本信息
                    doc.add_paragraph(f"客户: {quote.get('customer', '')}")
                    doc.add_paragraph(f"联系人: {quote.get('contact', '')}")
                    doc.add_paragraph(f"日期: {quote.get('date', '')}")
                    
                    # 表格
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '编号'
                    hdr_cells[1].text = '品名'
                    hdr_cells[2].text = '规格'
                    hdr_cells[3].text = '单价/码'
                    
                    for item in quote.get('items', []):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(item.get('编号', ''))
                        row_cells[1].text = item.get('品名', '')
                        row_cells[2].text = item.get('规格', '')
                        row_cells[3].text = str(item.get('单价/码', ''))
                    
                    # 备注
                    if quote.get('remark'):
                        doc.add_paragraph('\n备注:')
                        doc.add_paragraph(quote.get('remark'))
                    
                    doc.save(file_path)
                    messagebox.showinfo("成功", f"已导出到: {file_path}")
                except Exception as e:
                    messagebox.showerror("错误", f"导出失败: {str(e)}")
        
        def export_to_pdf():
            if not PDF_AVAILABLE:
                messagebox.showwarning("警告", "请先安装reportlab库: pip install reportlab")
                return
            
            # 构建文件名：客户+品名-报价单
            customer_name = quote.get('customer', '')
            # 获取品名（取第一个商品的品名）
            product_name = ''
            items = quote.get('items', [])
            if items and len(items) > 0:
                first_item = items[0]
                product_name = first_item.get('品名', '')
            
            # 生成文件名：客户+品名-报价单
            if customer_name and product_name:
                default_filename = f"{customer_name}{product_name}-报价单.pdf"
            elif customer_name:
                default_filename = f"{customer_name}-报价单.pdf"
            elif product_name:
                default_filename = f"{product_name}-报价单.pdf"
            else:
                default_filename = "报价单.pdf"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf", 
                filetypes=[("PDF文件", "*.pdf")],
                initialfile=default_filename
            )
            if file_path:
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.units import cm
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib import colors
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont
                    from reportlab.lib.enums import TA_CENTER, TA_LEFT
                    
                    # 注册中文字体（嵌入字体确保跨平台显示）
                    font_registered = False
                    font_paths = [
                        # Windows常见中文字体
                        'C:/Windows/Fonts/simhei.ttf',   # 黑体
                        'C:/Windows/Fonts/simsun.ttc',   # 宋体
                        'C:/Windows/Fonts/msyh.ttc',     # 微软雅黑
                        'C:/Windows/Fonts/simfang.ttf',  # 仿宋
                        'C:/Windows/Fonts/simkai.ttf',   # 楷体
                        # Linux常见中文字体
                        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                        # macOS常见中文字体
                        '/System/Library/Fonts/PingFang.ttc',
                        '/Library/Fonts/Arial Unicode.ttf',
                    ]
                    
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                # 尝试注册不同格式
                                if font_path.endswith('.ttc'):
                                    # TTC字体文件可能包含多个字体
                                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                                else:
                                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                                font_registered = True
                                break
                            except Exception as e:
                                print(f"注册字体 {font_path} 失败: {e}")
                                continue
                    
                    # 创建中文字体样式
                    if font_registered:
                        title_style = ParagraphStyle(
                            'ChineseTitle',
                            fontName='ChineseFont',
                            fontSize=18,
                            leading=24,
                            alignment=TA_CENTER,
                            spaceAfter=12
                        )
                        normal_style = ParagraphStyle(
                            'ChineseNormal',
                            fontName='ChineseFont',
                            fontSize=11,
                            leading=16,
                            alignment=TA_LEFT,
                            spaceAfter=6
                        )
                        table_header_style = ParagraphStyle(
                            'ChineseTableHeader',
                            fontName='ChineseFont',
                            fontSize=11,
                            leading=14,
                            alignment=TA_CENTER
                        )
                        table_cell_style = ParagraphStyle(
                            'ChineseTableCell',
                            fontName='ChineseFont',
                            fontSize=10,
                            leading=14,
                            alignment=TA_CENTER
                        )
                        bold_style = ParagraphStyle(
                            'ChineseBold',
                            fontName='ChineseFont',
                            fontSize=12,
                            leading=18,
                            alignment=TA_LEFT,
                            spaceAfter=6
                        )
                    else:
                        # 如果没有中文字体，使用默认样式并提示
                        title_style = ParagraphStyle('Title', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=12)
                        normal_style = ParagraphStyle('Normal', fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=6)
                        table_header_style = ParagraphStyle('Header', fontSize=11, leading=14, alignment=TA_CENTER)
                        table_cell_style = ParagraphStyle('Cell', fontSize=10, leading=14, alignment=TA_CENTER)
                        bold_style = ParagraphStyle('Bold', fontSize=12, leading=18, alignment=TA_LEFT, spaceAfter=6)
                    
                    doc = SimpleDocTemplate(
                        file_path, 
                        pagesize=A4,
                        encoding='utf-8'  # 明确指定UTF-8编码
                    )
                    elements = []
                    
                    # 标题 - 公司LOGO和名称（LOGO向右移动4厘米）
                    logo_path = getattr(self, 'logo_path', '')
                    if logo_path and os.path.exists(logo_path):
                        try:
                            from reportlab.platypus import Image as RLImage
                            logo_img = RLImage(logo_path, width=2*cm, height=2*cm)
                            logo_table = Table([[logo_img, Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style)]], colWidths=[3*cm, 12*cm])
                            logo_table.setStyle(TableStyle([
                                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('LEFTPADDING', (0, 0), (0, 0), 0),
                                ('RIGHTPADDING', (0, 0), (0, 0), 10),
                            ]))
                            elements.append(logo_table)
                        except Exception as e:
                            print(f"加载LOGO图片失败: {e}")
                            elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                    else:
                        elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                    elements.append(Paragraph(f"<b>报价单</b>", title_style))
                    elements.append(Spacer(1, 0.5*cm))
                    
                    # 信息 - 客户、联系人、日期在同一行，日期右对齐
                    info_data = [
                        [Paragraph(f"客户: {quote.get('customer', '')}", normal_style),
                         Paragraph(f"联系人: {quote.get('contact', '')}", normal_style),
                         Paragraph(f"日期: {quote.get('date', '')}", normal_style)]
                    ]
                    info_table = Table(info_data, colWidths=[6*cm, 4*cm, 5*cm])
                    info_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
                        ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ]))
                    elements.append(info_table)
                    elements.append(Spacer(1, 0.3*cm))
                    
                    # 表格数据（使用Paragraph支持中文）
                    table_data = []
                    header_row = [
                        Paragraph('编号', table_header_style),
                        Paragraph('品名', table_header_style),
                        Paragraph('规格', table_header_style),
                        Paragraph('单价/码', table_header_style)
                    ]
                    table_data.append(header_row)
                    
                    for item in quote.get('items', []):
                        table_data.append([
                            Paragraph(str(item.get('编号', '')), table_cell_style),
                            Paragraph(str(item.get('品名', '')), table_cell_style),
                            Paragraph(str(item.get('规格', '')), table_cell_style),
                            Paragraph(str(item.get('单价/码', '')), table_cell_style)
                        ])
                    
                    # 创建表格
                    t = Table(table_data, colWidths=[2*cm, 6*cm, 4*cm, 3*cm])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.white),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        ('TOPPADDING', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                        ('TOPPADDING', (0, 1), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#ffffff'), colors.HexColor('#f5f5f5')])
                    ]))
                    elements.append(t)
                    
                    # 备注
                    if quote.get('remark'):
                        elements.append(Spacer(1, 0.5*cm))
                        elements.append(Paragraph("<b>备注:</b>", bold_style))
                        remark_lines = quote.get('remark').split('\n')
                        for line in remark_lines:
                            if line.strip():
                                elements.append(Paragraph(line, normal_style))
                    
                    # 签名区域：客户负责人和报价人下移4行
                    elements.append(Spacer(1, 2*cm))
                    
                    # 签名行 - 客户负责人（左下角）和报价人（右下角）同一行
                    # 使用更大的列宽，让两者分布在页面左右两端，右移1.5厘米
                    sign_data = [[
                        Paragraph("客户负责人: ____________", normal_style),
                        Paragraph(f"报价人: {getattr(self, 'quoter', '') or '____________'}", normal_style)
                    ]]
                    sign_table = Table(sign_data, colWidths=[12*cm, 7*cm])
                    sign_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 1.5*cm),
                    ]))
                    elements.append(sign_table)


                    
                    # 如果字体未注册，添加提示
                    if not font_registered:
                        elements.append(Spacer(1, 0.5*cm))
                        warning_style = ParagraphStyle('Warning', fontSize=9, textColor=colors.orange)
                        elements.append(Paragraph("⚠️ 提示: 未检测到中文字体，中文可能无法正常显示。请安装中文字体后重新导出。", warning_style))
                    
                    doc.build(elements)
                    
                    # ✅ 使用PyMuPDF添加公章功能
                    stamp_path = getattr(self, 'stamp_path', '')
                    if stamp_path and os.path.exists(stamp_path):
                        try:
                            if not MUPDF_AVAILABLE:
                                messagebox.showwarning("提示", "请先安装PyMuPDF库以支持盖章功能\npip install pymupdf")
                            else:
                                import fitz  # PyMuPDF
                                
                                # ========== 关键控制逻辑 ==========
                                # 公章大小：4CM × 4CM
                                stamp_size = 4 * cm
                                
                                # 签名区域位置估算（基于A4纸张尺寸）
                                # 报价人签名区域在右下角
                                sign_x_approx = 15 * cm  # 签名x坐标（靠右）
                                sign_y_approx = 4 * cm   # 签名y坐标（距底部）
                                
                                # overlap控制重合度：0.3 表示印章30%的高度压在文字上
                                # stamp_y = sign_y + (stamp_size * overlap)
                                overlap = 0.3
                                stamp_y = sign_y_approx + (stamp_size * overlap)
                                
                                # stamp_x 用签名x坐标减去印章一半宽度，实现水平居中
                                stamp_x = sign_x_approx - (stamp_size / 2)
                                
                                # 获取原始PDF的页面尺寸
                                src_doc = fitz.open(file_path)
                                first_page = src_doc[0]
                                page_width = first_page.rect.width
                                page_height = first_page.rect.height
                                
                                # 确保印章在页面范围内
                                if stamp_x < 0:
                                    stamp_x = 1 * cm
                                if stamp_x + stamp_size > page_width:
                                    stamp_x = page_width - stamp_size - 1 * cm
                                if stamp_y < 0:
                                    stamp_y = 1 * cm
                                if stamp_y + stamp_size > page_height:
                                    stamp_y = page_height - stamp_size - 1 * cm
                                
                                # 使用PIL处理公章图片（保留透明通道）
                                from PIL import Image
                                stamp_img = Image.open(stamp_path).convert('RGBA')
                                
                                # 应用透明度：0.7 表示70%不透明度
                                alpha = stamp_img.split()[3]  # 获取Alpha通道
                                alpha = alpha.point(lambda p: int(p * 0.7))  # 70%强度
                                stamp_img.putalpha(alpha)
                                
                                # 保存透明图片到临时PNG文件
                                import tempfile
                                transparent_temp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                                transparent_path = transparent_temp.name
                                transparent_temp.close()
                                stamp_img.save(transparent_path, 'PNG')
                                
                                # 使用PyMuPDF添加公章
                                # stamp_x, stamp_y是左下角坐标，需要转换为PyMuPDF的左上角坐标
                                # PyMuPDF使用左下角为原点，y轴向上
                                mupdf_stamp_x = stamp_x
                                mupdf_stamp_y = page_height - stamp_y - stamp_size  # 转换y坐标
                                
                                # 创建图片矩形区域
                                stamp_rect = fitz.Rect(
                                    mupdf_stamp_x, 
                                    mupdf_stamp_y, 
                                    mupdf_stamp_x + stamp_size, 
                                    mupdf_stamp_y + stamp_size
                                )
                                
                                # 在PDF页面上插入图片（带透明度）
                                first_page.insert_image(stamp_rect, filename=transparent_path)
                                
                                # 保存PDF
                                src_doc.save(file_path + '.tmp')
                                src_doc.close()
                                
                                # 替换原文件
                                import shutil
                                shutil.move(file_path + '.tmp', file_path)
                                
                                # 删除临时透明图片
                                try:
                                    os.remove(transparent_path)
                                except:
                                    pass
                                
                                print(f"✓ PyMuPDF公章添加成功: overlap={overlap}, alpha=0.7")
                                print(f"  印章位置: x={stamp_x/cm:.2f}cm, y={stamp_y/cm:.2f}cm")
                                print(f"  印章大小: {stamp_size/cm:.1f}cm × {stamp_size/cm:.1f}cm")
                        except Exception as stamp_error:
                            print(f"添加公章失败: {stamp_error}")
                    
                    messagebox.showinfo("成功", f"已导出到: {file_path}\n编码: UTF-8")
                    # 导出成功后直接打开PDF
                    import subprocess
                    subprocess.Popen(f'"{file_path}"', shell=True)
                except Exception as e:
                    messagebox.showerror("错误", f"导出失败: {str(e)}")
        
        def print_quotation():
            """打印报价单 - 直接打印，不弹窗"""
            if not PDF_AVAILABLE:
                return
            
            try:
                import tempfile
                import subprocess
                import threading
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.enums import TA_CENTER, TA_LEFT
                
                # 注册中文字体
                font_registered = False
                font_paths = [
                    'C:/Windows/Fonts/simhei.ttf',
                    'C:/Windows/Fonts/simsun.ttc',
                    'C:/Windows/Fonts/msyh.ttc',
                    'C:/Windows/Fonts/simfang.ttf',
                    'C:/Windows/Fonts/simkai.ttf',
                    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                    '/System/Library/Fonts/PingFang.ttc',
                    '/Library/Fonts/Arial Unicode.ttf',
                ]
                
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            if font_path.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                            else:
                                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_registered = True
                            break
                        except:
                            continue
                
                # 创建中文字体样式
                if font_registered:
                    title_style = ParagraphStyle('ChineseTitle', fontName='ChineseFont', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=12)
                    normal_style = ParagraphStyle('ChineseNormal', fontName='ChineseFont', fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=6)
                    table_header_style = ParagraphStyle('ChineseTableHeader', fontName='ChineseFont', fontSize=11, leading=14, alignment=TA_CENTER)
                    table_cell_style = ParagraphStyle('ChineseTableCell', fontName='ChineseFont', fontSize=10, leading=14, alignment=TA_CENTER)
                    bold_style = ParagraphStyle('ChineseBold', fontName='ChineseFont', fontSize=12, leading=18, alignment=TA_LEFT, spaceAfter=6)
                else:
                    title_style = ParagraphStyle('Title', fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=12)
                    normal_style = ParagraphStyle('Normal', fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=6)
                    table_header_style = ParagraphStyle('Header', fontSize=11, leading=14, alignment=TA_CENTER)
                    table_cell_style = ParagraphStyle('Cell', fontSize=10, leading=14, alignment=TA_CENTER)
                    bold_style = ParagraphStyle('Bold', fontSize=12, leading=18, alignment=TA_LEFT, spaceAfter=6)
                
                # 创建临时PDF文件
                temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_path = temp_file.name
                temp_file.close()
                
                doc = SimpleDocTemplate(temp_path, pagesize=A4, encoding='utf-8')
                elements = []
                
                # 标题 - 公司LOGO和名称
                logo_path = getattr(self, 'logo_path', '')
                if logo_path and os.path.exists(logo_path):
                    try:
                        logo_img = RLImage(logo_path, width=2*cm, height=2*cm)
                        logo_table = Table([[logo_img, Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style)]], colWidths=[3*cm, 12*cm])
                        logo_table.setStyle(TableStyle([
                            ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('LEFTPADDING', (0, 0), (0, 0), 0),
                            ('RIGHTPADDING', (0, 0), (0, 0), 10),
                        ]))
                        elements.append(logo_table)
                    except:
                        elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                else:
                    elements.append(Paragraph(f"<b>{self.company_name or '公司'}</b>", title_style))
                elements.append(Paragraph(f"<b>报价单</b>", title_style))
                elements.append(Spacer(1, 0.5*cm))
                
                # 信息行
                info_data = [
                    [Paragraph(f"客户: {quote.get('customer', '')}", normal_style),
                     Paragraph(f"联系人: {quote.get('contact', '')}", normal_style),
                     Paragraph(f"日期: {quote.get('date', '')}", normal_style)]
                ]
                info_table = Table(info_data, colWidths=[6*cm, 4*cm, 5*cm])
                info_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'LEFT'),
                    ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ]))
                elements.append(info_table)
                elements.append(Spacer(1, 0.3*cm))
                
                # 表格数据
                table_data = []
                header_row = [
                    Paragraph('编号', table_header_style),
                    Paragraph('品名', table_header_style),
                    Paragraph('规格', table_header_style),
                    Paragraph('单价/码', table_header_style)
                ]
                table_data.append(header_row)
                
                for item in quote.get('items', []):
                    table_data.append([
                        Paragraph(str(item.get('编号', '')), table_cell_style),
                        Paragraph(str(item.get('品名', '')), table_cell_style),
                        Paragraph(str(item.get('规格', '')), table_cell_style),
                        Paragraph(str(item.get('单价/码', '')), table_cell_style)
                    ])
                
                # 创建表格
                t = Table(table_data, colWidths=[2*cm, 6*cm, 4*cm, 3*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.white),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, 0), 10),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#ffffff'), colors.HexColor('#f5f5f5')])
                ]))
                elements.append(t)
                
                # 备注
                if quote.get('remark'):
                    elements.append(Spacer(1, 0.5*cm))
                    elements.append(Paragraph("<b>备注:</b>", bold_style))
                    remark_lines = quote.get('remark').split('\n')
                    for line in remark_lines:
                        if line.strip():
                            elements.append(Paragraph(line, normal_style))
                
                # 签名区域 - 打印版本不加载公章图片，客户负责人左下角，报价人右下角
                elements.append(Spacer(1, 0.5*cm))
                
                # 客户负责人 - 左下角
                customer_manager = Paragraph("客户负责人: ____________", normal_style)
                customer_table = Table([[customer_manager]], colWidths=[8*cm])
                customer_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                ]))
                
                # 报价人 - 右下角
                quoter_name = getattr(self, 'quoter', '') or '____________'
                quoter = Paragraph(f"报价人: {quoter_name}", normal_style)
                quoter_table = Table([[quoter]], colWidths=[7*cm])
                quoter_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                ]))
                
                # 两行布局：客户负责人在上左，报价人在下右
                sign_table = Table([
                    [customer_table, quoter_table]
                ], colWidths=[8*cm, 7*cm])
                sign_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                ]))
                elements.append(sign_table)
                
                doc.build(elements)
                
                # 打开PDF并60秒后自动删除
                def open_and_delete():
                    import time as time_module
                    try:
                        subprocess.Popen(f'"{temp_path}"', shell=True)
                    except:
                        pass
                    time_module.sleep(60)
                    try:
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    except:
                        pass
                
                threading.Thread(target=open_and_delete, daemon=True).start()
            except:
                pass
        
        tk.Button(btn_left, text="🖨️ 打印", command=print_quotation,
                  font=("Microsoft YaHei", 10), bg='#3498db', fg='white',
                  relief=tk.FLAT, padx=12, pady=5).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(btn_left, text="📄 导出为PDF", command=export_to_pdf,
                  font=("Microsoft YaHei", 10), bg='#e74c3c', fg='white',
                  relief=tk.FLAT, padx=12, pady=5).pack(side=tk.LEFT, padx=(0, 5))

        # 公司LOGO和名称
        header_frame = tk.Frame(left_frame, bg='#ffffff')
        header_frame.pack(pady=(5, 5))
        if getattr(self, 'logo_path', '') and os.path.exists(self.logo_path):
            try:
                from PIL import Image, ImageTk
                img = Image.open(self.logo_path)
                img.thumbnail((50, 50), Image.Resampling.LANCZOS)
                logo_photo = ImageTk.PhotoImage(img)
                logo_label = tk.Label(header_frame, image=logo_photo, bg='#ffffff')
                logo_label.image = logo_photo
                logo_label.pack(side=tk.LEFT, padx=(0, 10))
            except Exception as e:
                print(f"加载LOGO失败: {e}")
        company_name = self.company_name or "公司名称"
        company_label = tk.Label(header_frame, text=company_name, font=("Microsoft YaHei", 16, "bold"),
                                bg='#ffffff', fg='#2c3e50')
        company_label.pack(side=tk.LEFT)

        # 报价单标题
        title_label = tk.Label(left_frame, text="报价单", font=("Microsoft YaHei", 18, "bold"),
                              bg='#ffffff', fg='#e74c3c')
        title_label.pack(pady=(0, 15))

        # 客户、联系人、日期区域
        info_frame = tk.Frame(left_frame, bg='#ffffff')
        info_frame.pack(fill=tk.X, pady=(0, 15))

        # 客户名称
        customer_frame = tk.Frame(info_frame, bg='#ffffff')
        customer_frame.pack(side=tk.LEFT, padx=(0, 20))
        tk.Label(customer_frame, text="客户:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        customer_value = tk.Label(customer_frame, text=quote.get('customer', ''), 
                                  font=("Microsoft YaHei", 11), bg='#ffffff', fg='#2c3e50')
        customer_value.pack(side=tk.LEFT, padx=(5, 0))

        # 联系人
        contact_frame = tk.Frame(info_frame, bg='#ffffff')
        contact_frame.pack(side=tk.LEFT, padx=(0, 20))
        tk.Label(contact_frame, text="联系人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        contact_value = tk.Label(contact_frame, text=quote.get('contact', ''), 
                                 font=("Microsoft YaHei", 11), bg='#ffffff')
        contact_value.pack(side=tk.LEFT, padx=(5, 0))

        # 日期
        date_frame = tk.Frame(info_frame, bg='#ffffff')
        date_frame.pack(side=tk.RIGHT)
        tk.Label(date_frame, text="日期:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        date_value = tk.Label(date_frame, text=quote.get('date', ''), 
                              font=("Microsoft YaHei", 11), bg='#ffffff')
        date_value.pack(side=tk.LEFT, padx=(5, 0))

        # 商品列表表格
        columns = ('编号', '品名', '规格', '单价/码')
        tree = ttk.Treeview(left_frame, columns=columns, show='headings', height=4)

        tree.heading('编号', text='编号')
        tree.heading('品名', text='品名')
        tree.heading('规格', text='规格')
        tree.heading('单价/码', text='单价/码')

        tree.column('编号', width=80, anchor='center')
        tree.column('品名', width=250, anchor='center')
        tree.column('规格', width=150, anchor='center')
        tree.column('单价/码', width=150, anchor='center')

        tree.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)

        # 填充数据
        for item in quote.get('items', []):
            tree.insert('', tk.END, values=(
                item.get('编号', ''),
                item.get('品名', ''),
                item.get('规格', ''),
                item.get('单价/码', '')
            ))

        # 上方备注文本框（单行）
        top_remark_frame = tk.Frame(left_frame, bg='#ffffff')
        top_remark_frame.pack(fill=tk.X, pady=(15, 10))
        tk.Label(top_remark_frame, text="备注:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        top_remark_value = tk.Label(top_remark_frame, text=quote.get('remark', '').split('\n')[0] if quote.get('remark') else '', 
                                    font=("Microsoft YaHei", 11), bg='#ffffff')
        top_remark_value.pack(side=tk.LEFT, padx=(5, 0))

        # 下方多行备注文本框
        remark_frame = tk.Frame(left_frame, bg='#ffffff')
        remark_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        remark_text = tk.Text(remark_frame, width=80, height=10, font=("Microsoft YaHei", 11), wrap=tk.WORD)
        remark_text.pack(fill=tk.BOTH, expand=True, padx=(0, 0))
        remark_text.insert('1.0', quote.get('remark', "1. 以上为本公司统一价格\nAll prices listed above are our company's standard quotations.\n\n2. 颜色可以根据客户要求生产，一般不另加价。\nColors can be produced as per customer's requirements, with generally no additional charge.\n\n3. 本公司负责为客户送到指定地点。\nWe will arrange delivery of the goods to the designated location specified by the customer.\n4. 本公司结算方式按约定。\nPayment terms shall be as mutually agreed upon by both parties.\n5. 本报价单一经客户确认后，即作为定单附件。\nOnce confirmed by the customer, this quotation shall become an integral part of the corresponding purchase order (PO)."))
        remark_text.config(state='disabled')

        # 公章图片区域
        stamp_frame = tk.Frame(left_frame, bg='#ffffff')
        stamp_frame.pack(fill=tk.X, pady=(30, 10))
        
        # 左侧客户负责人
        left_bottom = tk.Frame(stamp_frame, bg='#ffffff')
        left_bottom.pack(side=tk.LEFT)
        tk.Label(left_bottom, text="客户负责人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        manager_value = tk.Label(left_bottom, text="____________", font=("Microsoft YaHei", 11), bg='#ffffff')
        manager_value.pack(side=tk.LEFT, padx=(5, 0))

        # 右侧报价人和公章
        right_bottom = tk.Frame(stamp_frame, bg='#ffffff')
        right_bottom.pack(side=tk.RIGHT)
        
        # 报价人
        quoter_frame = tk.Frame(right_bottom, bg='#ffffff')
        quoter_frame.pack(side=tk.TOP)
        tk.Label(quoter_frame, text="报价人:", font=("Microsoft YaHei", 11), bg='#ffffff').pack(side=tk.LEFT)
        quoter_value = tk.Label(quoter_frame, text=getattr(self, 'quoter', '') or "____________", 
                                 font=("Microsoft YaHei", 11), bg='#ffffff')
        quoter_value.pack(side=tk.LEFT, padx=(5, 0))
        
        # 公章图片
        if getattr(self, 'stamp_path', '') and os.path.exists(self.stamp_path):
            try:
                from PIL import Image, ImageTk
                img = Image.open(self.stamp_path)
                img.thumbnail((80, 80), Image.Resampling.LANCZOS)
                stamp_photo = ImageTk.PhotoImage(img)
                stamp_label = tk.Label(right_bottom, image=stamp_photo, bg='#ffffff')
                stamp_label.image = stamp_photo
                stamp_label.pack(side=tk.TOP, pady=(10, 0))
            except Exception as e:
                print(f"加载公章图片失败: {e}")
    
    def _view_quotation_detail(self, idx, parent_window):
        """查看报价单详情（弹窗版本，保留兼容性）"""
        detail_window = tk.Toplevel(parent_window)
        quote = self.quotations[idx]
        detail_window.title(f"报价单详情 - {quote.get('quote_no', '')}")
        detail_window.geometry("800x500")
        info_frame = tk.Frame(detail_window, padx=20, pady=20)
        info_frame.pack(fill=tk.BOTH, expand=True)
        tk.Label(info_frame, text=f"报价单号: {quote.get('quote_no', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"客户名称: {quote.get('customer', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"联系人: {quote.get('contact', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"电话: {quote.get('phone', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"报价日期: {quote.get('date', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"有效期: {quote.get('validity', '')} 天", font=("Microsoft YaHei", 12)).pack(anchor=tk.W)
        tk.Label(info_frame, text=f"合计金额: ¥{quote.get('total', 0):.2f}", font=("Microsoft YaHei", 14, "bold"), fg='red').pack(anchor=tk.W)
        tk.Label(info_frame, text=f"备注: {quote.get('remark', '')}", font=("Microsoft YaHei", 11)).pack(anchor=tk.W)
        if quote.get('items'):
            tk.Label(info_frame, text="\n报价明细:", font=("Microsoft YaHei", 12, "bold")).pack(anchor=tk.W)
            items_text = "\n".join([f"{i+1}. {item['项目名称']} ({item['规格']}) x {item['数量']} {item['单位']} = ¥{item['合计']}" for i, item in enumerate(quote.get('items', []))])
            tk.Label(info_frame, text=items_text, font=("Microsoft YaHei", 10), justify=tk.LEFT).pack(anchor=tk.W)

    # -------------------- 付款管理 --------------------
    def show_payments(self):
        """显示付款列表"""
        self.status_label.config(text="正在加载付款列表...")
        list_window = tk.Toplevel(self.root)
        list_window.title("付款列表")
        list_window.geometry("1000x600")
        tk.Label(list_window, text="付款列表", font=("Microsoft YaHei", 16, "bold")).pack(pady=10)

        search_frame = tk.Frame(list_window)
        search_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(search_frame, text="搜索客户:").pack(side=tk.LEFT)
        search_entry = tk.Entry(search_frame, width=20)
        search_entry.pack(side=tk.LEFT, padx=5)
        def refresh_list():
            for item in tree.get_children():
                tree.delete(item)
            search_text = search_entry.get().lower()
            for idx, payment in enumerate(self.payments, 1):
                if not search_text or search_text in payment.get('customer', '').lower():
                    tree.insert('', tk.END, values=(
                        idx, payment.get('payment_no', ''), payment.get('customer', ''),
                        f"¥{payment.get('amount', 0):.2f}", payment.get('payment_date', ''),
                        payment.get('payment_method', ''), payment.get('status', '')
                    ))
        tk.Button(search_frame, text="刷新列表", command=refresh_list).pack(side=tk.LEFT)

        columns = ('序号', '付款编号', '客户', '金额', '付款日期', '付款方式', '状态')
        tree = ttk.Treeview(list_window, columns=columns, show='headings')
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=130)
        tree.column('序号', width=50)
        tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        scrollbar = ttk.Scrollbar(list_window, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        refresh_list()

        total_amount = sum(p.get('amount', 0) for p in self.payments)
        total_label = tk.Label(list_window, text=f"付款总额: ¥{total_amount:.2f}", font=("Microsoft YaHei", 12, "bold"), fg='red')
        total_label.pack(pady=5)

        def view_detail():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            item = tree.item(selection[0])
            idx = item['values'][0] - 1
            payment = self.payments[idx]
            detail_window = tk.Toplevel(list_window)
            detail_window.title(f"付款详情 - {payment.get('payment_no', '')}")
            detail_window.geometry("600x400")
            info_frame = tk.Frame(detail_window, padx=20, pady=20)
            info_frame.pack(fill=tk.BOTH, expand=True)
            tk.Label(info_frame, text=f"付款编号: {payment.get('payment_no', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"客户名称: {payment.get('customer', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"付款金额: ¥{payment.get('amount', 0):.2f}", font=("Microsoft YaHei", 14, "bold"), fg='red').pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"付款日期: {payment.get('payment_date', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"付款方式: {payment.get('payment_method', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"关联账单: {payment.get('related_invoice', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"状态: {payment.get('status', '')}", font=("Microsoft YaHei", 12)).pack(anchor=tk.W, pady=5)
            tk.Label(info_frame, text=f"备注: {payment.get('remark', '')}", font=("Microsoft YaHei", 11)).pack(anchor=tk.W, pady=5)

        def delete_payment():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            if messagebox.askyesno("确认", "确定要删除这条付款记录吗?"):
                item = tree.item(selection[0])
                idx = item['values'][0] - 1
                del self.payments[idx]
                self.save_payments()
                refresh_list()
                new_total = sum(p.get('amount', 0) for p in self.payments)
                total_label.config(text=f"付款总额: ¥{new_total:.2f}")
                messagebox.showinfo("成功", "删除成功!")

        btn_frame = tk.Frame(list_window)
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="查看详情", command=view_detail, width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="删除记录", command=delete_payment, width=15, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=list_window.destroy, width=15).pack(side=tk.LEFT, padx=5)

    def show_payments_page(self):
        """显示付款列表 - 页面内显示"""
        self.status_label.config(text="正在加载付款列表...")
        self.clear_main_content()
        
        container = tk.Frame(self.main_content_frame, bg='#ffffff', padx=20, pady=15)
        container.pack(fill=tk.BOTH, expand=True)
        
        # 顶部标题区域
        header_frame = tk.Frame(container, bg='#ffffff')
        header_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Label(header_frame, text="付款记录管理", font=("Microsoft YaHei", 18, "bold"), 
                bg='#ffffff', fg='#2c3e50').pack(side=tk.LEFT)
        tk.Button(header_frame, text="← 返回主页", command=self.show_welcome_page,
                  font=("Microsoft YaHei", 10), bg='#95a5a6', fg='white',
                  relief=tk.FLAT, padx=15, pady=5).pack(side=tk.RIGHT)
        
        # 搜索区域 - 年份下拉框 + 客户下拉框
        search_frame = tk.Frame(container, bg='#f5f5f5', padx=10, pady=10)
        search_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 年份下拉框
        tk.Label(search_frame, text="年份:", font=("Microsoft YaHei", 11), bg='#f5f5f5').pack(side=tk.LEFT)
        year_var = tk.StringVar()
        year_combo = ttk.Combobox(search_frame, textvariable=year_var, width=10, state='readonly', font=("Microsoft YaHei", 11))
        current_year = datetime.now().year
        year_list = [str(current_year), '全部'] + [str(y) for y in range(current_year - 5, current_year)]
        year_combo['values'] = year_list
        year_combo.current(0)
        year_combo.pack(side=tk.LEFT, padx=(5, 15))
        
        # 表格 - 15列：序号、年份、客户、1月-12月
        columns = ['序号', '年份', '客户']
        for i in range(1, 13):
            columns.append(f'{i}月')
        tree = ttk.Treeview(container, columns=columns, show='headings', height=15)
        tree.heading('序号', text='序号')
        tree.heading('年份', text='年份')
        tree.heading('客户', text='客户')
        tree.column('序号', width=50, anchor='center')
        tree.column('年份', width=70, anchor='center')
        tree.column('客户', width=100, anchor='center')
        for i in range(1, 13):
            tree.heading(f'{i}月', text=f'{i}月')
            tree.column(f'{i}月', width=70, anchor='center')
        tree.pack(fill=tk.BOTH, expand=True, pady=10)
        scrollbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        
        def refresh_list():
            for item in tree.get_children():
                tree.delete(item)
            # 获取筛选条件
            selected_year = year_var.get()
            
            # 同步加载客户数据 - 用于预填客户名
            try:
                customers_path = os.path.join(self.data_dir, 'customers.json')
                if os.path.exists(customers_path):
                    with open(customers_path, 'r', encoding='utf-8') as f:
                        customers_data = json.load(f)
                else:
                    customers_data = []
            except:
                customers_data = []
            
            # 构建客户名列表（与客户资料列表同步，优先使用客户资料列表）
            if customers_data:
                customer_names = [c.get('name', '') for c in customers_data if c.get('name')]
            else:
                # 如果客户资料为空，则使用付款记录中的客户名
                customer_names = list(set([p.get('customer', '') for p in self.payments if p.get('customer')]))

            # 按年份+客户分组汇总每月付款金额
            data_by_year_customer = {}
            for payment in self.payments:
                # 从付款日期中提取年份和月份
                payment_date = payment.get('payment_date', '')
                customer = payment.get('customer', '')
                try:
                    if '-' in payment_date:
                        year = int(payment_date.split('-')[0])
                        month = int(payment_date.split('-')[1])
                    elif '/' in payment_date:
                        parts = payment_date.split('/')
                        year = int(parts[0])
                        month = int(parts[1])
                    else:
                        year = datetime.now().year
                        month = 1
                except:
                    year = datetime.now().year
                    month = 1
                
                # 年份筛选
                if selected_year and selected_year != '全部' and str(year) != selected_year:
                    continue
                
                key = (year, customer)
                if key not in data_by_year_customer:
                    data_by_year_customer[key] = {i: 0.0 for i in range(1, 13)}
                if 1 <= month <= 12:
                    data_by_year_customer[key][month] += payment.get('amount', 0)
            
            # 显示数据 - 优先显示客户资料列表中的客户，按年份降序、客户名升序排序
            idx = 1
            
            # 获取需要显示的年份列表
            if selected_year and selected_year != '全部':
                year_list = [int(selected_year)]
            else:
                year_list = sorted(set([int(selected_year)] if selected_year and selected_year != '全部' else list(data_by_year_customer.keys())) | 
                                   {datetime.now().year}, reverse=True)
            
            # 如果选中了特定年份，只显示该年份；如果选了"全部"，显示当前年份及以上的所有年份
            if selected_year == '全部':
                # 显示所有年份，包括客户资料列表中的客户
                all_years = sorted(set([key[0] for key in data_by_year_customer.keys()]) | {datetime.now().year}, reverse=True)
            elif selected_year and selected_year != '全部':
                all_years = [int(selected_year)]
            else:
                all_years = sorted(set([key[0] for key in data_by_year_customer.keys()]) | {datetime.now().year}, reverse=True)
            
            # 显示有付款记录的客户
            for year in all_years:
                for customer in sorted(customer_names):
                    key = (year, customer)
                    if key in data_by_year_customer:
                        values = [idx, str(year), customer]
                        for month in range(1, 13):
                            amount = data_by_year_customer[key][month]
                            values.append(f"¥{amount:.2f}" if amount > 0 else '')
                        tree.insert('', tk.END, values=values)
                        idx += 1
                    elif selected_year != '全部':
                        # 如果选中了特定年份，且客户在该年份没有付款记录，也显示一行空记录
                        # 但只在选择了"全部"年份时显示没有付款记录的客户
                        pass
                
                # 对于选中了"全部"的年份，也要显示没有付款记录的客户（金额全为空）
                if selected_year == '全部':
                    # 获取该年份已有的客户
                    existing_customers = set([key[1] for key in data_by_year_customer.keys() if key[0] == year])
                    # 显示没有付款记录的客户
                    for customer in sorted(customer_names):
                        if customer not in existing_customers:
                            values = [idx, str(year), customer]
                            for month in range(1, 13):
                                values.append('')
                            tree.insert('', tk.END, values=values)
                            idx += 1
        
        # 绑定下拉框选择事件
        year_combo.bind('<<ComboboxSelected>>', lambda e: refresh_list())
        refresh_list()
        
        # 合计金额
        total_amount = sum(p.get('amount', 0) for p in self.payments)
        total_label = tk.Label(container, text=f"付款总额: ¥{total_amount:.2f}", 
                              font=("Microsoft YaHei", 14, "bold"), fg='#e74c3c', bg='#ffffff')
        total_label.pack(pady=10)
        
        # 按钮区域
        btn_frame = tk.Frame(container, bg='#ffffff')
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="导出记录", command=self.export_payments, width=15, height=2,
                 bg='#3498db', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)
        
        self.status_label.config(text="付款列表已加载")

    def _delete_payment_in_page(self, tree, refresh_callback, total_label):
        """删除付款记录（页面内版本）"""
        selection = tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一条记录")
            return
        if messagebox.askyesno("确认", "确定要删除这条付款记录吗?"):
            item = tree.item(selection[0])
            idx = item['values'][0] - 1
            del self.payments[idx]
            self.save_payments()
            refresh_callback()
            new_total = sum(p.get('amount', 0) for p in self.payments)
            total_label.config(text=f"付款总额: ¥{new_total:.2f}")
            messagebox.showinfo("成功", "删除成功!")

    def add_payment(self):
        """添加付款记录"""
        self.status_label.config(text="正在添加付款记录...")
        payment_window = tk.Toplevel(self.root)
        payment_window.title("添加付款记录")
        payment_window.geometry("500x550")
        form_frame = tk.Frame(payment_window, padx=30, pady=20)
        form_frame.pack(fill=tk.BOTH, expand=True)
        tk.Label(form_frame, text="添加付款记录", font=("Microsoft YaHei", 14, "bold")).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 15))

        tk.Label(form_frame, text="付款编号:").grid(row=1, column=0, sticky=tk.W, pady=8)
        payment_no_entry = tk.Entry(form_frame, width=30)
        payment_no_entry.insert(0, f"FK{datetime.now().strftime('%Y%m%d%H%M%S')}")
        payment_no_entry.grid(row=1, column=1, pady=8)

        tk.Label(form_frame, text="客户名称:").grid(row=2, column=0, sticky=tk.W, pady=8)
        customer_entry = tk.Entry(form_frame, width=30)
        customer_entry.grid(row=2, column=1, pady=8)

        tk.Label(form_frame, text="付款金额:").grid(row=3, column=0, sticky=tk.W, pady=8)
        amount_entry = tk.Entry(form_frame, width=30)
        amount_entry.grid(row=3, column=1, pady=8)

        tk.Label(form_frame, text="付款日期:").grid(row=4, column=0, sticky=tk.W, pady=8)
        date_entry = tk.Entry(form_frame, width=30)
        date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        date_entry.grid(row=4, column=1, pady=8)

        tk.Label(form_frame, text="付款方式:").grid(row=5, column=0, sticky=tk.W, pady=8)
        payment_method_var = tk.StringVar(value="银行转账")
        payment_method_combo = ttk.Combobox(form_frame, width=28, textvariable=payment_method_var)
        payment_method_combo['values'] = ('银行转账', '现金', '支票', '支付宝', '微信', '其他')
        payment_method_combo.grid(row=5, column=1, pady=8)

        tk.Label(form_frame, text="关联账单:").grid(row=6, column=0, sticky=tk.W, pady=8)
        related_invoice_entry = tk.Entry(form_frame, width=30)
        related_invoice_entry.grid(row=6, column=1, pady=8)

        tk.Label(form_frame, text="付款状态:").grid(row=7, column=0, sticky=tk.W, pady=8)
        status_var = tk.StringVar(value="已付")
        status_combo = ttk.Combobox(form_frame, width=28, textvariable=status_var)
        status_combo['values'] = ('未付', '已付', '部分付款', '已过期')
        status_combo.grid(row=7, column=1, pady=8)

        tk.Label(form_frame, text="备注:").grid(row=8, column=0, sticky=tk.W, pady=8)
        remark_entry = tk.Entry(form_frame, width=30)
        remark_entry.grid(row=8, column=1, pady=8)

        btn_frame = tk.Frame(form_frame)
        btn_frame.grid(row=9, column=0, columnspan=2, pady=20)
        def save_payment():
            if not customer_entry.get():
                messagebox.showwarning("警告", "请输入客户名称")
                return
            try:
                amount = float(amount_entry.get())
            except ValueError:
                messagebox.showwarning("警告", "请输入正确的金额!")
                return
            payment_data = {
                'payment_no': payment_no_entry.get(),
                'customer': customer_entry.get(),
                'amount': amount,
                'payment_date': date_entry.get(),
                'payment_method': payment_method_var.get(),
                'related_invoice': related_invoice_entry.get(),
                'status': status_var.get(),
                'remark': remark_entry.get(),
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            self.payments.append(payment_data)
            self.save_payments()
            messagebox.showinfo("成功", "付款记录保存成功!")
            payment_window.destroy()
            self.status_label.config(text="付款记录已保存")
        tk.Button(btn_frame, text="保存", command=save_payment, width=15, height=2, bg='#27ae60', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=payment_window.destroy, width=15, height=2, bg='#e74c3c', fg='white', font=("Microsoft YaHei", 11)).pack(side=tk.LEFT, padx=5)

    def export_payments(self):
        """导出付款记录"""
        if not self.payments:
            messagebox.showwarning("警告", "没有可导出的付款记录!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")], title="导出付款记录")
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("=" * 60 + "\n")
                    f.write("付款记录汇总\n")
                    f.write(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("=" * 60 + "\n\n")
                    for payment in self.payments:
                        f.write(f"编号: {payment.get('payment_no')}\n")
                        f.write(f"客户: {payment.get('customer')}\n")
                        f.write(f"金额: ¥{payment.get('amount', 0):.2f}\n")
                        f.write(f"日期: {payment.get('payment_date')}\n")
                        f.write(f"方式: {payment.get('payment_method')}\n")
                        f.write(f"状态: {payment.get('status')}\n")
                        f.write(f"备注: {payment.get('remark', '-')}\n")
                        f.write("-" * 40 + "\n")
                    total = sum(p.get('amount', 0) for p in self.payments)
                    f.write(f"\n总金额: ¥{total:.2f}\n")
                    f.write(f"付款记录: {len(self.payments)}条\n")
                messagebox.showinfo("成功", f"导出成功!\n文件保存至: {file_path}")
                self.status_label.config(text="付款记录已导出")
            except Exception as e:
                messagebox.showerror("错误", f"导出失败: {str(e)}")

    # -------------------- 客户管理 --------------------
    def show_customers(self):
        """显示客户列表"""
        self.status_label.config(text="正在加载客户列表...")
        list_window = tk.Toplevel(self.root)
        list_window.title("客户列表")
        list_window.geometry("600x400")
        list_window.update_idletasks()
        screen_width = list_window.winfo_screenwidth()
        screen_height = list_window.winfo_screenheight()
        x = (screen_width - 600) // 2
        y = (screen_height - 400) // 2
        list_window.geometry(f"600x400+{x}+{y}")

        tk.Label(list_window, text="客户列表", font=("Microsoft YaHei", 16, "bold")).pack(pady=10)

        search_frame = tk.Frame(list_window)
        search_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(search_frame, text="搜索:").pack(side=tk.LEFT)
        search_entry = tk.Entry(search_frame, width=20)
        search_entry.pack(side=tk.LEFT, padx=5)
        def refresh_list():
            for item in tree.get_children():
                tree.delete(item)
            search_text = search_entry.get().lower()
            for idx, customer in enumerate(self.customers, 1):
                if not search_text or search_text in customer.get('name', '').lower() or search_text in customer.get('contact', '').lower():
                    tree.insert('', tk.END, values=(idx, customer.get('name', ''), customer.get('contact', '')))
        tk.Button(search_frame, text="刷新列表", command=refresh_list).pack(side=tk.LEFT)

        columns = ('序号', '客户名称', '联系人')
        tree = ttk.Treeview(list_window, columns=columns, show='headings')
        tree.heading('序号', text='序号')
        tree.heading('客户名称', text='客户名称')
        tree.heading('联系人', text='联系人')
        tree.column('序号', width=80, anchor='center')
        tree.column('客户名称', width=250, anchor='center')
        tree.column('联系人', width=200, anchor='center')
        tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        scrollbar = ttk.Scrollbar(list_window, orient=tk.VERTICAL, command=tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar.set)
        refresh_list()

        def add_new_customer():
            add_window = tk.Toplevel(list_window)
            add_window.title("添加客户")
            add_window.geometry("400x250")
            form_frame = tk.Frame(add_window, padx=20, pady=20)
            form_frame.pack(fill=tk.BOTH, expand=True)
            tk.Label(form_frame, text="客户名称:").grid(row=0, column=0, sticky=tk.W, pady=10)
            name_entry = tk.Entry(form_frame, width=25)
            name_entry.grid(row=0, column=1, pady=10)
            tk.Label(form_frame, text="联系方式:").grid(row=1, column=0, sticky=tk.W, pady=10)
            contact_entry = tk.Entry(form_frame, width=25)
            contact_entry.grid(row=1, column=1, pady=10)
            def save_new():
                if not name_entry.get():
                    messagebox.showwarning("警告", "请输入客户名称")
                    return
                self.customers.append({'name': name_entry.get(), 'contact': contact_entry.get(), 'phone': '', 'address': '', 'remark': ''})
                self.save_customers()
                refresh_list()
                add_window.destroy()
            btn_frame2 = tk.Frame(form_frame)
            btn_frame2.grid(row=2, column=0, columnspan=2, pady=20)
            tk.Button(btn_frame2, text="保存", command=save_new, width=12, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame2, text="取消", command=add_window.destroy, width=12, bg='#95a5a6').pack(side=tk.LEFT, padx=5)

        def edit_customer():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            item = tree.item(selection[0])
            idx = item['values'][0] - 1
            edit_window = tk.Toplevel(list_window)
            edit_window.title("编辑客户")
            edit_window.geometry("400x200")
            customer = self.customers[idx]
            form_frame = tk.Frame(edit_window, padx=20, pady=20)
            form_frame.pack(fill=tk.BOTH, expand=True)
            tk.Label(form_frame, text="客户名称:").grid(row=0, column=0, sticky=tk.W, pady=10)
            name_entry = tk.Entry(form_frame, width=25)
            name_entry.insert(0, customer.get('name', ''))
            name_entry.grid(row=0, column=1, pady=10)
            tk.Label(form_frame, text="联系方式:").grid(row=1, column=0, sticky=tk.W, pady=10)
            contact_entry = tk.Entry(form_frame, width=25)
            contact_entry.insert(0, customer.get('contact', ''))
            contact_entry.grid(row=1, column=1, pady=10)
            def save_edit():
                self.customers[idx] = {'name': name_entry.get(), 'contact': contact_entry.get(), 'phone': '', 'address': '', 'remark': ''}
                self.save_customers()
                refresh_list()
                edit_window.destroy()
            btn_frame2 = tk.Frame(form_frame)
            btn_frame2.grid(row=2, column=0, columnspan=2, pady=20)
            tk.Button(btn_frame2, text="保存", command=save_edit, width=12, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame2, text="取消", command=edit_window.destroy, width=12, bg='#95a5a6').pack(side=tk.LEFT, padx=5)

        def delete_customer():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("警告", "请先选择一条记录")
                return
            if messagebox.askyesno("确认", "确定要删除这个客户吗?"):
                item = tree.item(selection[0])
                idx = item['values'][0] - 1
                del self.customers[idx]
                self.save_customers()
                refresh_list()
                messagebox.showinfo("成功", "删除成功!")

        btn_frame = tk.Frame(list_window)
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="添加客户", command=add_new_customer, width=15, bg='#3498db', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="编辑客户", command=edit_customer, width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="删除客户", command=delete_customer, width=15, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=list_window.destroy, width=15).pack(side=tk.LEFT, padx=5)

    # -------------------- 辅助功能 --------------------
    def update_stats(self):
        """更新统计信息"""
        if hasattr(self, 'stats_label'):
            self.stats_label.config(text=f"对账单数: {len(self.monthly_invoices)}\n报价单数: {len(self.quotations)}")

    def _analyze_delivery_photo_with_ai(self, photo_path):
        """使用AI模型分析送货单照片，提取信息
        
        识别内容：
        - 客户名称
        - 送货日期
        - 品名
        - 颜色
        - 规格
        - 数量
        - 单价
        - 备注中的染费
        """
        try:
            if not AI_AVAILABLE:
                raise ImportError("AI库未安装")
            
            # 检查是否有API配置
            model_name = getattr(self, 'doubao_model', '豆包2.0LITE')
            base_url = getattr(self, 'doubao_base_url', 'https://api.doubao.com')
            api_key = getattr(self, 'doubao_api_key', '')
            
            if not api_key:
                raise ValueError("请先在系统设置中配置API Key")
            
            # 读取图片并转换为base64
            with open(photo_path, 'rb') as f:
                image_data = f.read()
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # 构建请求
            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {api_key}'
            }
            
            # 根据模型选择不同的API端点
            if model_name == '豆包2.0LITE':
                model_endpoint = 'doubao-2.0-lite'
            elif model_name == '豆包2.0PRO':
                model_endpoint = 'doubao-2.0-pro'
            else:
                model_endpoint = 'doubao-2.0-lite'
            
            # 构建提示词
            prompt = """请分析这张送货单图片，提取以下信息并以JSON格式返回：
            1. 客户名称
            2. 送货日期
            3. 商品明细列表（每个商品包含：品名、规格、颜色、单位、数量、单价、金额）
            4. 染费（如果有）
            
            请确保返回格式为：
            {
                "客户名称": "客户名称",
                "送货日期": "送货日期",
                "items": [
                    {
                        "品名": "品名",
                        "规格": "规格",
                        "颜色": "颜色",
                        "单位": "单位",
                        "数量": "数量",
                        "单价": "单价",
                        "金额": "金额"
                    }
                ],
                "染费": "染费金额"
            }
            
            如果没有找到某项信息，请留空字符串。"""
            
            payload = {
                'model': model_endpoint,
                'messages': [
                    {
                        'role': 'user',
                        'content': [
                            {'type': 'text', 'text': prompt},
                            {'type': 'image_url', 'image_url': {'url': f'data:image/jpeg;base64,{image_base64}'}}
                        ]
                    }
                ],
                'max_tokens': 2000,
                'temperature': 0.1
            }
            
            # 发送请求
            response = requests.post(
                f'{base_url}/v1/chat/completions',
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code != 200:
                raise Exception(f"API请求失败: {response.status_code} - {response.text}")
            
            # 解析响应
            response_data = response.json()
            content = response_data['choices'][0]['message']['content']
            
            # 提取JSON部分
            import re
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                result = json.loads(json_str)
            else:
                # 如果没有找到JSON，尝试解析文本
                result = {'客户名称': '', 'items': [], '染费': ''}
                # 简单解析逻辑
                lines = content.split('\n')
                for line in lines:
                    if '客户' in line:
                        result['客户名称'] = line.replace('客户', '').replace(':', '').strip()
                    elif '染费' in line:
                        result['染费'] = line.replace('染费', '').replace(':', '').strip()
            
            return result
            
        except Exception as e:
            print(f"AI分析错误: {str(e)}")
            return None

    def _analyze_delivery_photo(self, photo_path):
        """分析送货单照片，提取信息（根据系统设置选择OCR或AI）
        
        识别内容：
        - 客户名称
        - 送货日期
        - 品名
        - 颜色
        - 规格
        - 数量
        - 单价
        - 备注中的染费
        """
        # 检查系统设置，决定使用OCR还是AI
        use_ai = False
        try:
            model_name = getattr(self, 'doubao_model', '豆包2.0LITE')
            api_key = getattr(self, 'doubao_api_key', '')
            
            # 如果有配置AI模型且API Key不为空，则使用AI
            if model_name and api_key and AI_AVAILABLE:
                use_ai = True
        except:
            use_ai = False
        
        if use_ai:
            # 使用AI模型分析
            self.status_label.config(text="正在使用AI模型分析图片...")
            result = self._analyze_delivery_photo_with_ai(photo_path)
            if result:
                self.status_label.config(text="AI分析完成")
                return result
            else:
                # AI分析失败，回退到OCR
                self.status_label.config(text="AI分析失败，尝试使用OCR...")
        
        # 使用OCR分析
        try:
            if not OCR_AVAILABLE:
                raise ImportError("OCR库未安装")
            img = Image.open(photo_path)
            
            # 使用中文+英文混合识别，提高准确率
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            result = {'客户名称': '', 'items': [], '染费': ''}
            lines = text.split('\n')

            # 提取客户名称 - 智能识别
            for i, line in enumerate(lines[:10]):
                line = line.strip()
                if not line:
                    continue
                # 匹配带标签的客户名称行
                if any(kw in line for kw in ['客户', '公司', '单位', '购货单位', '收货单位']):
                    match = re.search(r'[:：\s]+([^\n:：]+)', line)
                    if match:
                        result['客户名称'] = match.group(1).strip()
                        # 清理可能残留的序号等前缀
                        result['客户名称'] = re.sub(r'^[0-9一二三四五六七八九十]+[.、)）\s]+', '', result['客户名称'])
                        break
                # 备用：如果第一行看起来像客户名称
                elif i == 0 and len(line) > 2 and len(line) < 400 and not any(c.isdigit() for c in line):
                    if not any(kw in line for kw in ['送货单', '对账单', '发票', '日期', '订单', 'BABES']):
                        result['客户名称'] = line.strip()
                        break

            # 提取染费（从备注区域）- 增强版
            # 查找备注区域（通常在表格下方或右侧）
            remark_lines = []
            in_remark_section = False
            for line in lines:
                line_stripped = line.strip()
                # 检测备注开始
                if any(kw in line_stripped for kw in ['备注', '注', '说明', 'Memo', 'Note']):
                    in_remark_section = True
                    remark_lines.append(line_stripped)
                elif in_remark_section:
                    remark_lines.append(line_stripped)
            
            # 搜索染费关键词
            for line in lines + remark_lines:
                line = line.strip()
                # 查找染费相关信息
                if any(kw in line for kw in ['染费', '染色费', '加工费', '染整费']):
                    # 提取染费后面的数字
                    # 匹配模式：染费xxx数字 或 染费：xxx数字 等
                    fee_patterns = [
                        r'染费[：:\s]*[¥￥$]?\s*(\d+\.?\d*)',  # 染费：123 或 染费 123
                        r'染费[:：]?\s*(\d+\.?\d*)',  # 染费123
                        r'[染费染色费加工费染整费][：:\s]*[¥￥$]?\s*(\d+\.?\d*)',  # 其他染费关键词
                        r'(\d+\.?\d*)\s*元',  # 数字+元（可能是染费金额）
                    ]
                    for pattern in fee_patterns:
                        fee_match = re.search(pattern, line)
                        if fee_match:
                            result['染费'] = fee_match.group(1)
                            break
                    if result['染费']:
                        break

            # 提取商品明细 - 增强版
            for line in lines:
                line = line.strip()
                if not line or len(line) < 3:
                    continue

                # 跳过标题行和备注行
                skip_keywords = ['客户', '送货单', '对账单', '公司名称', '地址', '电话', '银行', '账号', 
                               '注:', '注意', '以上', '一经', '本公司', 'COLOR', 'SIZE', '品名', '规格',
                               '一、', '二、', '三、', '四、', '五、', '1.', '2.', '3.', '4.', '5.',
                               '编号', '序号', '送货日期', '送货单号']
                if any(line.startswith(kw) or line.endswith(kw) for kw in skip_keywords if len(kw) > 2):
                    continue

                item = {}

                # 送货日期 - 多种格式支持
                date_patterns = [
                    r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?)',
                    r'(\d{8})',  # 20240101格式
                    r'(\d{4}\.\d{1,2}\.\d{1,2})'
                ]
                for pattern in date_patterns:
                    date_match = re.search(pattern, line)
                    if date_match:
                        item['送货日期'] = date_match.group(1).replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-')
                        break

                # 品名 - 增强匹配
                product_match = re.search(r'(?:品名|产品|名称|Name)[:：]?\s*([\u4e00-\u9fa5A-Za-z0-9\-]+)', line, re.IGNORECASE)
                if not product_match:
                    # 尝试匹配常见的纺织品名称
                    product_keywords = ['纱', '线', '布', '面料', '布料', '纤维', '棉', '丝', '毛', '涤', '锦', '氨', '纶', '纱线', '坯布', '面料', '牛仔布', '针织布', '梭织布', '毛呢', '羊绒', '羊毛', '真丝', '仿真丝']
                    for kw in product_keywords:
                        if kw in line:
                            # 提取包含产品名的部分
                            product_match = re.search(rf'([\u4e00-\u9fa5A-Za-z0-9\-]*{kw}[\u4e00-\u9fa5A-Za-z0-9\-]*)', line)
                            if product_match:
                                item['品名'] = product_match.group(1)
                                break
                    if not product_match:
                        # 尝试直接匹配产品描述（如 32S棉纱、40D氨纶等）
                        product_match = re.search(r'(\d+[A-Za-z][\u4e00-\u9fa5A-Za-z]+)', line)
                
                if product_match:
                    item['品名'] = product_match.group(1)

                # 颜色 - 独立提取
                color_value = ''
                color_keywords = ['红', '蓝', '绿', '白', '黑', '黄', '紫', '粉', '灰', '棕', '橙', '驼', '藏青', '牛仔', '彩', '米', '杏', '卡其', '军绿', '酒红', '宝蓝', '天蓝', '深蓝', '浅蓝', '深灰', '浅灰', '墨绿', '草绿', '翠绿', '咖啡', '褐色']
                for kw in color_keywords:
                    if kw in line:
                        # 提取颜色描述
                        color_match = re.search(rf'([\u4e00-\u9fa5]{1,4}(?:色|兰)?)', line)
                        if color_match:
                            color_value = color_match.group(1)
                        else:
                            color_value = kw + '色'
                        break

                # 规格 - 独立提取
                spec_value = ''
                spec_patterns = [
                    r'(?:规格|型号|Spec|SIZE|MODEL)[:：]?\s*([A-Za-z0-9\-]+)',
                    r'\b(\d+[\*×]\d+)\b',  # 如 100*200
                    r'\b(\d+[A-Za-z])\b',   # 如 32S, 40D
                    r'(S|M|L|XL|XXL|XXXL|xs)',
                    r'\b(\d+\.?\d*)D\b',  # 如 40D
                    r'\b(\d+)S\b',  # 如 32S
                    r'(\d+[A-Za-z]{1,3})',  # 通用规格格式
                ]
                for pattern in spec_patterns:
                    spec_match = re.search(pattern, line, re.IGNORECASE)
                    if spec_match:
                        spec_value = spec_match.group(1)
                        break

                # 颜色和规格分别存储
                if color_value:
                    item['颜色'] = color_value
                if spec_value:
                    item['规格'] = spec_value
                if color_value or spec_value:
                    item['颜色规格'] = f"{color_value} {spec_value}".strip() if color_value else spec_value

                # 数量 - 增强版
                qty_patterns = [
                    r'(?:数量|Qty|Q)[:：]?\s*(\d+\.?\d*)',
                    r'(?:x|×|X)\s*(\d+\.?\d*)',
                    r'(\d+\.?\d*)\s*(?:个|件|套|米|码|箱|PCS|pc|KG|kg|码|YDS|码)',
                    r'^(\d+\.?\d*)\s',
                    r'[-–]\s*(\d+\.?\d*)\s*(?:个|件|套|米|码|箱|PCS|pc|KG|kg|码|YDS|码)',  # -100码
                ]
                for pattern in qty_patterns:
                    qty_match = re.search(pattern, line, re.IGNORECASE)
                    if qty_match:
                        item['数量'] = qty_match.group(1)
                        break

                # 单价 - 增强版
                price_patterns = [
                    r'(?:单价|Price)[:：]?\s*[¥￥$]?\s*(\d+\.?\d*)',
                    r'[¥￥$]\s*(\d+\.?\d*)',  # 直接跟货币符号
                    r'(\d+\.?\d*)\s*/\s*(?:码|米|个|件|套|PCS|kg|KG)',  # 数字/码
                ]
                for pattern in price_patterns:
                    price_match = re.search(pattern, line, re.IGNORECASE)
                    if price_match:
                        item['单价'] = price_match.group(1)
                        break

                # 金额
                amount_match = re.search(r'(?:金额|合计|Total|Sum)[:：]?\s*[¥￥$]?\s*(\d+\.?\d*)', line, re.IGNORECASE)
                if amount_match:
                    item['金额'] = amount_match.group(1)

                # 单位
                unit_match = re.search(r'(?:单位|Unit)[:：]?\s*([\u4e00-\u9fa5A-Za-z]+)', line)
                if unit_match:
                    item['单位'] = unit_match.group(1)
                elif re.search(r'\d+\.?\d*\s*(?:个|件|套|米|码|箱|PCS|pc|KG|kg|YDS)', line):
                    unit_map = {'个': '个', '件': '件', '套': '套', '米': '米', '码': '码', '箱': '箱', 'PCS': 'PCS', 'pc': '个', 'KG': 'KG', 'kg': 'KG', 'YDS': 'YDS'}
                    for u, name in unit_map.items():
                        if u in line:
                            item['单位'] = name
                            break

                # 如果解析到了有效数据项，则保存
                valid_fields = [k for k, v in item.items() if v]
                if len(valid_fields) >= 2:  # 至少有两个有效字段
                    for field in ['送货日期', '品名', '颜色', '规格', '颜色规格', '单位', '数量', '单价', '金额']:
                        if field not in item:
                            item[field] = ''
                    result['items'].append(item)

            return result
        except Exception as e:
            print(f"OCR分析错误: {str(e)}")
            return None

    # -------------------- 系统设置 --------------------
    def show_system_settings(self):
        """显示系统设置页面"""
        self.status_label.config(text="正在打开系统设置...")
        self.clear_main_content()

        container = tk.Frame(self.main_content_frame, bg='#ffffff', padx=30, pady=20)
        container.pack(fill=tk.BOTH, expand=True)

        header_frame = tk.Frame(container, bg='#ffffff')
        header_frame.pack(fill=tk.X, pady=(0, 15))
        tk.Button(header_frame, text="← 返回主页", command=self.show_welcome_page,
                  font=("Microsoft YaHei", 10), bg='#95a5a6', fg='white',
                  relief=tk.FLAT, padx=15, pady=5).pack(side=tk.LEFT)
        tk.Label(header_frame, text="系统设置", font=("Microsoft YaHei", 18, "bold"), bg='#ffffff', fg='#2c3e50').pack(side=tk.LEFT, padx=20)

        content_frame = tk.Frame(container, bg='#ffffff')
        content_frame.pack(fill=tk.BOTH, expand=True)

        left_frame = tk.Frame(content_frame, bg='#ffffff')
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))

        # 基本设置
        basic_frame = tk.LabelFrame(left_frame, text="系统基本设置", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        basic_frame.pack(fill=tk.X, pady=(0, 15))

        row0_1 = tk.Frame(basic_frame, bg='#ffffff')
        row0_1.pack(fill=tk.X, pady=5)
        tk.Label(row0_1, text="系统名称:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        system_name_entry = tk.Entry(row0_1, width=30)
        system_name_entry.pack(side=tk.LEFT, padx=10)
        system_name_entry.insert(0, getattr(self, 'system_name', ''))

        row0_2 = tk.Frame(basic_frame, bg='#ffffff')
        row0_2.pack(fill=tk.X, pady=5)
        tk.Label(row0_2, text="公司名称:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        company_name_entry = tk.Entry(row0_2, width=30)
        company_name_entry.pack(side=tk.LEFT, padx=10)
        company_name_entry.insert(0, self.company_name)

        row0_3 = tk.Frame(basic_frame, bg='#ffffff')
        row0_3.pack(fill=tk.X, pady=5)
        tk.Label(row0_3, text="LOGO图片:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        logo_path_var = tk.StringVar(value=getattr(self, 'logo_path', ''))
        logo_entry = tk.Entry(row0_3, width=25, textvariable=logo_path_var)
        logo_entry.pack(side=tk.LEFT, padx=10)
        def select_logo():
            file_path = filedialog.askopenfilename(title="选择LOGO图片", filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif"), ("所有文件", "*.*")])
            if file_path:
                logo_path_var.set(file_path)
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(file_path)
                    img_preview = img.resize((120, 120), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img_preview)
                    logo_preview_label.config(image=photo, text="")
                    logo_preview_label.image = photo
                except:
                    pass
        tk.Button(row0_3, text="选择图片", command=select_logo, bg='#3498db', fg='white', relief='flat', padx=10).pack(side=tk.LEFT)

        # 公章图片选择行
        stamp_row = tk.Frame(basic_frame, bg='#ffffff')
        stamp_row.pack(fill=tk.X, pady=5)
        tk.Label(stamp_row, text="公章图片:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        stamp_path_var = tk.StringVar(value=getattr(self, 'stamp_path', ''))
        stamp_entry = tk.Entry(stamp_row, width=25, textvariable=stamp_path_var)
        stamp_entry.pack(side=tk.LEFT, padx=10)
        def select_stamp():
            file_path = filedialog.askopenfilename(title="选择公司公章图片", filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif"), ("所有文件", "*.*")])
            if file_path:
                stamp_path_var.set(file_path)
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(file_path)
                    img_preview = img.resize((120, 120), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img_preview)
                    stamp_preview_label.config(image=photo, text="")
                    stamp_preview_label.image = photo
                except Exception as e:
                    print(f"预览公章失败: {e}")
        tk.Button(stamp_row, text="选择图片", command=select_stamp, bg='#3498db', fg='white', relief='flat', padx=10).pack(side=tk.LEFT)

        # 预览框对齐行 - LOGO预览和公章预览并排对齐
        preview_row = tk.Frame(basic_frame, bg='#ffffff')
        preview_row.pack(fill=tk.X, pady=(10, 5))
        
        # LOGO预览区域
        tk.Label(preview_row, text="LOGO预览:", width=12, bg='#ffffff').pack(side=tk.LEFT, anchor='n')
        preview_frame = tk.Frame(preview_row, bg='#f0f0f0', width=120, height=120, relief='solid', bd=1)
        preview_frame.pack(side=tk.LEFT, padx=(10, 30))
        preview_frame.pack_propagate(False)
        logo_preview_label = tk.Label(preview_frame, bg='#f0f0f0', text="选择LOGO\n图片预览", fg='#999999')
        logo_preview_label.pack(fill=tk.BOTH, expand=True)
        if getattr(self, 'logo_path', ''):
            try:
                from PIL import Image, ImageTk
                img = Image.open(self.logo_path)
                img_preview = img.resize((120, 120), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img_preview)
                logo_preview_label.config(image=photo, text="")
                logo_preview_label.image = photo
            except:
                pass

        # 公章预览区域 - 与LOGO预览对齐
        tk.Label(preview_row, text="公章预览:", width=12, bg='#ffffff').pack(side=tk.LEFT, anchor='n')
        stamp_preview_frame = tk.Frame(preview_row, bg='#f0f0f0', relief='solid', bd=1, width=120, height=120)
        stamp_preview_frame.pack(side=tk.LEFT, padx=10)
        stamp_preview_frame.pack_propagate(False)
        stamp_preview_label = tk.Label(stamp_preview_frame, bg='#f0f0f0', text="选择公章\n图片预览", fg='#999999', cursor='hand2')
        stamp_preview_label.pack(fill=tk.BOTH, expand=True)
        
        # 加载已有公章预览
        if getattr(self, 'stamp_path', ''):
            try:
                from PIL import Image, ImageTk
                img = Image.open(self.stamp_path)
                img_preview = img.resize((120, 120), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img_preview)
                stamp_preview_label.config(image=photo, text="")
                stamp_preview_label.image = photo
            except Exception as e:
                print(f"加载公章预览失败: {e}")

        # 右侧区域 - 公司信息、打印设置、银行信息
        right_frame = tk.Frame(content_frame, bg='#ffffff')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        # 公司信息
        company_frame = tk.LabelFrame(right_frame, text="公司信息", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        company_frame.pack(fill=tk.X, pady=(0, 15))
        row2 = tk.Frame(company_frame, bg='#ffffff')
        row2.pack(fill=tk.X, pady=5)
        tk.Label(row2, text="公司地址:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        company_addr_entry = tk.Entry(row2, width=30)
        company_addr_entry.pack(side=tk.LEFT, padx=10)
        company_addr_entry.insert(0, getattr(self, 'company_address', ''))
        row3 = tk.Frame(company_frame, bg='#ffffff')
        row3.pack(fill=tk.X, pady=5)
        tk.Label(row3, text="联系电话:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        company_phone_entry = tk.Entry(row3, width=30)
        company_phone_entry.pack(side=tk.LEFT, padx=10)
        company_phone_entry.insert(0, getattr(self, 'company_phone', ''))

        # 打印设置
        print_frame = tk.LabelFrame(right_frame, text="打印设置", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        print_frame.pack(fill=tk.X, pady=(0, 15))
        row6 = tk.Frame(print_frame, bg='#ffffff')
        row6.pack(fill=tk.X, pady=5)
        tk.Label(row6, text="纸张大小:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        paper_size_var = tk.StringVar(value=getattr(self, 'paper_size', 'A4'))
        tk.Radiobutton(row6, text="A4", variable=paper_size_var, value='A4', bg='#ffffff').pack(side=tk.LEFT, padx=10)
        tk.Radiobutton(row6, text="A5", variable=paper_size_var, value='A5', bg='#ffffff').pack(side=tk.LEFT)

        # 银行信息
        bank_frame = tk.LabelFrame(right_frame, text="银行信息", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        bank_frame.pack(fill=tk.X)
        row4 = tk.Frame(bank_frame, bg='#ffffff')
        row4.pack(fill=tk.X, pady=5)
        tk.Label(row4, text="银行账号:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        bank_account_entry = tk.Entry(row4, width=30)
        bank_account_entry.pack(side=tk.LEFT, padx=10)
        bank_account_entry.insert(0, getattr(self, 'bank_account', ''))
        row5 = tk.Frame(bank_frame, bg='#ffffff')
        row5.pack(fill=tk.X, pady=5)
        tk.Label(row5, text="开户银行:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        bank_name_entry = tk.Entry(row5, width=30)
        bank_name_entry.pack(side=tk.LEFT, padx=10)
        bank_name_entry.insert(0, getattr(self, 'bank_name', ''))

        # 报价人和经办人
        person_frame = tk.LabelFrame(right_frame, text="人员信息", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        person_frame.pack(fill=tk.X, pady=(15, 0))
        row7 = tk.Frame(person_frame, bg='#ffffff')
        row7.pack(fill=tk.X, pady=5)
        tk.Label(row7, text="报价人:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        quoter_entry = tk.Entry(row7, width=30)
        quoter_entry.pack(side=tk.LEFT, padx=10)
        quoter_entry.insert(0, getattr(self, 'quoter', ''))
        row8 = tk.Frame(person_frame, bg='#ffffff')
        row8.pack(fill=tk.X, pady=5)
        tk.Label(row8, text="经办人:", width=12, bg='#ffffff').pack(side=tk.LEFT)
        handler_entry = tk.Entry(row8, width=30)
        handler_entry.pack(side=tk.LEFT, padx=10)
        handler_entry.insert(0, getattr(self, 'handler', ''))

        # AI模型设置
        ai_model_frame = tk.LabelFrame(left_frame, text="AI模型设置", font=("Microsoft YaHei", 12), padx=15, pady=15, bg='#ffffff')
        ai_model_frame.pack(fill=tk.X, pady=(15, 0))
        
        # 预设配置
        default_base_url = 'https://ark.cn-beijing.volces.com/api/v3'
        default_pro_model = 'doubao-seed-2-0-pro-260215'
        default_lite_model = 'doubao-seed-2-0-lite-260215'
        
        # 模型选择（单选按钮）
        model_select_row = tk.Frame(ai_model_frame, bg='#ffffff')
        model_select_row.pack(fill=tk.X, pady=(0, 10))
        tk.Label(model_select_row, text="选择模型:", width=12, font=("Microsoft YaHei", 11, "bold"), bg='#ffffff', fg='#2c3e50').pack(side=tk.LEFT)
        
        selected_model_var = tk.StringVar(value='pro')  # 默认选择PRO模型
        
        def on_model_changed(*args):
            """当模型选择改变时，自动填入对应的模型名称"""
            if selected_model_var.get() == 'pro':
                model_name_entry.delete(0, tk.END)
                model_name_entry.insert(0, default_pro_model)
            else:  # lite
                model_name_entry.delete(0, tk.END)
                model_name_entry.insert(0, default_lite_model)
        
        # PRO模型单选按钮
        pro_radio = tk.Radiobutton(model_select_row, text="豆包2.0 PRO ✨（推荐）", 
                                   variable=selected_model_var, value='pro',
                                   bg='#ffffff', activebackground='#ffffff',
                                   font=("Microsoft YaHei", 10), fg='#27ae60',
                                   command=on_model_changed)
        pro_radio.pack(side=tk.LEFT, padx=5)
        
        # LITE模型单选按钮
        lite_radio = tk.Radiobutton(model_select_row, text="豆包2.0 LITE 💡（快速）", 
                                    variable=selected_model_var, value='lite',
                                    bg='#ffffff', activebackground='#ffffff',
                                    font=("Microsoft YaHei", 10), fg='#f39c12',
                                    command=on_model_changed)
        lite_radio.pack(side=tk.LEFT, padx=5)
        
        # API Key 配置（用户需要填入）
        api_key_row = tk.Frame(ai_model_frame, bg='#ffffff')
        api_key_row.pack(fill=tk.X, pady=5)
        tk.Label(api_key_row, text="API Key:", width=12, font=("Microsoft YaHei", 10), bg='#ffffff').pack(side=tk.LEFT)
        api_key_entry = tk.Entry(api_key_row, width=40, font=("Microsoft YaHei", 10))
        api_key_entry.pack(side=tk.LEFT, padx=10)
        api_key_entry.insert(0, getattr(self, 'doubao_api_key', ''))
        
        # Base URL 配置（预填好，设为只读）
        base_url_row = tk.Frame(ai_model_frame, bg='#ffffff')
        base_url_row.pack(fill=tk.X, pady=5)
        tk.Label(base_url_row, text="Base URL:", width=12, font=("Microsoft YaHei", 10), bg='#ffffff').pack(side=tk.LEFT)
        base_url_var = tk.StringVar(value=default_base_url)
        base_url_entry = tk.Entry(base_url_row, textvariable=base_url_var, width=40, font=("Microsoft YaHei", 10), state='readonly')
        base_url_entry.pack(side=tk.LEFT, padx=10)
        
        # 模型名称配置（根据选择的模型自动填入，设为只读）
        model_name_row = tk.Frame(ai_model_frame, bg='#ffffff')
        model_name_row.pack(fill=tk.X, pady=5)
        tk.Label(model_name_row, text="模型名称:", width=12, font=("Microsoft YaHei", 10), bg='#ffffff').pack(side=tk.LEFT)
        model_name_var = tk.StringVar(value=default_pro_model)
        model_name_entry = tk.Entry(model_name_row, textvariable=model_name_var, width=40, font=("Microsoft YaHei", 10), state='readonly')
        model_name_entry.pack(side=tk.LEFT, padx=10)
        
        # 测试配置按钮
        test_btn_row = tk.Frame(ai_model_frame, bg='#ffffff')
        test_btn_row.pack(fill=tk.X, pady=(10, 5))
        tk.Button(test_btn_row, text="🔗 测试连接",
                 command=lambda: self.test_model_connection(
                     model_name_entry.get(),
                     base_url_entry.get(),
                     api_key_entry.get(),
                     test_log_text
                 ),
                 bg='#27ae60', fg='white', font=("Microsoft YaHei", 11, "bold"), 
                 relief='flat', padx=20, pady=8, cursor='hand2').pack(side=tk.LEFT, padx=(0, 10))
        
        # 测试日志区域
        tk.Label(ai_model_frame, text="📋 测试日志:", font=("Microsoft YaHei", 10, "bold"), bg='#ffffff', fg='#2c3e50').pack(anchor='w', pady=(10, 5))
        test_log_frame = tk.Frame(ai_model_frame, bg='#1e1e1e', relief='solid', bd=1)
        test_log_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
        test_log_text = tk.Text(test_log_frame, width=60, height=8, font=("Consolas", 10), 
                               bg='#1e1e1e', fg='#00ff00', wrap=tk.WORD, relief=tk.FLAT)
        test_log_text.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        # 滚动条
        test_scrollbar = ttk.Scrollbar(test_log_frame, command=test_log_text.yview)
        test_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        test_log_text.config(yscrollcommand=test_scrollbar.set)
        
        # 预设提示
        hint_label = tk.Label(ai_model_frame, 
                             text=f"💡 提示: 选择模型后自动填入模型名称，只需填入API Key即可测试",
                             font=("Microsoft YaHei", 9), bg='#ffffff', fg='#3498db')
        hint_label.pack(anchor='w', pady=(5, 0))

        # 按钮
        btn_frame = tk.Frame(container, bg='#ffffff')
        btn_frame.pack(pady=20)
        def save_settings():
            self.system_name = system_name_entry.get()
            self.company_name = company_name_entry.get()
            self.company_address = company_addr_entry.get()
            self.company_phone = company_phone_entry.get()
            self.bank_account = bank_account_entry.get()
            self.bank_name = bank_name_entry.get()
            self.paper_size = paper_size_var.get()
            self.logo_path = logo_path_var.get()
            self.stamp_path = stamp_path_var.get()
            self.quoter = quoter_entry.get()
            self.handler = handler_entry.get()
            # 保存AI模型设置 - 同步模型选择
            selected_model = selected_model_var.get()
            if selected_model == 'pro':
                self.doubao_model = '豆包2.0PRO'
                self.doubao_model_name = default_pro_model
            else:
                self.doubao_model = '豆包2.0LITE'
                self.doubao_model_name = default_lite_model
            self.doubao_base_url = base_url_entry.get()
            self.doubao_api_key = api_key_entry.get()
            self.save_settings_to_file()
            self.update_title()
            self.status_label.config(text="✔️ 设置已保存")
        tk.Button(btn_frame, text="保存设置", command=save_settings, width=15, height=2, bg='#27ae60', fg='white', relief='flat', pady=5).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="返回主页", command=self.show_welcome_page, width=15, height=2, bg='#95a5a6', fg='white', relief='flat', pady=5).pack(side=tk.LEFT, padx=10)

    # -------------------- AI模型测试连接 --------------------
    def test_model_connection(self, model_name, base_url, api_key, result_text_widget):
        """测试模型连接

        Args:
            model_name: 模型名称（从UI文本框读取）
            base_url: BaseURL（从UI文本框读取）
            api_key: API Key（从UI文本框读取）
            result_text_widget: 测试日志文本框组件
        """
        # 清除之前的日志
        result_text_widget.delete('1.0', tk.END)
        
        # 直接使用传入的参数值（从UI文本框读取）
        if not api_key:
            result_text_widget.insert(tk.END, "❌ 错误：请先在API Key字段中输入您的API密钥\n")
            return
        
        if not model_name:
            result_text_widget.insert(tk.END, "❌ 错误：请先选择或输入模型名称\n")
            return
        
        # 构建请求头
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # 测试模型调用 - 使用 OpenAI 兼容格式
        test_payload = {
            "model": model_name,
            "messages": [
                {
                    "role": "user",
                    "content": "你好，请回复'测试成功'"
                }
            ],
            "max_tokens": 100,
            "temperature": 0.1
        }
        
        import requests
        import json
        from datetime import datetime
        
        try:
            # 移除base_url末尾的斜杠
            base_url = base_url.strip().rstrip('/')
            # 根据base_url格式确定正确的端点
            if base_url.endswith('/v3'):
                # Base URL已经包含/v3，直接添加/chat/completions
                endpoint = f"{base_url}/chat/completions"
            else:
                # 否则添加/v1/chat/completions
                endpoint = f"{base_url}/v1/chat/completions"
            
            result_text_widget.insert(tk.END, f"🔄 正在测试连接...\n")
            result_text_widget.insert(tk.END, f"端点: {endpoint}\n")
            result_text_widget.insert(tk.END, f"模型: {model_name}\n")
            result_text_widget.update_idletasks()
            
            # 发送请求
            resp = requests.post(endpoint, 
                               json=test_payload, 
                               headers=headers, 
                               timeout=15)
            
            if resp.status_code == 200:
                response_data = resp.json()
                result_text_widget.insert(tk.END, f"\n✅ {model_name} 连接测试成功！\n")
                result_text_widget.insert(tk.END, f"响应时间: {resp.elapsed.total_seconds()*1000:.0f}ms\n")
                
                # 显示模型响应
                if 'choices' in response_data and len(response_data['choices']) > 0:
                    content = response_data['choices'][0].get('message', {}).get('content', '')
                    result_text_widget.insert(tk.END, f"\n📝 模型回复: {content}\n")
                
            else:
                error_info = resp.text
                try:
                    error_json = resp.json()
                    if 'error' in error_json:
                        error_info = error_json['error'].get('message', error_json['error'].get('code', ''))
                    else:
                        error_info = json.dumps(error_json, indent=2, ensure_ascii=False)
                except:
                    pass
                
                result_text_widget.insert(tk.END, f"\n❌ 连接失败 (状态码: {resp.status_code})\n")
                result_text_widget.insert(tk.END, f"错误信息: {error_info}\n")
                result_text_widget.insert(tk.END, f"\n💡 提示: 请检查:\n")
                result_text_widget.insert(tk.END, f"  1. Base URL 是否正确\n")
                result_text_widget.insert(tk.END, f"  2. 模型名称是否正确\n")
                result_text_widget.insert(tk.END, f"  3. API Key 是否有效\n")
                
        except requests.exceptions.Timeout:
            result_text_widget.insert(tk.END, "\n❌ 连接超时：请求超过15秒未响应\n")
            result_text_widget.insert(tk.END, "请检查网络连接或API端点是否正确\n")
        except requests.exceptions.ConnectionError as e:
            result_text_widget.insert(tk.END, f"\n❌ 连接错误：无法连接到服务器\n")
            result_text_widget.insert(tk.END, f"错误详情: {str(e)}\n")
            result_text_widget.insert(tk.END, f"\n💡 提示: 请检查 Base URL 是否正确\n")
        except Exception as e:
            result_text_widget.insert(tk.END, f"\n❌ 请求异常：{str(e)}\n")
        
        # 添加时间戳
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        result_text_widget.insert(tk.END, f"\n测试时间: {timestamp}\n")

    # -------------------- 帮助与关于 --------------------
    def show_about(self):
        about_text = """账单管理系统

版本: 2.0

功能特点:
- 对账单管理
- 报价单管理
- 付款跟踪
- 客户管理
- 数据导出
- 图片OCR识别
- AI模型集成

技术栈: Python + Tkinter

© 2025 All Rights Reserved"""
        messagebox.showinfo("关于", about_text)

    def show_help(self):
        help_text = """使用说明:

1. 创建对账单:
   点击"创建对账单"按钮或菜单栏中的"对账单 -> 创建对账单"

2. 管理报价单:
   点击"创建新报价"按钮或菜单栏中的"报价管理 -> 创建新报价"

3. 付款管理
   在菜单栏"付款管理"中可以查看付款列表、添加付款记录、导出付款

4. 客户管理
   在菜单栏"客户管理"中可以管理客户信息

5. 数据存储
   所有数据自动保存在JSON文件中，方便备份和迁移

常见问题:
- 如何导出数据? 在各列表窗口中有导出按钮
- 如何添加客户? 客户管理 -> 添加客户
- 如何查看历史账单? 对账单 -> 查看对账单列表"""
        messagebox.showinfo("使用帮助", help_text)

    # -------------------- 运行 --------------------
    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = InvoiceManager()
    app.run()